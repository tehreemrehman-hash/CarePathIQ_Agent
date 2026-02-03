import streamlit as st
# Version info sidebar caption (admin-only)
import streamlit.components.v1 as components
import os
import json
import pandas as pd
import altair as alt
import urllib.request
import urllib.parse
import re
import time
import base64
from io import BytesIO
import datetime
from datetime import date, timedelta
import copy
import xml.etree.ElementTree as ET
from contextlib import contextmanager
import requests
import hashlib
import textwrap
from google import genai
from google.genai import types

# Import Gemini function declarations and helpers
from gemini_functions import (
    PRIMARY_MODEL, MODEL_CASCADE,
    GENERATE_PATHWAY_NODES, DEFINE_PATHWAY_SCOPE, CREATE_IHI_CHARTER,
    GRADE_EVIDENCE, ANALYZE_HEURISTICS, APPLY_HEURISTICS,
    GENERATE_BETA_TEST_SCENARIOS, ANALYZE_AUDIENCE,
    get_tool, get_tools, get_generation_config, extract_function_call_result,
    DEFAULT_THINKING_CONFIG, COMPLEX_THINKING_CONFIG, LIGHT_THINKING_CONFIG
)

# Clinical pathway generation modules
try:
    from pathway_generator import (
        PathwayGenerator, Order, EvidenceBasedAddition,
        DispositionCriteria, DispositionType,
        create_mermaid_from_nodes, create_dot_from_nodes,
        export_pathway_markdown
    )
    PATHWAY_GENERATOR_AVAILABLE = True
except ImportError:
    PATHWAY_GENERATOR_AVAILABLE = False

try:
    from llm_prompt_templates import (
        format_comprehensive_prompt, format_refinement_prompt,
        build_evidence_context, build_pathway_summary
    )
    LLM_TEMPLATES_AVAILABLE = True
except ImportError:
    LLM_TEMPLATES_AVAILABLE = False

def _get_query_param(name: str):
    try:
        # Streamlit >= 1.33
        if hasattr(st, "query_params"):
            val = st.query_params.get(name)
        else:
            val = st.experimental_get_query_params().get(name)
        if isinstance(val, list):
            return val[0] if val else None
        return val
    except Exception:
        return None

def is_admin():
    # Env flag still enables admin mode for server-side checks
    if os.environ.get("CPQ_SHOW_VERSION", "").lower() in ("1", "true", "yes", "y"):
        return True
    # Optional shared code via env or secrets + query param `?admin=CODE`
    code = os.environ.get("CPQ_ADMIN_CODE", "")
    try:
        code = code or st.secrets.get("ADMIN_CODE", "")
    except Exception:
        pass
    if code:
        return (_get_query_param("admin") == code)
    return False

def is_debug():
    # Enable with env `CPQ_DEBUG` or query param `?debug=1`
    if os.environ.get("CPQ_DEBUG", "").lower() in ("1", "true", "yes", "y"):
        return True
    val = _get_query_param("debug")
    return str(val).lower() in ("1", "true", "yes", "y")

def debug_log(msg: str):
    try:
        if is_debug():
            st.sidebar.caption(str(msg))
    except Exception:
        pass


# --- GOOGLE GENAI CLIENT HELPER ---
# Using official google-genai SDK per https://ai.google.dev/gemini-api/docs/quickstart
def get_genai_client():
    return st.session_state.get("genai_client")

# Column helper that prefers top alignment but gracefully falls back
def columns_top(spec, **kwargs):
    try:
        # Streamlit >= 1.30 supports vertical_alignment
        return st.columns(spec, vertical_alignment="top", **kwargs)
    except TypeError:
        # Older Streamlit versions without vertical_alignment
        return st.columns(spec, **kwargs)

# Unified status UI for AI tasks
@contextmanager
def ai_activity(label="Working with the AI agent…"):
    with st.status(label, expanded=False) as status:
        try:
            yield status
            status.update(label="Ready!", state="complete")
        except Exception:
            status.update(label="There was a problem completing this step.", state="error")
            st.error("Please try again or adjust your input.")
            return

def regenerate_nodes_with_refinement(nodes, refine_text, heuristics_data=None):
    """Regenerate Phase 3 nodes based on user refinement notes and optional heuristics context."""
    refine_text = (refine_text or "").strip()
    if not refine_text:
        return None

    cond = st.session_state.data['phase1'].get('condition') or "Pathway"
    setting = st.session_state.data['phase1'].get('setting') or "care setting"
    evidence_list = st.session_state.data['phase2'].get('evidence', [])

    ev_context = "\n".join([
        f"- PMID {e['id']}: {e['title']} | Abstract: {e.get('abstract', 'N/A')[:200]}"
        for e in evidence_list[:20]
    ])

    heuristics_summary = ""
    if heuristics_data:
        bullet_lines = []
        for key, val in heuristics_data.items():
            try:
                label = val.get('label', key)
                recs = val.get('recommendations', [])
                if recs:
                    bullet_lines.append(f"- {label}: {recs[0]}")
            except Exception:
                continue
        if bullet_lines:
            heuristics_summary = "\nHeuristics guidance (PRESERVE clinical complexity while applying):\n" + "\n".join(bullet_lines[:5])

    prompt = f"""
    Act as a CLINICAL DECISION SCIENTIST. Refine the EXISTING pathway based on the user's request.

    CRITICAL: Apply refinements while PRESERVING and potentially ENHANCING clinical complexity.
    
    Current pathway for {cond} in {setting}:
    {json.dumps(nodes, indent=2)}

    Available Evidence:
    {ev_context}

    User's refinement request: "{refine_text}"
    {heuristics_summary}

    MANDATORY PRESERVATION RULES:
    
    1. MAINTAIN DECISION SCIENCE FRAMEWORK:
       - Keep CGT/Ad/it principles and Medical Decision Analysis structure intact
       - Preserve actionable clinical notes (red flags, thresholds, monitoring parameters)
       - Maintain evidence-based reasoning (cite PMIDs)
    
    2. PRESERVE DECISION DIVERGENCE (Minimum Separation Rule):
       - Do NOT collapse multiple branches into linear sequences
       - Each branch from a Decision must have 2-3 unique steps BEFORE any convergence
       - Eventual convergence to shared End nodes or late Process steps is OK after meaningful divergence
       - NEVER allow immediate reconvergence (both branches pointing to same next node)
       - If refining convergence points, ensure at least 3 steps of unique pathway before merge
    
    3. PRESERVE CLINICAL COVERAGE:
       - All 4 stages must remain: Initial Evaluation, Diagnosis/Treatment, Re-evaluation, Final Disposition
       - Do NOT remove edge cases or special population considerations
       - EXPAND specificity when refining (apply patterns below, adapted to the clinical condition):
         * Validated clinical scores relevant to this condition with specific numerical thresholds
         * Age-adjusted or population-specific calculations where established in literature
         * Special population screening (pregnancy before radiation/teratogens, renal function before contrast/drugs, contraindications)
         * Medication specificity: Brand AND generic names, exact dosing, timing, route, location IN THE NODE LABEL
         * CRITICAL: Medication administration is a CLINICAL ACTION—create Process nodes with doses visible in flowchart
         * Example node label: "Administer vancomycin 15-20 mg/kg IV q8-12h, adjust for CrCl <30"
         * Example node notes: "Red flags: fever >38.5°C, rigors, hypotension. Monitor: trough levels before 4th dose."
         * Insurance/cost considerations: "Ensure Rx covered; provide assistance program links if available"
         * Resource contingencies: "If [preferred test] NOT available → [Alternative approach]"
         * Follow-up timing: "[Provider type] within [timeframe]", "Virtual care if [provider] unavailable"
    
    4. ENHANCE DEPTH, NOT REDUCE:
       - When applying refinements, consider adding detail (more nodes, more branches)
       - If user asks to "simplify," interpret as "make more understandable" (clearer labels, better organization)
       - NOT "remove clinical branches" or reduce node count
       - Prioritize clinical completeness over arbitrary node limits
    
    5. MAINTAIN DAG STRUCTURE:
       - No cycles, backward loops, or reconvergent branches
       - Escalation moves forward (ED → ICU, not back)
       - All paths must terminate in explicit End nodes
    
    6. SOPHISTICATED CLINICAL REALISM (Apply when refining relevant sections):
       - Risk stratification: Use validated scores BEFORE diagnostic tests (not generic "assess risk")
       - Contraindication checks: Explicit "Absolute contraindications?" decision nodes BEFORE treatment
       - Resource availability: "Preferred imaging available?" branches with alternatives specified
       - Educational content: Note hyperlink candidates (score calculators, drug info, evidence citations)
       - Disposition specificity: Never vague "discharge" - specify follow-up provider, timing, virtual alternatives
    
    OUTPUT: Complete revised JSON array of nodes with fields: type, label, evidence, (optional) notes
    Rules:
    - type: "Start" | "Decision" | "Process" | "End"
    - First node: type "Start", label "patient present to {setting} with {cond}"
    - NO artificial node count limit—maintain complexity needed for clinical accuracy
    - End nodes must be TERMINAL single outcomes (no "or" phrasing)
    - Consecutive Decision nodes are allowed and encouraged for true clinical branching
    - Notes field: Actionable clinical details like red flag signs, specific thresholds, monitoring parameters
    - Evidence citations (PMIDs) on clinically important steps
    - Apply sophisticated patterns above to make pathway immediately implementable by clinicians
    """

    # Use native function calling for structured output (with json_mode fallback)
    result = get_gemini_response(
        prompt, 
        function_declaration=GENERATE_PATHWAY_NODES,
        thinking_budget=2048  # Complex pathway generation needs more reasoning
    )
    
    # Extract nodes from function call result or fall back to legacy parsing
    if isinstance(result, dict) and 'arguments' in result:
        return result['arguments'].get('nodes', [])
    elif isinstance(result, dict) and 'nodes' in result:
        return result.get('nodes', [])
    elif isinstance(result, list):
        return result
    
    # Fallback to json_mode if function calling didn't work
    return get_gemini_response(prompt, json_mode=True)

# --- LIBRARY HANDLING ---
try:
    from docx import Document
    from docx.shared import Inches as DocxInches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import matplotlib.pyplot as plt
    import matplotlib.dates as mdates
    import graphviz  # for diagram exports (SVG/PNG/DOT)
except ImportError:
    # Safe fallback to prevent crash if libraries are missing
    Document = None
    plt = None
    mdates = None
    graphviz = None

# ==========================================
# 1. PAGE CONFIGURATION & STYLING
# ==========================================
st.set_page_config(
    page_title="CarePathIQ AI Agent",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Hide detailed Python tracebacks from end users
try:
    st.set_option('client.showErrorDetails', False)
except Exception:
    pass

# --- CUSTOM CSS ---
# Force-clear Streamlit caches once per user session to avoid stale views
if "cleared_cache_once" not in st.session_state:
    try:
        st.cache_data.clear()
    except Exception:
        pass
    try:
        st.cache_resource.clear()
    except Exception:
        pass
    st.session_state.cleared_cache_once = True

st.markdown("""
<style>
    /* HIDE HEADER DEPLOYMENT LINKS BUT KEEP CONTENT ANCHOR LINKS */
    [data-testid="stHeaderAction"] { display: none !important; visibility: hidden !important; opacity: 0 !important; }
    
    /* Make anchor links visible for content headings */
    .stMarkdown h1 a.anchor-link, 
    .stMarkdown h2 a.anchor-link, 
    .stMarkdown h3 a.anchor-link { 
        display: inline-block !important; 
        opacity: 1 !important;
        visibility: visible !important;
        height: auto !important; 
        width: auto !important; 
        margin-left: 0.5rem !important;
    }
    h1 > a.anchor-link, h2 > a.anchor-link, h3 > a.anchor-link { 
        display: inline-block !important; 
        opacity: 1 !important;
        visibility: visible !important;
    }
    
    /* BUTTONS */
    /* Style secondary (non-active) buttons as brown. Avoid overriding primary. */
    button[kind="secondary"],
    div.stButton > button:not([kind="primary"]),
    div[data-testid="stButton"] > button:not([kind="primary"]) {
        background-color: #5D4037 !important; 
        color: white !important;
        border: 1px solid #5D4037 !important;
        border-radius: 5px !important;
        font-weight: 600 !important;
    }
    button[kind="secondary"]:hover,
    div.stButton > button:not([kind="primary"]):hover,
    div[data-testid="stButton"] > button:not([kind="primary"]):hover {
        background-color: #3E2723 !important; 
        border-color: #3E2723 !important;
        color: white !important;
    }

    /* FORM SUBMIT BUTTONS (e.g., "Regenerate" inside forms) */
    /* Streamlit renders these differently from st.button; target explicitly */
    div[data-testid="stFormSubmitButton"] > button,
    form[data-testid="stForm"] button[type="submit"] {
        background-color: #5D4037 !important; 
        color: white !important;
        border: 1px solid #5D4037 !important;
        border-radius: 5px !important;
        font-weight: 600 !important;
    }
    div[data-testid="stFormSubmitButton"] > button:hover,
    form[data-testid="stForm"] button[type="submit"]:hover {
        background-color: #3E2723 !important; 
        border-color: #3E2723 !important;
        color: white !important;
    }
    
    /* PRIMARY BUTTONS (Current Phase) */
    div.stButton > button[kind="primary"],
    button[kind="primary"] {
        background-color: #FFB0C9 !important; 
        color: black !important;
        border: 1px solid black !important;
        border-radius: 5px !important;
        font-weight: 600 !important;
    }
    div.stButton > button[kind="primary"]:hover,
    button[kind="primary"]:hover {
        background-color: #FF9BB8 !important; 
        border-color: black !important;
        color: black !important;
    }
    div.stButton > button:disabled {
        background-color: #eee !important;
        color: #999 !important;
        border: 1px solid #ccc !important;
    }

    /* DOWNLOAD BUTTONS */
    div.stDownloadButton > button {
        background-color: #5D4037 !important; 
        color: white !important;
        border: none !important;
        border-radius: 5px !important;
    }
    div.stDownloadButton > button:hover {
        background-color: #3E2723 !important;
    }

    /* LINK BUTTONS */
    a[kind="secondary"] {
        background-color: #5D4037 !important;
        border-color: #5D4037 !important;
        color: white !important;
        text-decoration: none;
        padding: 0.5rem 1rem;
        border-radius: 5px;
    }
    a[kind="secondary"]:hover {
        background-color: #3E2723 !important;
        border-color: #3E2723 !important;
        color: white !important;
    }

    /* SIDEBAR BUTTONS */
    section[data-testid="stSidebar"] div.stButton > button {
        background-color: #A9EED1 !important; 
        color: #5D4037 !important;
        border: none !important;
        font-weight: bold;
    }
    section[data-testid="stSidebar"] div.stButton > button:hover {
        background-color: #8FD9BC !important; 
        color: #3E2723 !important;
    }
    
    /* RADIO BUTTONS */
    div[role="radiogroup"] label > div:first-child {
        background-color: white !important;
        border: 2px solid #5D4037 !important;
        border-radius: 50% !important;
        width: 20px !important;
        height: 20px !important;
        margin-right: 8px !important;
        box-shadow: 0 1px 2px rgba(0,0,0,0.05);
    }
    div[role="radiogroup"] label[data-checked="true"] > div:first-child {
        background-color: #FFB0C9 !important;
        border: 2px solid black !important;
        box-shadow: 0 2px 4px rgba(255,176,201,0.35);
    }
    
    /* SPINNER */
    .stSpinner > div {
        border-top-color: #5D4037 !important;
    }
    
    /* EXPANDERS */
    div[data-testid="stExpander"] details summary {
        background-color: #5D4037 !important;
        color: white !important;
        border-radius: 5px;
        margin-bottom: 5px;
    }
    div[data-testid="stExpander"] details summary:hover {
        color: #A9EED1 !important;
    }
    div[data-testid="stExpander"] details summary svg {
        color: white !important;
    }

    /* Sidebar chat expander matches landing banner */
    section[data-testid="stSidebar"] div[data-testid="stExpander"]:first-of-type details summary {
        background-color: #FFB0C9 !important;
        color: black !important;
        border: 1px solid black !important;
        box-shadow: 0 3px 10px rgba(0,0,0,0.08);
        font-weight: 700 !important;
    }
    section[data-testid="stSidebar"] div[data-testid="stExpander"]:first-of-type details summary:hover {
        background-color: #FF9BB8 !important;
        color: #3E2723 !important;
    }
    section[data-testid="stSidebar"] div[data-testid="stExpander"]:first-of-type details summary svg {
        color: black !important;
    }

    /* HEADERS */
    h1, h2, h3 { color: #00695C; font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif; }
    
    /* TOOLTIPS */
    div[data-testid="stTooltipContent"] {
        background-color: white !important;
        color: #333 !important;
        border: 1px solid #ddd !important;
    }
    /* Help icon spacing right after labels (e.g., GRADE) */
    [data-testid="stTooltipIcon"] {
        margin-left: 4px !important;
        display: inline-block !important;
        vertical-align: middle !important;
    }
    /* Heuristic title tooltips */
    .heuristic-title {
        cursor: help;
        font-weight: bold;
        color: #00695C;
        text-decoration: underline dotted;
        font-size: 1.05em;
    }
    /* Inline info icon for section headers */
    .tooltip-info {
        display: inline-block;
        margin-left: 6px;
        color: #5D4037;
        background: #A9EED1;
        border: 1px solid #5D4037;
        border-radius: 50%;
        width: 18px;
        height: 18px;
        line-height: 16px;
        text-align: center;
        font-size: 12px;
        cursor: help;
    }
    
    /* GRADE FILTER MULTISELECT */
    div[data-testid="stMultiSelect"] label:has(+ div[data-baseweb="select"]) {
        background-color: #FFB0C9 !important;
        color: black !important;
        padding: 8px 12px !important;
        border-radius: 5px !important;
        border: 1px solid black !important;
        display: inline-block !important;
        margin-bottom: 8px !important;
        font-weight: 500 !important;
    }

    /* IMPROVE BUTTON AND TEXT VERTICAL ALIGNMENT */
    /* Top-align items in columns to keep headers aligned horizontally */
    div[data-testid="stColumn"] {
        display: flex !important;
        align-items: flex-start !important;
    }

    /* For horizontal button rows (e.g., bottom navigation), center-align contents */
    [data-testid="stHorizontalBlock"] div[data-testid="stColumn"] {
        align-items: center !important;
    }
    /* Ensure the inner vertical block centers content within horizontal rows */
    [data-testid="stHorizontalBlock"] div[data-testid="stColumn"] > div > div[data-testid="stVerticalBlock"] {
        justify-content: center !important;
        align-items: center !important;
    }
    
    button, input, select, textarea, label {
        vertical-align: middle !important;
    }
    
    /* PHASE 2 EXPORT SECTION ALIGNMENT */
    div[data-testid="stColumn"] > div > div[data-testid="stVerticalBlock"] {
        display: flex !important;
        flex-direction: column !important;
        justify-content: flex-start !important;
    }
    
    /* SUBHEADER AND SELECTBOX ALIGNMENT */
    div[data-testid="stColumn"] h3 {
        margin-bottom: 0 !important;
        line-height: 1.2 !important;
    }
    /* Ensure Streamlit subheaders (h2) sit close to content for alignment */
    div[data-testid="stColumn"] h2 {
        margin-bottom: 0.35rem !important;
        line-height: 1.2 !important;
    }
    
    div[data-testid="stColumn"] div[data-baseweb="select"] {
        margin-top: 0 !important;
    }
    
    /* RADIO BUTTON LABEL ALIGNMENT */
    div[role="radiogroup"] {
        display: flex !important;
        align-items: center !important;
        gap: 1.5rem !important;
        width: 50% !important;
    }
    
    div[role="radiogroup"] label {
        display: flex !important;
        align-items: center !important;
        margin-bottom: 0 !important;
        margin-right: 0 !important;
        line-height: 1 !important;
        white-space: nowrap !important;
        gap: 0.5rem !important;
    }
    
    div[role="radiogroup"] label > span {
        vertical-align: middle !important;
    }
    
    div[role="radiogroup"] label > div {
        display: flex !important;
        align-items: center !important;
        justify-content: center !important;
    }
    
    /* DOWNLOAD BUTTON CONSISTENT HEIGHT */
    div.stDownloadButton > button,
    div.stButton > button {
        min-height: 38px !important;
        display: flex !important;
        align-items: center !important;
        justify-content: center !important;
    }

    /* Global link-button styling to match Streamlit buttons */
    .cpq-link-button {
        display: flex; align-items: center; justify-content: center;
        width: 100%; padding: 10px 14px; border-radius: 5px;
        background-color: #5D4037; color: #fff; text-decoration: none;
        border: 1px solid #5D4037; font-weight: 600; height: 42px;
    }
    .cpq-link-button:hover { background-color: #3E2723; border-color: #3E2723; color: #fff; }

    /* FILE UPLOADER & TEXTAREA — compact spacing */
    div[data-testid="stFileUploader"] section {
        padding-top: 0 !important;
        padding-bottom: 0 !important;
    }
    div[data-testid="stFileUploaderDropzone"] {
        padding: 12px !important;
        border-radius: 8px !important;
    }
    textarea {
        line-height: 1.35 !important;
    }

    /* PHASE 2 — highlight rows marked as new via the first checkbox column */
    div[data-testid="stDataFrame"] tbody tr:has(input[type="checkbox"]:checked) td,
    div[data-testid="stTable"] tbody tr:has(input[type="checkbox"]:checked) td {
        background-color: #FFB0C9 !important;
    }

    /* DATA TABLES — enable text wrapping for easier scanning */
    div[data-testid="stDataFrame"] table,
    div[data-testid="stTable"] table {
        table-layout: fixed !important;
    }
    div[data-testid="stDataFrame"] td, div[data-testid="stDataFrame"] th,
    div[data-testid="stTable"] td, div[data-testid="stTable"] th {
        white-space: normal !important;
        word-break: break-word !important;
    }
    
    /* BOTTOM NAVIGATION */
    /* Hide only Streamlit's default footer, not custom navigation */
    footer[data-testid="stBottom"] {
        visibility: hidden;
    }
    
    /* ENHANCED PHASE NAVIGATION */
    .phase-nav-container {
        background: linear-gradient(to right, #FFB0C9 0%, #A9EED1 100%);
        padding: 10px 20px;
        border-radius: 10px;
        margin-bottom: 20px;
        box-shadow: 0 2px 8px rgba(0,0,0,0.1);
    }
    
    .phase-progress-bar {
        height: 6px;
        background: #e0e0e0;
        border-radius: 3px;
        overflow: hidden;
        margin-bottom: 15px;
    }
    
    .phase-progress-fill {
        height: 100%;
        background: linear-gradient(to right, #5D4037, #00695C);
        transition: width 0.5s ease;
    }
    
    .phase-nav-label {
        font-size: 0.8rem;
        color: #5D4037;
        font-weight: 600;
        text-align: center;
        margin-bottom: 5px;
    }
    
    /* Make navigation buttons more compact */
    [data-testid="stHorizontalBlock"] > div[data-testid="stColumn"] button {
        font-size: 0.9rem !important;
        padding: 10px 8px !important;
        white-space: nowrap !important;
        overflow: hidden !important;
        text-overflow: ellipsis !important;
        line-height: 1.3 !important;
        min-height: 42px !important;
        max-height: 42px !important;
        height: 42px !important;
    }
    
    /* Ensure all button text stays horizontal */
    [data-testid="stHorizontalBlock"] button div,
    [data-testid="stHorizontalBlock"] button p,
    [data-testid="stHorizontalBlock"] button span {
        writing-mode: horizontal-tb !important;
        text-orientation: mixed !important;
        transform: none !important;
        display: inline !important;
    }

    /* Responsive: stack columns and make paired controls full-width on narrow screens */
    @media (max-width: 560px) {
        /* Stack all Streamlit columns */
        div[data-testid="stColumn"] {
            width: 100% !important;
            flex: 0 0 100% !important;
        }
        /* Make link-button and download button fill width */
        .cpq-link-button { width: 100% !important; }
        div.stDownloadButton > button { width: 100% !important; }
    }
</style>
""", unsafe_allow_html=True)

# Integrating phase 5 helper functions into the Streamlit app

# Importing necessary functions from phase5_helpers
from phase5_helpers import CAREPATHIQ_COLORS, CAREPATHIQ_FOOTER, SHARED_CSS, ensure_carepathiq_branding

# CONSTANTS
COPYRIGHT_MD = "\n\n---\n**© 2024 CarePathIQ by Tehreem Rehman.** Licensed under [CC BY-SA 4.0](https://creativecommons.org/licenses/by-sa/4.0/)."
COPYRIGHT_HTML_FOOTER = """
<div style="text-align: center; margin-top: 40px; padding-top: 20px; border-top: 1px solid #ddd; font-size: 0.85em; color: #666;">
    <p>
        <a href="https://www.carepathiq.org" target="_blank" style="text-decoration:none; color:#4a4a4a; font-weight:bold;">CarePathIQ</a> 
        © 2024 by 
        <a href="https://www.tehreemrehman.com" target="_blank" style="text-decoration:none; color:#4a4a4a; font-weight:bold;">Tehreem Rehman</a> 
        is licensed under 
        <a href="https://creativecommons.org/licenses/by-sa/4.0/" target="_blank" style="text-decoration:none; color:#4a4a4a;">CC BY-SA 4.0</a>
        <img src="https://mirrors.creativecommons.org/presskit/icons/cc.svg" alt="" style="max-width: 1em;max-height:1em;margin-left: .2em;">
        <img src="https://mirrors.creativecommons.org/presskit/icons/by.svg" alt="" style="max-width: 1em;max-height:1em;margin-left: .2em;">
        <img src="https://mirrors.creativecommons.org/presskit/icons/sa.svg" alt="" style="max-width: 1em;max-height:1em;margin-left: .2em;">
    </p>
</div>
"""
HEURISTIC_DEFS = {
    "H1": "Visibility of system status: The design should always keep users informed about what is going on.",
    "H2": "Match between system and real world: Speak the users' language, avoiding jargon.",
    "H3": "User control and freedom: Provide clearly marked 'emergency exits' to leave unwanted states.",
    "H4": "Consistency and standards: Users shouldn't have to wonder if different words mean the same thing.",
    "H5": "Error prevention: Good error messages are important, but preventing problems is better.",
    "H6": "Recognition rather than recall: Minimize memory load; making elements and options visible.",
    "H7": "Flexibility and efficiency of use: Accelerators for experts while remaining usable for novices.",
    "H8": "Aesthetic and minimalist design: Interfaces should not contain irrelevant information.",
    "H9": "Help users recognize, diagnose, and recover from errors: Error messages in plain language.",
    "H10": "Help and documentation: Provide concise, concrete documentation focused on user tasks."
}

# Heuristics categorized by applicability to clinical pathways
HEURISTIC_CATEGORIES = {
    "pathway_actionable": {
        "H2": "Language clarity (replace medical jargon with patient-friendly terms where appropriate)",
        "H4": "Consistency (standardize terminology and node types across pathway)",
        "H5": "Error prevention (add critical alerts, validation rules, and edge case handling)",
        "H9": "Error recovery (surface critical checks earlier; add recovery steps in the flow)"
    },
    "ui_design_only": {
        "H1": "Status visibility (implement progress indicators and highlighting in the interface)",
        "H3": "User control (add escape routes and undo/skip options in UI)",
        "H6": "Recognition not recall (use visual icons and clear labels instead of hidden menus)",
        "H7": "Efficiency accelerators (add keyboard shortcuts and quick actions for power users)",
        "H8": "Minimalist design (remove clutter and non-essential information from interface)",
        "H10": "Help & docs (provide in-app tooltips, FAQs, and guided walkthroughs)"
    }
}

ROLE_COLORS = {
    "Physician": "#E3F2FD",
    "Doctor": "#E3F2FD",
    "MD": "#E3F2FD",
    "Nurse": "#E8F5E9",
    "RN": "#E8F5E9",
    "Pharmacist": "#F3E5F5",
    "PharmD": "#F3E5F5",
    "Patient": "#FFF3E0",
    "Care Coordinator": "#FFF8E1",
    "Process": "#FFFDE7",
}
PHASES = [
    "Define Scope & Charter",
    "Appraise Evidence",
    "Build Decision Tree",
    "Design User Interface",
    "Operationalize & Deploy"
]

PROVIDER_OPTIONS = {
    "google": "Google Forms (user account)",
    "html": "HTML Preview (no submission)",
}

# ==========================================
# 2. HELPER FUNCTIONS
# ==========================================

import secrets
import urllib.parse

# --- NAVIGATION CONTROLLER ---
def change_phase(new_phase):
    st.session_state.current_phase_label = new_phase
    st.session_state.top_nav_radio = new_phase  # Sync radio button value

def render_bottom_navigation():
    """Renders Previous/Next buttons at the bottom of the page with phase indicators."""
    if "current_phase_label" in st.session_state and st.session_state.current_phase_label in PHASES:
        current_idx = PHASES.index(st.session_state.current_phase_label)
        st.divider()
        # Add a small spacer to keep bottom nav comfortably separated from content above
        st.markdown("<div style='height: 10px'></div>", unsafe_allow_html=True)
        
        try:
            # Try Streamlit >= 1.30 with vertical_alignment
            col_prev, col_middle, col_next = st.columns([1, 1, 1], vertical_alignment="center")
        except TypeError:
            # Fallback for older Streamlit versions
            col_prev, col_middle, col_next = st.columns([1, 1, 1])
        
        if current_idx > 0:
            prev_phase = PHASES[current_idx - 1]
            with col_prev:
                if st.button(f"← {prev_phase}", key=f"bottom_prev_{current_idx}", type="secondary", use_container_width=True):
                    st.session_state.current_phase_label = prev_phase
                    st.rerun()
        else:
            with col_prev:
                st.write("")  # Empty space for alignment
        
        with col_middle:
            # Simple phase counter
            st.markdown(
                f"<div style='text-align: center; padding: 10px; color: #5D4037; font-weight: bold;'>Phase {current_idx + 1} of {len(PHASES)}</div>",
                unsafe_allow_html=True
            )
        
        if current_idx < len(PHASES) - 1:
            next_phase = PHASES[current_idx + 1]
            with col_next:
                if st.button(f"{next_phase} →", key=f"bottom_next_{current_idx}", type="primary", use_container_width=True):
                    st.session_state.current_phase_label = next_phase
                    st.rerun()
        # Always render brand/licensing footer before any phase stop()
        try:
            st.markdown(COPYRIGHT_HTML_FOOTER, unsafe_allow_html=True)
        except Exception:
            pass

def calculate_granular_progress():
    """
    Calculate progress across all 5 phases with balanced weighting.
    Each phase = 20% of total progress (fair distribution).
    Within each phase, fields are weighted equally.
    Returns: float between 0.0 and 1.0
    """
    if 'data' not in st.session_state:
        return 0.0
    
    data = st.session_state.data
    phase_progress = {}
    
    # --- PHASE 1: 20% (6 fields, each 1/6 of phase 1) ---
    p1 = data.get('phase1', {})
    p1_fields = ['condition', 'setting', 'inclusion', 'exclusion', 'problem', 'objectives']
    p1_completed = sum(1 for k in p1_fields if p1.get(k))
    phase_progress['p1'] = p1_completed / len(p1_fields)
    
    # --- PHASE 2: 20% (2 fields, each 1/2 of phase 2) ---
    p2 = data.get('phase2', {})
    p2_fields = ['mesh_query', 'evidence']
    p2_completed = sum(1 for k in p2_fields if p2.get(k))
    phase_progress['p2'] = p2_completed / len(p2_fields)
    
    # --- PHASE 3: 20% (1 field, worth full phase 3) ---
    p3 = data.get('phase3', {})
    phase_progress['p3'] = 1.0 if p3.get('nodes') else 0.0
    
    # --- PHASE 4: 20% (1 field, worth full phase 4) ---
    p4 = data.get('phase4', {})
    phase_progress['p4'] = 1.0 if p4.get('heuristics_data') else 0.0
    
    # --- PHASE 5: 20% (3 fields, each 1/3 of phase 5) ---
    p5 = data.get('phase5', {})
    p5_fields = ['beta_html', 'expert_html', 'edu_html']
    p5_completed = sum(1 for k in p5_fields if p5.get(k))
    phase_progress['p5'] = p5_completed / len(p5_fields)
    
    # Each phase worth 20% of total
    overall_progress = sum(phase_progress.values()) / 5
    return min(1.0, overall_progress)

def styled_info(text):
    formatted_text = text.replace("Tip:", "<b>Tip:</b>")
    st.markdown(f"""
    <div style="background-color: #FFB0C9; color: black; padding: 10px; border-radius: 5px; border: 1px solid black; margin-bottom: 10px;">
        {formatted_text}
    </div>""", unsafe_allow_html=True)

def upload_and_review_file(uploaded_file, phase_key: str, context: str = ""):
    """
    Upload a file to Gemini, auto-review it, and return markdown summary + file URI.
    Args:
        uploaded_file: Streamlit UploadedFile object
        phase_key: Unique key for session state (e.g., 'p1_refine')
        context: Optional context about what the file is for (e.g., 'clinical pathway')
    Returns:
        dict with 'review' (markdown), 'file_uri', and 'filename'
    """
    if not uploaded_file:
        return None

    client = get_genai_client()
    if not client:
        st.error("API connection error. Cannot upload file.")
        return None

    try:
        # Read file bytes
        file_bytes = uploaded_file.read()
        
        # Convert DOCX to markdown, then upload as .md file
        if uploaded_file.name.lower().endswith('.docx'):
            try:
                from io import BytesIO
                from docx import Document
                doc = Document(BytesIO(file_bytes))
                
                # Convert to markdown with structure preserved
                md_lines = []
                for para in doc.paragraphs:
                    text = para.text.strip()
                    if not text:
                        continue
                    # Detect headings by style
                    style_name = para.style.name.lower() if para.style else ""
                    if 'heading 1' in style_name:
                        md_lines.append(f"# {text}")
                    elif 'heading 2' in style_name:
                        md_lines.append(f"## {text}")
                    elif 'heading 3' in style_name:
                        md_lines.append(f"### {text}")
                    elif 'title' in style_name:
                        md_lines.append(f"# {text}")
                    else:
                        md_lines.append(text)
                
                md_content = "\n\n".join(md_lines)
                file_bytes = md_content.encode('utf-8')
                mime_type = 'text/markdown'
                display_name = uploaded_file.name.replace('.docx', '.md').replace('.DOCX', '.md')
            except Exception as docx_err:
                st.error(f"Could not convert DOCX file: {docx_err}")
                return None
        else:
            # Build MIME type for other supported files
            mime_type = uploaded_file.type or "application/octet-stream"
            if uploaded_file.name.endswith('.pdf'):
                mime_type = 'application/pdf'
            elif uploaded_file.name.endswith('.txt'):
                mime_type = 'text/plain'
            elif uploaded_file.name.endswith('.md'):
                mime_type = 'text/markdown'
            display_name = uploaded_file.name

        # Upload to Gemini Files API
        # New SDK expects: file=bytes/path, config with mimeType
        import tempfile
        import os
        
        # Determine file extension for temp file
        ext = '.md' if uploaded_file.name.lower().endswith('.docx') else os.path.splitext(uploaded_file.name)[1]
        
        # Write to temp file (SDK expects file path or bytes)
        with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as tmp:
            tmp.write(file_bytes)
            tmp_path = tmp.name
        
        try:
            uploaded = client.files.upload(
                file=tmp_path,
                config=types.UploadFileConfig(
                    mime_type=mime_type,
                    display_name=display_name
                )
            )
        finally:
            # Clean up temp file
            os.unlink(tmp_path)

        review_text = review_document(uploaded.uri, context)
        st.session_state[f"file_{phase_key}"] = f"File: {uploaded_file.name} ({uploaded.uri})"
        return {
            "review": review_text,
            "file_uri": uploaded.uri,
            "filename": uploaded_file.name,
        }
    except Exception as e:
        st.error(f"File upload failed: {e}")
        return None


def review_document(file_uri: str, context: str = "") -> str:
    """
    Review an uploaded file using Gemini API and return a markdown summary.
    
    Args:
        file_uri: URI of the file uploaded to Gemini Files API
        context: Optional context about what the file is for (e.g., 'clinical pathway')
    
    Returns:
        str: Markdown-formatted review summary
    """
    client = get_genai_client()
    if not client:
        return "⚠️ API connection unavailable. Could not review file."
    
    try:
        context_phrase = f" for {context}" if context else ""
        prompt = f"""Review this document{context_phrase}. Provide a concise summary in markdown format including:
        
- **Purpose**: Main topic and intent
- **Key Points**: 3-5 most important takeaways
- **Relevance**: How this relates to clinical pathway development{context_phrase}

Keep the summary brief (3-5 sentences per section)."""
        
        # Use the file URI in the content
        contents = [
            {
                "parts": [
                    {"text": prompt},
                    {"file_data": {"file_uri": file_uri}}
                ]
            }
        ]
        
        response = client.models.generate_content(
            model="gemini-flash-latest",
            contents=contents,
            config=get_generation_config(enable_thinking=True, thinking_budget=512)
        )
        
        return response.text if response and response.text else "✓ File uploaded successfully. Content available for AI analysis."
    
    except Exception as e:
        return f"✓ File uploaded. *(Auto-review unavailable: {str(e)})*"


def create_formsubmit_html(form_html: str) -> str:
    """
    Ensure FormSubmit.co compatibility in generated HTML.
    Adds dynamic form action based on recipient email field.
    
    Args:
        form_html: Generated HTML form content
        
    Returns:
        str: HTML with FormSubmit integration and dynamic submission
    """
    # Add JavaScript to handle dynamic form submission
    formsubmit_script = r'''
    <script>
    document.addEventListener('DOMContentLoaded', function() {
        const form = document.querySelector('form');
        if (form) {
            // Add recipient email field at the top if it doesn't exist
            const hasRecipientField = form.querySelector('input[name="recipient_email"]');
            if (!hasRecipientField) {
                const fieldHTML = `
                    <div style="margin-bottom: 20px; padding: 15px; background-color: #f0f0f0; border-radius: 5px;">
                        <label for="recipient_email" style="font-weight: bold; display: block; margin-bottom: 5px;">
                            Where should this form submission be sent? *
                        </label>
                        <input type="email" id="recipient_email" name="recipient_email" required 
                               placeholder="your-email@hospital.org"
                               style="width: 100%; padding: 8px; border: 1px solid #ccc; border-radius: 3px;">
                        <small style="display: block; margin-top: 5px; color: #666;">
                            Form responses will be sent to this email address via FormSubmit.co
                        </small>
                    </div>
                `;
                form.insertAdjacentHTML('afterbegin', fieldHTML);
            }
            
            // Handle form submission
            form.addEventListener('submit', function(e) {
                e.preventDefault();
                st.divider()

                # Determine current phase selection from navigation
                # Use the authoritative `current_phase_label` to avoid desync loops
                const recipientEmail = form.querySelector('input[name="recipient_email"]').value;
                if (recipientEmail) {
                    form.action = `https://formsubmit.co/${recipientEmail}`;
                    form.method = 'POST';
                    // Add FormSubmit configuration fields
                    if (!form.querySelector('input[name="_subject"]')) {
                        const subjectInput = document.createElement('input');
                        subjectInput.type = 'hidden';
                        subjectInput.name = '_subject';
                        subjectInput.value = document.title || 'Form Submission';
                        form.appendChild(subjectInput);
                    }
                    if (!form.querySelector('input[name="_template"]')) {
                        const templateInput = document.createElement('input');
                        templateInput.type = 'hidden';
                        templateInput.name = '_template';
                        templateInput.value = 'box';
                        form.appendChild(templateInput);
                    }
                    if (!form.querySelector('input[name="_captcha"]')) {
                        const captchaInput = document.createElement('input');
                        captchaInput.type = 'hidden';
                        captchaInput.name = '_captcha';
                        captchaInput.value = 'false';
                        form.appendChild(captchaInput);
                    }
                    form.submit();
                } else {
                    alert('Please provide a recipient email address.');
                }
            });
        }
    });
    </script>
    '''
    
    # Insert script before closing body tag
    if '</body>' in form_html:
        form_html = form_html.replace('</body>', formsubmit_script + '</body>')
    else:
        form_html += formsubmit_script
    
    # Add honeypot field if not present
    if '_honey' not in form_html and 'bot-field' not in form_html:
        honeypot = '<input type="text" name="_honey" style="display:none">'
        if '<form' in form_html:
            form_html = form_html.replace('<form', f'<form>{honeypot}', 1)
    
    return form_html


def make_preview_only(form_html: str) -> str:
    """Convert a generated form to preview-only by hiding email/submit fields and blocking submission."""
    try:
        if not form_html:
            return form_html
        banner = (
            '<div style="background:#FFF3CD;border:1px solid #F0AD4E;color:#8A6D3B;'
            'padding:12px;border-radius:6px;margin:12px 0;">\n'
            '  Preview only — responses are not collected here. Use Google Forms to collect submissions.\n'
            '</div>\n'
        )
        style_js = (
            '<style id="cpq-preview-only">\n'
            'form input[type="email"],\n'
            'form input[type="submit"],\n'
            'form button[type="submit"],\n'
            'form button.submit,\n'
            'form .submit,\n'
            'form .cpq-submit { display: none !important; }\n'
            '</style>\n'
            '<script>\n'
            "document.addEventListener('DOMContentLoaded', function() {\n"
            "  document.querySelectorAll('form').forEach(function(f) {\n"
            "    f.addEventListener('submit', function(e) {\n"
            "      e.preventDefault();\n"
            "      alert('Preview only — responses are not collected here.');\n"
            "    });\n"
            "  });\n"
            "});\n"
            '</script>\n'
        )
        injection = style_js + banner
        if "</body>" in form_html:
            return form_html.replace("</body>", injection + "</body>")
        return injection + form_html
    except Exception:
        return form_html


def mark_provider_connected(provider: str):
    if "oauth" not in st.session_state:
        st.session_state["oauth"] = {}
    st.session_state["oauth"][provider] = True


def render_provider_card(name: str, provider_key: str, description: str, configured: bool, bullets: list, button_key: str):
    """Render a provider card with status and a mock connect CTA (placeholder until real OAuth)."""
    connected = configured or ("oauth" in st.session_state and st.session_state["oauth"].get(provider_key))
    status_color = "#2e7d32" if connected else "#9e9e9e"
    status_text = "Connected" if connected else "Not connected"
    st.markdown(f"<div style='border:1px solid #ddd;border-radius:8px;padding:14px;margin-bottom:8px;'>\n"
                f"<div style='display:flex;justify-content:space-between;align-items:center;'>"
                f"<strong>{name}</strong>"
                f"<span style='padding:4px 10px;border-radius:12px;background:{status_color};color:#fff;font-size:12px;'>{status_text}</span>"
                f"</div>\n"
                f"<div style='color:#444;margin:6px 0;'>{description}</div>\n"
                + "".join([f"<div style='color:#555;font-size:13px;'>• {b}</div>" for b in bullets]) +
                "</div>", unsafe_allow_html=True)
    if st.button(f"Connect {name.split()[0]}", width="stretch", key=button_key):
        mark_provider_connected(provider_key)
        st.success(f"{name} connected for this session (placeholder; add OAuth to persist).")


def is_configured(keys) -> bool:
    """Return True if any of the given keys exist in env or Streamlit secrets."""
    try:
        for k in keys:
            if os.getenv(k):
                return True
            if hasattr(st, "secrets") and k in st.secrets:
                return True
    except Exception:
        return False
    return False


def is_google_forms_configured() -> bool:
    # Minimal heuristic: presence of OAuth creds or service account JSON
    keys = [
        "GOOGLE_FORMS_CLIENT_ID",
        "GOOGLE_FORMS_CLIENT_SECRET",
        "GOOGLE_FORMS_REFRESH_TOKEN",
        "GOOGLE_APPLICATION_CREDENTIALS",
    ]
    has_env = is_configured(keys)
    has_session = "oauth" in st.session_state and st.session_state["oauth"].get("google")
    return bool(has_env or has_session)


def _get_secret(key: str, default: str | None = None) -> str | None:
    try:
        if hasattr(st, "secrets") and key in st.secrets:
            return st.secrets[key]
    except Exception:
        pass
    return os.getenv(key, default)


def build_google_auth_url() -> str | None:
    client_id = _get_secret("GOOGLE_FORMS_CLIENT_ID")
    redirect_uri = _get_secret("GOOGLE_REDIRECT_URI", _get_secret("OAUTH_REDIRECT_URI", "http://localhost:8501/"))
    scopes = [
        "openid",
        "email",
        "profile",
        "https://www.googleapis.com/auth/forms.body",
        "https://www.googleapis.com/auth/forms.body.readonly",
    ]
    if not client_id or not redirect_uri:
        return None
    state = secrets.token_urlsafe(16)
    if "oauth" not in st.session_state:
        st.session_state["oauth"] = {}
    st.session_state["oauth"]["google_state"] = state
    params = {
        "client_id": client_id,
        "redirect_uri": redirect_uri,
        "response_type": "code",
        "scope": " ".join(scopes),
        "access_type": "offline",
        "include_granted_scopes": "true",
        "prompt": "consent",
        "state": state,
    }
    return "https://accounts.google.com/o/oauth2/v2/auth?" + urllib.parse.urlencode(params)


def complete_google_oauth(params: dict) -> bool:
    code = params.get("code", [None])[0]
    state = params.get("state", [None])[0]
    expected = st.session_state.get("oauth", {}).get("google_state")
    if not code or not state or state != expected:
        return False
    client_id = _get_secret("GOOGLE_FORMS_CLIENT_ID")
    client_secret = _get_secret("GOOGLE_FORMS_CLIENT_SECRET")
    redirect_uri = _get_secret("GOOGLE_REDIRECT_URI", _get_secret("OAUTH_REDIRECT_URI", "http://localhost:8501/"))
    token_url = "https://oauth2.googleapis.com/token"
    data = {
        "code": code,
        "client_id": client_id,
        "client_secret": client_secret,
        "redirect_uri": redirect_uri,
        "grant_type": "authorization_code",
    }
    try:
        import requests
        resp = requests.post(token_url, data=data, timeout=10)
        if resp.status_code == 200:
            token = resp.json()
            st.session_state.setdefault("oauth_tokens", {})["google"] = token
            st.session_state["oauth"]["google"] = True
            return True
    except Exception:
        return False
    return False


def render_connect_link(provider_key: str):
    if provider_key != "google":
        return
    url = build_google_auth_url()
    if url:
        st.link_button("Open consent window", url, width="stretch")
    else:
        st.warning("Missing client configuration. Please set client ID/redirect URI in secrets or env.")


def ensure_carepathiq_footer(html: str) -> str:
    """Append a CarePathIQ-branded footer with inline logo if missing."""
    try:
        if not html or "carepathiq-brand-footer" in html:
            return html
        footer_html = """
<div class=\"carepathiq-brand-footer\" style=\"text-align:center;margin-top:32px;padding-top:16px;border-top:1px solid #ddd;color:#444;\">
  <div style=\"font-weight:600;\">CarePathIQ</div>
  <svg viewBox=\"0 0 300 60\" xmlns=\"http://www.w3.org/2000/svg\" role=\"img\" aria-label=\"CarePathIQ logo\" style=\"width:160px;height:auto;display:block;margin:8px auto;\">
    <rect x=\"0\" y=\"0\" width=\"300\" height=\"60\" fill=\"transparent\"/>
    <text x=\"60\" y=\"38\" font-family=\"Segoe UI, Tahoma, Verdana, sans-serif\" font-size=\"28\" fill=\"#3E2723\">Care</text>
    <text x=\"130\" y=\"38\" font-family=\"Segoe UI, Tahoma, Verdana, sans-serif\" font-size=\"28\" fill=\"#5D4037\">Path</text>
    <text x=\"210\" y=\"38\" font-family=\"Segoe UI, Tahoma, Verdana, sans-serif\" font-size=\"28\" fill=\"#00695C\">IQ</text>
    <circle cx=\"25\" cy=\"30\" r=\"10\" fill=\"#A9EED1\" stroke=\"#3E2723\" stroke-width=\"2\"/>
  </svg>
  <div style=\"font-size:12px;color:#666;\">© CarePathIQ</div>
</div>
"""
        if "</body>" in html:
            return html.replace("</body>", footer_html + "</body>")
        return html + footer_html
    except Exception:
        return html

def fix_edu_certificate_html(html: str) -> str:
    """Ensure the education module's certificate uses landscape layout, brand colors,
    updated approval text, and includes a CarePathIQ logo at the bottom.
    Safe to call multiple times; idempotent-ish.
    """
    try:
        if not html:
            return html
        updated = html.replace("Verified Education Credit", "Approved by CarePathIQ")
        style = """
<style id="cpq-certificate-style">
@page { size: landscape; margin: 1in; }
:root { --cpq-brown:#5D4037; --cpq-brown-dark:#3E2723; --cpq-teal:#A9EED1; }
.cpq-certificate {
  width: 100%;
  max-width: 1100px;
  margin: 20px auto;
  background: #fff;
  border: 6px solid var(--cpq-brown);
  padding: 32px 40px;
  box-shadow: 0 4px 12px rgba(0,0,0,0.08);
}
.cpq-cert-header { text-align:center; color: var(--cpq-brown-dark); border-bottom: 3px solid var(--cpq-teal); padding-bottom: 8px; margin-bottom: 16px; }
.cpq-cert-badge { display:inline-block; background: var(--cpq-teal); color: var(--cpq-brown-dark); font-weight:700; padding: 4px 10px; border-radius: 4px; margin-top: 6px; }
.cpq-cert-body { color: #333; line-height: 1.5; }
.cpq-cert-footer { text-align:center; margin-top: 24px; color: #444; }
.cpq-cert-logo { display:block; margin: 12px auto 0 auto; width: 160px; height: auto; }
</style>
"""
        if "cpq-certificate-style" not in updated:
            if "</head>" in updated:
                updated = updated.replace("</head>", style + "</head>")
            else:
                updated = style + updated

        # Append a real logo image (base64) footer if available; fallback to inline SVG
        if "cpq-cert-logo" not in updated and "Approved by CarePathIQ" in updated:
            logo_html = ""
            try:
                logo_path = os.path.join(os.getcwd(), "CarePathIQ_Logo.png")
                if os.path.exists(logo_path):
                    with open(logo_path, "rb") as f:
                        logo_b64 = base64.b64encode(f.read()).decode("utf-8")
                    logo_html = f"""
<div class=\"cpq-cert-footer\">\n  <div>Approved by CarePathIQ</div>\n  <img class=\"cpq-cert-logo\" alt=\"CarePathIQ logo\" src=\"data:image/png;base64,{logo_b64}\" />\n  <div style=\"font-size:12px;color:#666;\">© CarePathIQ</div>\n</div>\n"""
            except Exception:
                logo_html = ""
            if not logo_html:
                logo_html = """
<div class=\"cpq-cert-footer\">\n  <div>Approved by CarePathIQ</div>\n  <svg class=\"cpq-cert-logo\" viewBox=\"0 0 300 60\" xmlns=\"http://www.w3.org/2000/svg\" role=\"img\" aria-label=\"CarePathIQ logo\">\n    <rect x=\"0\" y=\"0\" width=\"300\" height=\"60\" fill=\"transparent\"/>\n    <text x=\"60\" y=\"38\" font-family=\"Segoe UI, Tahoma, Verdana, sans-serif\" font-size=\"28\" fill=\"#3E2723\">Care</text>\n    <text x=\"130\" y=\"38\" font-family=\"Segoe UI, Tahoma, Verdana, sans-serif\" font-size=\"28\" fill=\"#5D4037\">Path</text>\n    <text x=\"210\" y=\"38\" font-family=\"Segoe UI, Tahoma, Verdana, sans-serif\" font-size=\"28\" fill=\"#00695C\">IQ</text>\n    <circle cx=\"25\" cy=\"30\" r=\"10\" fill=\"#A9EED1\" stroke=\"#3E2723\" stroke-width=\"2\"/>\n  </svg>\n  <div style=\"font-size:12px;color:#666;\">© CarePathIQ</div>\n</div>\n"""
            if "</body>" in updated:
                updated = updated.replace("</body>", logo_html + "</body>")
            else:
                updated = updated + logo_html

        return updated
    except Exception:
        return html

def render_refine_suggestions(target_key: str, suggestions: list[str]):
    """Render clickable suggestion buttons that pre-populate a refine text area.
    Clicking a suggestion sets st.session_state[target_key] and reruns.
    """
    if not suggestions:
        return
    st.caption("Quick suggestions:")
    # Render in rows of 3 buttons
    for i in range(0, len(suggestions), 3):
        row = suggestions[i:i+3]
        cols = st.columns(len(row))
        for idx, s in enumerate(row):
            with cols[idx]:
                if st.button(s, key=f"{target_key}_sugg_{i+idx}", width="stretch"):
                    st.session_state[target_key] = s
                    st.rerun()

@st.cache_data(ttl=3600)
def generate_gantt_image(schedule):
    if not schedule: return None
    # Safety check if matplotlib is not installed/imported
    if 'plt' not in globals() or 'mdates' not in globals() or plt is None: return None
    try:
        df = pd.DataFrame(schedule)
        df['Start'] = pd.to_datetime(df['Start'])
        df['End'] = pd.to_datetime(df['End'])
        df['Duration'] = (df['End'] - df['Start']).dt.days
        fig, ax = plt.subplots(figsize=(8, 4))
        y_pos = range(len(df))
        ax.barh(y_pos, df['Duration'], left=mdates.date2num(df['Start']), align='center', color='#5D4037', alpha=0.8)
        ax.set_yticks(y_pos)
        # Use 'Stage' column for labels if available
        labels = df['Stage'].tolist() if 'Stage' in df.columns else df['Phase'].tolist()
        ax.set_yticklabels(labels, fontsize=9)
        ax.invert_yaxis()
        ax.xaxis_date()
        ax.xaxis.set_major_formatter(mdates.DateFormatter('%b %Y'))
        ax.set_xlabel('Timeline')
        ax.set_title('Project Schedule', fontsize=12, fontweight='bold', color='#5D4037')
        plt.tight_layout()
        img_buffer = BytesIO()
        plt.savefig(img_buffer, format='png', dpi=150)
        plt.close(fig)
        img_buffer.seek(0)
        return img_buffer
    except Exception as e:
        return None

def create_word_docx(data):
    if Document is None: return None
    try:
        doc = Document()
        doc.add_heading(f"Project Charter: {data.get('condition', 'Untitled')}", 0)
        ihi = data.get('ihi_content', {})
        
        doc.add_heading('What are we trying to accomplish?', level=1)
        doc.add_heading('Problem', level=2)
        doc.add_paragraph(ihi.get('problem', data.get('problem', '')))
        
        doc.add_heading('Project Description', level=2)
        doc.add_paragraph(ihi.get('project_description', ''))
        
        doc.add_heading('Rationale', level=2)
        doc.add_paragraph(ihi.get('rationale', ''))
        
        doc.add_heading('Expected Outcomes', level=2)
        doc.add_paragraph(ihi.get('expected_outcomes', ''))
        
        doc.add_heading('Aim Statement', level=2)
        doc.add_paragraph(ihi.get('aim_statement', data.get('objectives', '')))

        doc.add_heading('How will we know that a change is an improvement?', level=1)
        doc.add_heading('Outcome Measure(s)', level=2)
        for m in ihi.get('outcome_measures', []): doc.add_paragraph(m, style='List Bullet')
        
        doc.add_heading('Process Measure(s)', level=2)
        for m in ihi.get('process_measures', []): doc.add_paragraph(m, style='List Bullet')
        
        doc.add_heading('Balancing Measure(s)', level=2)
        for m in ihi.get('balancing_measures', []): doc.add_paragraph(m, style='List Bullet')

        doc.add_heading('What changes can we make?', level=1)
        doc.add_heading('Initial Activities', level=2)
        doc.add_paragraph(ihi.get('initial_activities', ''))
        
        doc.add_heading('Change Ideas', level=2)
        for c in ihi.get('change_ideas', []): doc.add_paragraph(c, style='List Bullet')
        
        doc.add_heading('Key Stakeholders', level=2)
        doc.add_paragraph(ihi.get('stakeholders', ''))
        
        doc.add_heading('Barriers', level=2)
        doc.add_paragraph(ihi.get('barriers', ''))
        
        doc.add_heading('Boundaries', level=2)
        # Check both cases for boundaries key
        boundaries = ihi.get('boundaries', ihi.get('Boundaries', ''))
        if isinstance(boundaries, dict):
            # Format dictionary nicely (In Scope / Out of Scope)
            in_scope = boundaries.get('in_scope', boundaries.get('In Scope', ''))
            out_scope = boundaries.get('out_of_scope', boundaries.get('Out of Scope', ''))
            
            if in_scope:
                p = doc.add_paragraph()
                run = p.add_run("In Scope: ")
                run.bold = True
                p.add_run(str(in_scope))
                
            if out_scope:
                p = doc.add_paragraph()
                run = p.add_run("Out of Scope: ")
                run.bold = True
                p.add_run(str(out_scope))
        else:
            doc.add_paragraph(str(boundaries))
        
        doc.add_heading('Project Timeline', level=1)
        schedule = data.get('schedule', [])
        if schedule:
            gantt_img = generate_gantt_image(schedule)
            if gantt_img:
                doc.add_picture(gantt_img, width=DocxInches(6))
                doc.add_paragraph("")
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Stage'; hdr_cells[1].text = 'Owner'; hdr_cells[2].text = 'Start Date'; hdr_cells[3].text = 'End Date'
            for item in schedule:
                row_cells = table.add_row().cells
                row_cells[0].text = str(item.get('Stage', ''))
                row_cells[1].text = str(item.get('Owner', ''))
                row_cells[2].text = str(item.get('Start', ''))
                row_cells[3].text = str(item.get('End', ''))
        section = doc.sections[0]; footer = section.footer; p = footer.paragraphs[0]
        p.text = "Adapted from IHI QI Project Charter: https://www.ihi.org/library/tools/qi-project-charter\nCarePathIQ © 2024 by Tehreem Rehman is licensed under CC BY-SA 4.0"; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        buffer = BytesIO(); doc.save(buffer); buffer.seek(0)
        return buffer
    except Exception as e:
        st.error(f"Error generating document: {str(e)}")
        return None

def generate_p1_charter():
    """
    Auto-generate Phase 1 Project Charter using IHI Model for Improvement.
    Updates st.session_state.data['phase1']['ihi_content'] and triggers download.
    """
    d = st.session_state.data['phase1']
    if not d.get('condition') or not d.get('problem'):
        st.warning("Please fill in Condition and Problem before generating charter.")
        return
    
    with st.status("Generating Project Charter...", expanded=True) as status:
        st.write("Building project charter based on IHI Quality Improvement framework...")
        
        p_ihi = f"""You are a Quality Improvement Advisor using IHI's Model for Improvement.

Build a project charter for managing "{d['condition']}" in "{d.get('setting', '') or 'care setting'}".

**Phase 1 Context:**
Problem: {d['problem']}
Inclusion: {d.get('inclusion', '')}
Exclusion: {d.get('exclusion', '')}
Objectives: {d.get('objectives', '')}

Return JSON with keys:
- project_description (string)
- rationale (string)
- expected_outcomes (string)
- aim_statement (string)
- outcome_measures (string)
- process_measures (string)
- balancing_measures (string)
- initial_activities (string)
- change_ideas (array of strings)
- stakeholders (string)
- barriers (string)
- boundaries: object with in_scope (string) and out_of_scope (string)

Return clean JSON ONLY. No markdown, no explanation."""
        
        # Use native function calling for reliable structured output
        result = get_gemini_response(
            p_ihi, 
            function_declaration=CREATE_IHI_CHARTER,
            thinking_budget=1024
        )
        # Extract from function call or fall back
        if isinstance(result, dict) and 'arguments' in result:
            res = result['arguments']
        elif isinstance(result, dict):
            res = result
        else:
            res = get_gemini_response(p_ihi, json_mode=True)
        if res:
            st.session_state.data['phase1']['ihi_content'] = res
            doc = create_word_docx(st.session_state.data['phase1'])
            if doc:
                status.update(label="Ready!", state="complete")
                st.download_button(
                    "Download Project Charter (.docx)",
                    doc,
                    f"Project_Charter_{d['condition']}.docx",
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            else:
                status.update(label="Word export unavailable. Please ensure python-docx is installed.", state="error")
        else:
            status.update(label="Failed to generate charter.", state="error")

def create_exec_summary_docx(summary_text, condition):
    if Document is None: return None
    doc = Document()
    doc.add_heading(f"Executive Summary: {condition}", 0)
    for line in summary_text.split('\n'):
        if line.strip():
            if line.startswith('###'): doc.add_heading(line.replace('###', '').strip(), level=3)
            elif line.startswith('##'): doc.add_heading(line.replace('##', '').strip(), level=2)
            elif line.startswith('#'): doc.add_heading(line.replace('#', '').strip(), level=1)
            else: doc.add_paragraph(line.strip().replace('**', ''))
    section = doc.sections[0]; footer = section.footer; p = footer.paragraphs[0]
    p.text = "CarePathIQ © 2024 by Tehreem Rehman is licensed under CC BY-SA 4.0"; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    buffer = BytesIO(); doc.save(buffer); buffer.seek(0)
    return buffer

def auto_grade_evidence_list(evidence_list: list):
    """
    Auto-grade a list of evidence items using GRADE criteria via Gemini API.
    Updates each item in the list with 'grade' and 'rationale' fields.
    
    Args:
        evidence_list: List of evidence dictionaries with at least 'id' and 'title' keys
    """
    if not evidence_list:
        return
    
    try:
        prompt = (
            "Assign GRADE quality of evidence (use EXACTLY one of: 'High (A)', 'Moderate (B)', 'Low (C)', or 'Very Low (D)') "
            "and provide a brief Rationale (1-2 sentences) for each article. "
            f"{json.dumps([{k:v for k,v in e.items() if k in ['id','title']} for e in evidence_list])}. "
            "Return ONLY valid JSON object where keys are PMID strings and values are objects with 'grade' and 'rationale' fields. "
            '{\"12345678\": {\"grade\": \"High (A)\", \"rationale\": \"text here\"}}'
        )
        
        # Use native function calling for reliable structured output
        result = get_gemini_response(
            prompt, 
            function_declaration=GRADE_EVIDENCE,
            thinking_budget=1024
        )
        # Extract grades from function call or fall back
        if isinstance(result, dict) and 'arguments' in result:
            grades = result['arguments'].get('grades', {})
        elif isinstance(result, dict) and 'grades' not in result:
            grades = result  # Already in expected format
        else:
            # Fallback to json_mode
            grades = get_gemini_response(prompt, json_mode=True)
        
        if grades and isinstance(grades, dict):
            for e in evidence_list:
                pmid_str = str(e.get('id', ''))
                if pmid_str in grades:
                    grade_data = grades[pmid_str]
                    if isinstance(grade_data, dict):
                        e['grade'] = grade_data.get('grade', 'Un-graded')
                        e['rationale'] = grade_data.get('rationale', 'Not provided.')
                    else:
                        e['grade'] = 'Un-graded'
                        e['rationale'] = 'Not provided.'
                else:
                    e.setdefault('grade', 'Un-graded')
                    e.setdefault('rationale', 'Not yet evaluated.')
        else:
            # If API call fails, set defaults
            for e in evidence_list:
                e.setdefault('grade', 'Un-graded')
                e.setdefault('rationale', 'Auto-grading unavailable.')
    
    except Exception as ex:
        # On error, set defaults and log
        for e in evidence_list:
            e.setdefault('grade', 'Un-graded')
            e.setdefault('rationale', f'Auto-grading error: {str(ex)}')

def format_citation_line(entry, style="APA"):
    """Lightweight formatter for citation strings based on available PubMed fields.
    Supports preset styles (APA, MLA, Vancouver) and custom style names."""
    authors = (entry.get("authors") or "Unknown").rstrip(".")
    title = (entry.get("title") or "Untitled").rstrip(".")
    journal = (entry.get("journal") or "Journal").rstrip(".")
    year = entry.get("year") or "n.d."
    pmid = entry.get("id") or ""
    if style == "MLA":
        return f"{authors}. \"{title}.\" {journal}, {year}. PMID {pmid}."
    if style == "Vancouver":
        return f"{authors}. {title}. {journal}. {year}. PMID:{pmid}."
    # Default APA (also used for custom style names)
    return f"{authors} ({year}). {title}. {journal}. PMID: {pmid}."

def apply_pathway_heuristic_improvements(nodes, heuristics_data, extra_ui_insights=None):
    """
    Intelligently apply feasible heuristics (H1-H10) to pathway nodes.
    CRITICAL: Preserves and enhances clinical complexity, does NOT reduce it.
    The AI evaluates each heuristic and applies only those that improve without compromising decision science.
    Returns: (updated_nodes, applied_heuristics_list, summary_text) or (None, [], "")
    """
    if not heuristics_data:
        return None, [], ""

    # Include all heuristics in the analysis
    insights_text = "\n".join([f"{k}: {v}" for k, v in sorted(heuristics_data.items())])

    guardrails = """Safety rules (MANDATORY):
1) PRESERVE all clinical complexity and decision branches—never simplify away clinical logic
2) Do NOT reduce decision divergence or collapse distinct pathways
3) Do NOT remove edge cases or special population considerations
4) Preserve all node IDs, branching structure, and evidence citations
5) Add detail and specificity—do NOT generalize clinical steps
6) Only modify text clarity, add safety annotations, or improve organization
7) Add new nodes ONLY if critical for safety or decision clarity
8) For each heuristic applied, explain how it enhances (not reduces) the pathway"""

    prompt = f"""You are a CLINICAL DECISION SCIENTIST with expertise in Medical Decision Analysis and Nielsen's Usability Heuristics.

TASK: Apply feasible heuristics to improve this clinical decision pathway while PRESERVING AND ENHANCING decision-science integrity.

CRITICAL PRINCIPLE: This is NOT about simplification. Improvements should make the pathway MORE usable, more complete, and more clinically rigorous—not less complex.

Current pathway ({len(nodes)} nodes):
{json.dumps(nodes)}

Heuristic Assessment:
{insights_text}

APPLICATION STRATEGY:
- H1 (Status visibility): ADD checkpoint descriptions and alarm thresholds (enhances clinical specificity)
- H2 (Language clarity): Clarify terminology WITHOUT removing medical precision needed for safety
- H3 (User control): ADD escape routes/alternative pathways (increases decision options)
- H4 (Consistency): Standardize decision structures AND expand them uniformly
- H5 (Error prevention): ADD validation rules and edge case handling (increases complexity beneficially)
- H6 (Recognition not recall): Improve labeling clarity while preserving all decision detail
- H7 (Efficiency): Remove ONLY redundant steps; keep clinical content and decision branches
- H8 (Minimalist): Consolidate presentation, NOT clinical content; respect DAG and decision divergence
- H9 (Error recovery): Move critical checks EARLIER and ADD recovery pathways (increases safety)
- H10 (Help & docs): ADD evidence citations and rationale annotations to decision nodes

{guardrails}

BEFORE/AFTER RULE:
- Evaluate: Does this improvement ADD clinical value, clarity, or safety?
- If yes: Apply it and include in applied_heuristics list
- If no: Skip it
- NEVER: Reduce complexity, remove branches, generalize clinical steps, or simplify decision trees

Return ONLY valid JSON:
{{
  "updated_nodes": [array of modified node objects with enhanced detail and annotations],
  "applied_heuristics": ["H2", "H4", "H5", ...list of heuristics that genuinely improved the pathway],
  "applied_summary": "Detailed explanation of improvements made and how each enhances clinical decision-making"
}}

VALIDATION CHECKLIST BEFORE RETURNING:
- Node count maintained or INCREASED (not decreased)
- All Decision node branches still present and distinct
- Evidence citations preserved on all applicable nodes
- No "or" statements in End nodes
- DAG structure maintained (no cycles)
- All 4 clinical stages still represented: Initial Evaluation, Diagnosis/Treatment, Re-evaluation, Final Disposition
- Clinical depth enhanced, not reduced"""

    # Use native function calling for reliable structured output
    result = get_gemini_response(
        prompt, 
        function_declaration=APPLY_HEURISTICS,
        thinking_budget=2048  # Complex heuristics application
    )
    # Extract from function call or fall back
    if isinstance(result, dict) and 'arguments' in result:
        response = result['arguments']
    elif isinstance(result, dict) and 'updated_nodes' in result:
        response = result
    else:
        response = get_gemini_response(prompt, json_mode=True)
    
    if response and isinstance(response, dict):
        updated_nodes = response.get("updated_nodes")
        applied = response.get("applied_heuristics", [])
        summary = response.get("applied_summary", "")
        
        if updated_nodes and isinstance(updated_nodes, list) and isinstance(applied, list):
            # Optionally validate the returned nodes meet our standards
            validation = validate_decision_science_pathway(updated_nodes)
            if validation['complexity']['complexity_level'] != 'comprehensive':
                # If heuristics reduced complexity, log warning but still return
                st.warning(f"⚠️ Heuristic application reduced pathway complexity. Original: {len(nodes)} nodes, Updated: {len(updated_nodes)} nodes. Review the changes.")
            return updated_nodes, applied, summary
    
    return None, [], ""


def apply_actionable_heuristics_incremental(nodes, heuristics_data):
    """
    Intelligently apply all feasible heuristics to the pathway.
    The AI evaluates all H1-H10 and applies only those that work for pathway nodes.
    Returns: (updated_nodes, applied_heuristics_list, summary_text)
    """
    updated_nodes, applied_list, summary = apply_pathway_heuristic_improvements(nodes, heuristics_data, {})
    return updated_nodes, applied_list, summary

def create_references_docx(citations, style="APA"):
    """Create Word document with citations. Supports preset styles and custom style names."""
    if Document is None or not citations:
        return None
    doc = Document()
    # Use the style name (preset or custom) in the heading
    heading_text = f"References ({style})" if style else "References"
    doc.add_heading(heading_text, 0)
    for idx, entry in enumerate(citations, start=1):
        line = format_citation_line(entry, style)
        doc.add_paragraph(f"{idx}. {line}")
    # Add licensing/footer similar to other DOCX exports
    try:
        section = doc.sections[0]
        footer = section.footer
        p = footer.paragraphs[0]
        p.text = "CarePathIQ © 2024 by Tehreem Rehman is licensed under CC BY-SA 4.0"
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception:
        pass
    buffer = BytesIO(); doc.save(buffer); buffer.seek(0)
    return buffer

def harden_nodes(nodes_list):
    """Validate and fix node structure, ensuring Decision nodes have proper branches.
    
    IMPORTANT: Preserves existing branch structure when valid. Only creates default
    branches when none exist or when branches are malformed.
    """
    if not isinstance(nodes_list, list): return []
    validated = []
    n = len(nodes_list)
    
    for i, node in enumerate(nodes_list):
        if not isinstance(node, dict): continue
        # Ensure required fields
        if 'id' not in node or not node['id']: 
            node['id'] = f"{node.get('type', 'P')[0].upper()}{i+1}"
        if 'type' not in node: 
            node['type'] = 'Process'
        if 'label' not in node or not node.get('label'):
            node['label'] = f"Step {i+1}"
        
        # Clean up label text - remove literal \n, collapse whitespace
        label = str(node.get('label', ''))
        # Remove literal backslash-n sequences (from AI generation)
        label = label.replace('\\n', ' ').replace('\n', ' ')
        # Collapse multiple spaces
        import re
        label = re.sub(r'\s+', ' ', label).strip()
        node['label'] = label
        
        # Also clean notes field if present
        for notes_field in ['notes', 'detail']:
            if notes_field in node and node[notes_field]:
                notes = str(node[notes_field])
                notes = notes.replace('\\n', ' ').replace('\n', ' ')
                notes = re.sub(r'\s+', ' ', notes).strip()
                node[notes_field] = notes
        
        # Validate Decision nodes have branches - PRESERVE existing valid branches
        if node['type'] == 'Decision':
            existing_branches = node.get('branches', [])
            
            # Check if existing branches are valid
            valid_branches = []
            if isinstance(existing_branches, list):
                for branch in existing_branches:
                    if isinstance(branch, dict):
                        target = branch.get('target')
                        # Check if target is valid (numeric and in range)
                        if isinstance(target, (int, float)) and 0 <= int(target) < n:
                            # Keep this branch, just ensure it has a label
                            if 'label' not in branch or not branch.get('label'):
                                branch['label'] = 'Option'
                            valid_branches.append(branch)
                        elif isinstance(target, (int, float)):
                            # Target out of range - clamp it
                            branch['target'] = max(0, min(int(target), n - 1))
                            if 'label' not in branch or not branch.get('label'):
                                branch['label'] = 'Option'
                            valid_branches.append(branch)
            
            if len(valid_branches) >= 2:
                # Keep the existing valid branches - PRESERVE branching logic
                node['branches'] = valid_branches
            elif len(valid_branches) == 1:
                # Only one valid branch - try to find another appropriate target
                existing_target = int(valid_branches[0].get('target', i + 1))
                # Find an End node or another branch point
                alt_targets = []
                for j in range(i + 1, n):
                    if j != existing_target:
                        alt_targets.append(j)
                if alt_targets:
                    alt_idx = alt_targets[0]  # Use first available alternative
                    node['branches'] = valid_branches + [{'label': 'No', 'target': alt_idx}]
                else:
                    # No alternative found, create default
                    node['branches'] = valid_branches + [{'label': 'No', 'target': min(i + 1, n - 1)}]
            else:
                # No valid branches - create default divergent branches
                # Look for End nodes to create meaningful branches
                end_nodes = [j for j in range(i + 1, n) if nodes_list[j].get('type') == 'End']
                
                if len(end_nodes) >= 2:
                    # Branch to different End nodes
                    node['branches'] = [
                        {'label': 'Yes', 'target': end_nodes[0]}, 
                        {'label': 'No', 'target': end_nodes[1]}
                    ]
                elif len(end_nodes) == 1 and i + 1 < n and i + 1 != end_nodes[0]:
                    # Branch to next node and End node
                    node['branches'] = [
                        {'label': 'Yes', 'target': i + 1}, 
                        {'label': 'No', 'target': end_nodes[0]}
                    ]
                else:
                    # Fallback: sequential branches (less ideal but functional)
                    next_idx = min(i + 1, n - 1)
                    alt_idx = min(i + 2, n - 1) if i + 2 < n else next_idx
                    node['branches'] = [
                        {'label': 'Yes', 'target': next_idx}, 
                        {'label': 'No', 'target': alt_idx}
                    ]
        
        validated.append(node)
    return validated

def validate_pathway_flow(nodes_list):
    """
    Validate pathway for common flow issues:
    - Unreachable nodes (orphaned)
    - Invalid branch targets
    - Missing End nodes
    - Cycles (if DAG-only enforcement needed)
    Returns: (is_valid, issues_list)
    """
    if not isinstance(nodes_list, list) or len(nodes_list) == 0:
        return False, ["Empty pathway"]
    
    issues = []
    n = len(nodes_list)
    reachable = set([0])  # Start node is always index 0
    
    # Build reachability graph
    for i, node in enumerate(nodes_list):
        if not isinstance(node, dict):
            issues.append(f"Node {i}: Invalid node structure")
            continue
        
        node_type = node.get('type', 'Process')
        
        if node_type == 'Decision':
            branches = node.get('branches', [])
            if not branches:
                issues.append(f"Node {i} ({node.get('label', 'N/A')}): Decision node has no branches")
            for branch in branches:
                target = branch.get('target')
                if not isinstance(target, (int, float)):
                    issues.append(f"Node {i}: Branch '{branch.get('label', 'N/A')}' has invalid target: {target}")
                elif not (0 <= int(target) < n):
                    issues.append(f"Node {i}: Branch '{branch.get('label', 'N/A')}' points to out-of-bounds index: {target}")
                else:
                    reachable.add(int(target))
        elif i + 1 < n:
            # Non-decision nodes implicitly connect to next node
            reachable.add(i + 1)
    
    # Check for unreachable nodes (except End nodes which are terminal)
    unreachable = []
    for i, node in enumerate(nodes_list):
        if i not in reachable and node.get('type') != 'End':
            unreachable.append(f"Node {i} ({node.get('label', 'N/A')})")
    
    if unreachable:
        issues.append(f"Unreachable nodes: {', '.join(unreachable)}")
    
    # Check for at least one End node
    has_end = any(node.get('type') == 'End' for node in nodes_list)
    if not has_end:
        issues.append("Pathway has no End nodes")
    
    # Check for Start node
    if nodes_list[0].get('type') != 'Start':
        issues.append(f"First node should be Start, found: {nodes_list[0].get('type')}")
    
    return len(issues) == 0, issues

def fix_decision_flow_issues(nodes_list):
    """
    Fix common AI generation issues:
    1. Remove any nodes that appear after End nodes (End nodes must be terminal)
    2. Validate Decision node branches point to valid indices
    3. Detect if Decision branches artificially reconverge (all branches point to same outcome)
    4. Flag and warn about reconvergence for user review
    """
    if not isinstance(nodes_list, list) or len(nodes_list) == 0:
        return nodes_list
    
    # Find FIRST End node - but don't truncate yet; just validate structure
    # Multiple End nodes are allowed for different pathways
    end_indices = [i for i, node in enumerate(nodes_list) if isinstance(node, dict) and node.get('type') == 'End']
    
    # Validate Decision node branches and FIX reconvergence issues
    for i, node in enumerate(nodes_list):
        if not isinstance(node, dict):
            continue
        if node.get('type') == 'Decision' and 'branches' in node:
            valid_branches = []
            for branch in node.get('branches', []):
                target = branch.get('target')
                if isinstance(target, (int, float)) and 0 <= int(target) < len(nodes_list):
                    valid_branches.append(branch)
            
            if valid_branches:
                node['branches'] = valid_branches
                
                # CHECK FOR RECONVERGENCE: Do all branches point to the same target?
                branch_targets = [int(b.get('target', -1)) for b in valid_branches]
                unique_targets = set(branch_targets)
                
                if len(unique_targets) == 1 and len(branch_targets) > 1:
                    # All branches point to same node - FIX this by finding alternative targets
                    original_target = branch_targets[0]
                    
                    # Find alternative End nodes or Process nodes to create actual divergence
                    alternative_targets = []
                    for j in range(i + 1, len(nodes_list)):
                        if j != original_target:
                            # Prefer End nodes, then other Process nodes
                            alternative_targets.append(j)
                    
                    if alternative_targets:
                        # Reassign branches to create true divergence
                        for branch_idx, branch in enumerate(node['branches']):
                            if branch_idx == 0:
                                branch['target'] = original_target  # Keep first branch on original
                            else:
                                # Assign to different targets
                                alt_idx = (branch_idx - 1) % len(alternative_targets)
                                branch['target'] = alternative_targets[alt_idx]
    
    return nodes_list

def normalize_or_logic(nodes_list):
    """
    Automatically detect and convert OR logic in End nodes to proper Decision nodes with branches.
    This runs silently as backend auto-normalization with no user notification.
    
    Example: "Discharge or Admit to Cardiology" (End node)
    Becomes: "Disposition decision?" (Decision) with branches to "Discharge" and "Admit to Cardiology" (End nodes)
    """
    if not isinstance(nodes_list, list) or len(nodes_list) == 0:
        return nodes_list
    
    normalized = []
    added_nodes = []
    
    for idx, node in enumerate(nodes_list):
        if not isinstance(node, dict):
            normalized.append(node)
            continue
        
        # Check if this is an End node with OR logic
        label = node.get('label', '').lower()
        if node.get('type') == 'End' and ' or ' in label:
            # Extract the original label for clarity
            original_label = node.get('label', '')
            
            # Split by ' or ' to get individual outcomes
            outcomes = [o.strip() for o in original_label.split(' or ')]
            
            if len(outcomes) > 1:
                # Convert this End node to a Decision node
                decision_label = f"{outcomes[0].split()[0]} vs {outcomes[1].split()[0] if len(outcomes[1].split()) > 0 else 'other'}?"
                decision_node = {
                    'type': 'Decision',
                    'label': decision_label,
                    'evidence': node.get('evidence', 'N/A'),
                    'branches': []
                }
                
                # Create End nodes for each outcome
                end_node_targets = []
                for outcome_idx, outcome in enumerate(outcomes):
                    end_node = {
                        'type': 'End',
                        'label': outcome,
                        'evidence': 'N/A'
                    }
                    added_nodes.append(end_node)
                    # Calculate target index: current position + branches + previously added nodes
                    target_idx = len(normalized) + 1 + outcome_idx
                    end_node_targets.append(target_idx)
                    
                    # Add branch to decision node
                    decision_node['branches'].append({
                        'label': outcome.split()[0],  # First word of outcome as branch label
                        'target': target_idx
                    })
                
                normalized.append(decision_node)
            else:
                # No actual OR split possible, keep as-is
                normalized.append(node)
        else:
            normalized.append(node)
    
    # Append all newly created End nodes at the end
    result = normalized + added_nodes
    
    # Update all target indices to account for insertions
    # Rebuild with proper sequential indexing
    final_result = []
    for node in result:
        final_node = dict(node)
        if final_node.get('type') == 'Decision' and 'branches' in final_node:
            # Branches already set correctly during creation
            pass
        final_result.append(final_node)
    
    return final_result

def assess_clinical_complexity(nodes_list):
    """
    Assess whether a pathway has appropriate clinical complexity per decision science standards.
    
    Returns: dict with complexity metrics
    {
        'node_count': int,
        'complexity_level': 'minimal' | 'moderate' | 'comprehensive',
        'decision_count': int,
        'decision_divergence_ratio': float (how much branches stay separate),
        'evidence_coverage': float (% nodes with PMIDs),
        'clinical_stage_coverage': dict (which 4 stages are represented),
        'recommendations': [str]
    }
    """
    if not isinstance(nodes_list, list) or len(nodes_list) == 0:
        return {'complexity_level': 'minimal', 'recommendations': ['Pathway is empty']}
    
    metrics = {
        'node_count': len(nodes_list),
        'decision_count': 0,
        'process_count': 0,
        'end_count': 0,
        'decision_divergence_ratio': 0.0,
        'evidence_coverage': 0.0,
        'clinical_stage_coverage': {},
        'recommendations': []
    }
    
    # Count node types and stage coverage
    stages = {
        'initial_evaluation': False,
        'diagnosis_treatment': False,
        're_evaluation': False,
        'final_disposition': False
    }
    
    stage_keywords = {
        'initial_evaluation': ['initial', 'assess', 'vital', 'exam', 'presentation', 'triage'],
        'diagnosis_treatment': ['diagnos', 'treat', 'interven', 'medic', 'workup', 'order'],
        're_evaluation': ['recheck', 'monitor', 'response', 'follow', 'escalat', 're-evaluat'],
        'final_disposition': ['discharg', 'admit', 'transfer', 'disposition', 'prescri', 'referral']
    }
    
    pmid_nodes = 0
    for node in nodes_list:
        if not isinstance(node, dict):
            continue
        
        ntype = node.get('type', 'Process')
        if ntype == 'Decision':
            metrics['decision_count'] += 1
        elif ntype == 'Process':
            metrics['process_count'] += 1
        elif ntype == 'End':
            metrics['end_count'] += 1
        
        # Check stage coverage
        label_lower = (node.get('label', '') or '').lower()
        for stage, keywords in stage_keywords.items():
            if any(kw in label_lower for kw in keywords):
                stages[stage] = True
        
        # Check evidence coverage
        if node.get('evidence') and node.get('evidence') != 'N/A':
            pmid_nodes += 1
    
    metrics['clinical_stage_coverage'] = stages
    metrics['evidence_coverage'] = pmid_nodes / len(nodes_list) if nodes_list else 0.0
    
    # Assess divergence: check if Decision nodes lead to distinct downstream paths
    divergent_decisions = 0
    for i, node in enumerate(nodes_list):
        if not isinstance(node, dict) or node.get('type') != 'Decision':
            continue
        branches = node.get('branches', [])
        if len(branches) >= 2:
            # Check if branches lead to truly different sequences (not immediate reconvergence)
            branch_targets = [b.get('target') for b in branches if isinstance(b.get('target'), (int, float))]
            if len(set(branch_targets)) == len(branch_targets):  # All unique targets
                divergent_decisions += 1
    
    metrics['decision_divergence_ratio'] = divergent_decisions / max(1, metrics['decision_count'])
    
    # Determine complexity level
    if metrics['node_count'] < 12:
        metrics['complexity_level'] = 'minimal'
        metrics['recommendations'].append('⚠️ Pathway may be oversimplified. Consider adding more decision branches and edge cases.')
    elif metrics['node_count'] < 20:
        metrics['complexity_level'] = 'moderate'
        metrics['recommendations'].append('Pathway has moderate complexity. Consider adding edge cases or special populations.')
    else:
        metrics['complexity_level'] = 'comprehensive'
        metrics['recommendations'].append('✓ Pathway has appropriate complexity for evidence-based decision science.')
    
    # Check stage coverage
    stages_covered = sum(1 for v in stages.values() if v)
    if stages_covered < 4:
        metrics['recommendations'].append(f'⚠️ Missing clinical stages: {", ".join([k.replace("_", " ").title() for k, v in stages.items() if not v])}')
    
    # Check evidence coverage
    if metrics['evidence_coverage'] < 0.3:
        metrics['recommendations'].append(f'⚠️ Low evidence coverage ({metrics["evidence_coverage"]:.0%}). Add PMID citations to key clinical steps.')
    
    # Check decision divergence
    if metrics['decision_divergence_ratio'] < 0.5:
        metrics['recommendations'].append('⚠️ Many Decision nodes reconverge quickly. Consider keeping branches more distinct.')
    
    if metrics['end_count'] < 2:
        metrics['recommendations'].append('⚠️ Pathway has few distinct end points. Clinical reality typically has multiple outcomes.')
    
    return metrics

def assess_decision_science_integrity(nodes_list):
    """
    Assess whether pathway follows decision science best practices per Medical Decision Analysis framework.
    
    Returns: dict with integrity metrics and violations
    {
        'is_dag': bool (directed acyclic graph - no cycles),
        'terminal_end_nodes': bool (all End nodes are terminal),
        'no_or_logic': bool (no "or" statements in End nodes),
        'actionable_notes': bool (nodes have actionable clinical notes like red flags, thresholds),
        'evidence_cited': bool (key steps have PMIDs),
        'violations': [str]
    }
    """
    if not isinstance(nodes_list, list) or len(nodes_list) == 0:
        return {'violations': ['Empty pathway']}
    
    integrity = {
        'is_dag': True,
        'terminal_end_nodes': True,
        'no_or_logic': True,
        'actionable_notes': False,
        'evidence_cited': False,
        'violations': []
    }
    
    # Check for cycles (simple reachability check)
    visited = set()
    def has_cycle(node_idx, path):
        if node_idx in path:
            return True
        if node_idx in visited:
            return False
        visited.add(node_idx)
        path.add(node_idx)
        
        node = nodes_list[node_idx] if 0 <= node_idx < len(nodes_list) else None
        if not node or not isinstance(node, dict):
            return False
        
        if node.get('type') == 'Decision':
            for branch in node.get('branches', []):
                target = branch.get('target')
                if isinstance(target, (int, float)) and has_cycle(int(target), path.copy()):
                    return True
        elif node_idx + 1 < len(nodes_list):
            if has_cycle(node_idx + 1, path.copy()):
                return True
        
        return False
    
    if has_cycle(0, set()):
        integrity['is_dag'] = False
        integrity['violations'].append('🔄 Cycle detected: pathway has backward loops')
    
    # Check End nodes are terminal (nothing after them)
    for i, node in enumerate(nodes_list):
        if not isinstance(node, dict):
            continue
        if node.get('type') == 'End' and i + 1 < len(nodes_list):
            next_node = nodes_list[i + 1]
            if isinstance(next_node, dict) and next_node.get('type') not in ('End', None):
                integrity['terminal_end_nodes'] = False
                integrity['violations'].append(f"Node {i} is End but followed by {next_node.get('type')} at index {i+1}")
    
    # Check for OR logic in End nodes
    for i, node in enumerate(nodes_list):
        if not isinstance(node, dict):
            continue
        if node.get('type') == 'End':
            label = (node.get('label') or '').lower()
            if ' or ' in label:
                integrity['no_or_logic'] = False
                integrity['violations'].append(f"Node {i} (End): Contains 'or' logic: '{node.get('label')}'. Split into Decision branches.")
    
    # Check for actionable clinical notes (red flags, thresholds, monitoring parameters)
    actionable_keywords = ['red flag', 'threshold', 'monitor', 'escalate', 'alert', 'warning', 'if', 'when', '>', '<', '≥', '≤']
    notes_count = 0
    for node in nodes_list:
        if not isinstance(node, dict):
            continue
        notes = (node.get('notes', '') or node.get('detail', '') or '').lower()
        if notes and any(kw in notes for kw in actionable_keywords):
            notes_count += 1
    
    if notes_count >= max(1, len(nodes_list) // 5):
        integrity['actionable_notes'] = True
    else:
        integrity['violations'].append(f"⚠️ Few nodes have actionable notes ({notes_count}). Consider adding red flags, thresholds, or monitoring parameters.")
    
    # Check evidence coverage
    pmid_count = sum(1 for n in nodes_list if isinstance(n, dict) and n.get('evidence') and n.get('evidence') != 'N/A')
    if pmid_count >= len(nodes_list) * 0.3:
        integrity['evidence_cited'] = True
    else:
        integrity['violations'].append(f"⚠️ Low PMID coverage ({pmid_count}/{len(nodes_list)}). Cite evidence for key clinical steps.")
    
    return integrity

def validate_decision_science_pathway(nodes_list):
    """
    Comprehensive validation: combines complexity assessment and integrity check.
    Returns: dict with detailed diagnostic info
    """
    complexity = assess_clinical_complexity(nodes_list)
    integrity = assess_decision_science_integrity(nodes_list)
    flow_valid, flow_issues = validate_pathway_flow(nodes_list)
    
    return {
        'complexity': complexity,
        'integrity': integrity,
        'flow_valid': flow_valid,
        'flow_issues': flow_issues,
        'overall_quality': sum([
            complexity['complexity_level'] == 'comprehensive',
            integrity['is_dag'],
            integrity['terminal_end_nodes'],
            integrity['no_or_logic'],
            integrity['evidence_cited'],
            flow_valid
        ]) / 6
    }

def generate_mermaid_code(nodes, orientation="TD"):
    """
    Generate Mermaid flowchart code from pathway nodes.
    Uses pathway_generator module if available, otherwise falls back to DOT.
    
    Args:
        nodes: List of pathway node dictionaries
        orientation: "TD" (top-down) or "LR" (left-right)
    
    Returns:
        Mermaid diagram source code (or DOT code as fallback)
    """
    if PATHWAY_GENERATOR_AVAILABLE:
        try:
            return create_mermaid_from_nodes(nodes, include_styling=True)
        except Exception as e:
            debug_log(f"Mermaid generation error: {e}")
            return dot_from_nodes(nodes, orientation)
    return dot_from_nodes(nodes, orientation)

# --- GRAPH EXPORT HELPERS (Graphviz/DOT) ---
def _escape_label(text: str) -> str:
    if text is None:
        return ""
    # Escape quotes and backslashes for DOT labels
    # Note: \n in DOT is the line break character, don't double-escape it
    s = str(text).replace("\\", "\\\\").replace('"', "'")
    # Convert actual newlines to graphviz newlines
    s = s.replace("\n", "\\n")
    return s

def _wrap_label(text: str, width: int = 22, max_width: int = None) -> str:
    if not text:
        return ""
    # Clean up any literal \n sequences before wrapping
    clean_text = str(text).replace('\\n', ' ').replace('\n', ' ')
    # Collapse multiple spaces
    import re
    clean_text = re.sub(r'\s+', ' ', clean_text).strip()
    wrapped = textwrap.wrap(clean_text, width=max_width or width)
    return "\n".join(wrapped) if wrapped else clean_text

def _role_fill(role: str, default_fill: str) -> str:
    if not role:
        return default_fill
    return ROLE_COLORS.get(role, ROLE_COLORS.get(str(role).title(), default_fill))

def dot_from_nodes(nodes, orientation="TD") -> str:
    """Generate Graphviz DOT source from pathway nodes with clean decision tree layout.
    
    LAYOUT PRINCIPLES:
    1. Start node at top-left (rank=source)
    2. Clear top-to-bottom flow (TB rankdir)
    3. Decision nodes create true branching with distinct paths
    4. End nodes at bottom (rank=sink)
    5. No swimlane clustering (disrupts natural flow)
    """
    if not nodes:
        return "digraph G {\n  // No nodes\n}"
    valid_nodes = harden_nodes(nodes)
    
    # Identify special nodes for layout control
    start_node_idx = None
    end_node_indices = []
    decision_node_indices = []
    
    for i, n in enumerate(valid_nodes):
        ntype = n.get('type', 'Process')
        if ntype == 'Start' and start_node_idx is None:
            start_node_idx = i
        elif ntype == 'End':
            end_node_indices.append(i)
        elif ntype == 'Decision':
            decision_node_indices.append(i)
    
    rankdir = 'TB' if orientation == 'TD' else 'LR'
    
    # Build DOT with graph attributes for clean decision tree layout
    lines = [
        "digraph G {",
        f"  rankdir={rankdir};",
        "  splines=ortho;",  # Orthogonal edges for cleaner routing
        "  nodesep=0.8;",   # Horizontal spacing between nodes
        "  ranksep=1.0;",   # Vertical spacing between ranks
        "  node [fontname=Helvetica, fontsize=11];",
        "  edge [fontname=Helvetica, fontsize=10];",
    ]
    
    node_id_map = {}
    notes_list = []  # Collect notes for numbered legend
    notes_node_map = {}  # Map node index to note number
    
    # First pass: identify all nodes with notes and assign note numbers
    note_counter = 1
    for i, n in enumerate(valid_nodes):
        notes_text = n.get('notes', '') or n.get('detail', '')
        if notes_text and str(notes_text).strip():
            notes_list.append((note_counter, str(notes_text).strip()))
            notes_node_map[i] = note_counter
            note_counter += 1
    
    # Create all nodes WITHOUT swimlane clustering (cleaner layout)
    for i, n in enumerate(valid_nodes):
        nid = f"N{i}"
        node_id_map[i] = nid
        
        # Build label
        raw_label = n.get('label', 'Step')
        raw_label = str(raw_label).replace('\\n', ' ').replace('\n', ' ')
        wrapped_label = _wrap_label(raw_label, width=25)  # Slightly wider for readability
        
        # Add note reference if applicable
        if i in notes_node_map:
            wrapped_label = wrapped_label + f"\n(Note {notes_node_map[i]})"
        
        full_label = _escape_label(wrapped_label)
        ntype = n.get('type', 'Process')
        
        # Node styling based on type
        if ntype == 'Decision':
            shape, fill = 'diamond', '#F8CECC'  # Red/pink for decisions
        elif ntype == 'Start':
            shape, fill = 'oval', '#D5E8D4'  # Green for start
        elif ntype == 'End':
            shape, fill = 'oval', '#D5E8D4'  # Green for end
        elif ntype == 'Reevaluation':
            shape, fill = 'box', '#FFCC80'  # Orange for reevaluation
        else:  # Process
            shape, fill = 'box', '#FFF2CC'  # Yellow for process
        
        # Apply role-based coloring if role is specified
        role = n.get('role', '')
        if role:
            fill = _role_fill(role, fill)
        
        lines.append(f'  {nid} [label="{full_label}", shape={shape}, style=filled, fillcolor="{fill}"];')
    
    lines.append("")
    
    # LAYOUT CONSTRAINTS for proper decision tree structure
    # 1. Start node at source rank (top)
    if start_node_idx is not None:
        start_nid = node_id_map.get(start_node_idx)
        if start_nid:
            lines.append(f"  {{ rank=source; {start_nid}; }}")
    
    # 2. End nodes at sink rank (bottom)
    if end_node_indices:
        end_nids = [node_id_map.get(i) for i in end_node_indices if node_id_map.get(i)]
        if end_nids:
            lines.append(f"  {{ rank=sink; {'; '.join(end_nids)}; }}")
    
    # 3. Group Decision nodes with their immediate branches for better alignment
    for dec_idx in decision_node_indices:
        dec_node = valid_nodes[dec_idx]
        branches = dec_node.get('branches', [])
        if len(branches) >= 2:
            # Get target nodes for each branch
            branch_targets = []
            for b in branches:
                t = b.get('target')
                if isinstance(t, (int, float)) and 0 <= int(t) < len(valid_nodes):
                    target_nid = node_id_map.get(int(t))
                    if target_nid:
                        branch_targets.append(target_nid)
            # Put branch targets at same rank for proper side-by-side branching
            if len(branch_targets) >= 2:
                lines.append(f"  {{ rank=same; {'; '.join(branch_targets)}; }}")
    
    lines.append("")
    
    # Add notes legend at the bottom
    if notes_list:
        legend_lines = ["NOTES:"]
        for note_num, note_text in notes_list:
            wrapped_note = _wrap_label(note_text, max_width=60)
            legend_lines.append(f"[{note_num}] {wrapped_note}")
        legend_text = _escape_label("\n".join(legend_lines))
        lines.append(f'  NotesLegend [label="{legend_text}", shape=box, style=filled, fillcolor="#B3D9FF", fontsize=10];')
        lines.append("  { rank=max; NotesLegend; }")
    
    lines.append("")
    
    # EDGES - Critical for proper branching visualization
    for i, n in enumerate(valid_nodes):
        src = node_id_map.get(i)
        if not src:
            continue
        
        ntype = n.get('type', 'Process')
        
        if ntype == 'Decision' and 'branches' in n:
            # Decision nodes: use explicit branch targets with labels
            for b in n.get('branches', []):
                t = b.get('target')
                lbl = _escape_label(b.get('label', ''))
                if isinstance(t, (int, float)) and 0 <= int(t) < len(valid_nodes):
                    dst = node_id_map.get(int(t))
                    if dst:
                        if lbl:
                            lines.append(f'  {src} -> {dst} [label="{lbl}"];')
                        else:
                            lines.append(f'  {src} -> {dst};')
        elif ntype == 'End':
            # End nodes are terminal - no outgoing edges
            pass
        else:
            # Process/Start nodes: check for explicit target first, then sequential
            explicit_target = n.get('target')
            if explicit_target is not None and isinstance(explicit_target, (int, float)):
                target_idx = int(explicit_target)
                if 0 <= target_idx < len(valid_nodes):
                    dst = node_id_map.get(target_idx)
                    if dst:
                        lines.append(f'  {src} -> {dst};')
            elif i + 1 < len(valid_nodes):
                # Default: connect to next node (sequential flow)
                next_node = valid_nodes[i + 1]
                # Skip if next node is an End (we're at a branch point) and this isn't explicitly targeted
                # This prevents linear connections that bypass decision logic
                if next_node.get('type') != 'End' or len(valid_nodes) <= i + 2:
                    dst = node_id_map.get(i + 1)
                    if dst:
                        lines.append(f'  {src} -> {dst};')
    
    lines.append("}")
    return "\n".join(lines)

def build_graphviz_from_nodes(nodes, orientation="TD"):
    """Build a graphviz.Digraph from nodes with clean decision tree layout.
    
    LAYOUT PRINCIPLES (same as dot_from_nodes):
    1. Start node at top (rank=source)
    2. Clear top-to-bottom flow
    3. Decision nodes create true branching with distinct paths
    4. End nodes at bottom (rank=sink)
    5. No swimlane clustering (disrupts natural flow)
    """
    if graphviz is None:
        return None
    valid_nodes = harden_nodes(nodes or [])
    
    # Identify special nodes for layout control
    start_node_idx = None
    end_node_indices = []
    decision_node_indices = []
    
    for i, n in enumerate(valid_nodes):
        ntype = n.get('type', 'Process')
        if ntype == 'Start' and start_node_idx is None:
            start_node_idx = i
        elif ntype == 'End':
            end_node_indices.append(i)
        elif ntype == 'Decision':
            decision_node_indices.append(i)
    
    rankdir = 'TB' if orientation == 'TD' else 'LR'
    
    g = graphviz.Digraph(format='svg')
    g.attr(rankdir=rankdir)
    g.attr(splines='ortho')      # Orthogonal edges for cleaner routing
    g.attr(nodesep='0.8')        # Horizontal spacing
    g.attr(ranksep='1.0')        # Vertical spacing
    g.attr('node', fontname='Helvetica', fontsize='11')
    g.attr('edge', fontname='Helvetica', fontsize='10')
    
    node_id_map = {}
    notes_list = []
    notes_node_map = {}
    
    # First pass: identify all nodes with notes
    note_counter = 1
    for i, n in enumerate(valid_nodes):
        notes_text = n.get('notes', '') or n.get('detail', '')
        if notes_text and str(notes_text).strip():
            notes_list.append((note_counter, str(notes_text).strip()))
            notes_node_map[i] = note_counter
            note_counter += 1
    
    # Create all nodes WITHOUT swimlane clustering
    for i, n in enumerate(valid_nodes):
        nid = f"N{i}"
        node_id_map[i] = nid
        
        raw_label = n.get('label', 'Step')
        raw_label = str(raw_label).replace('\\n', ' ').replace('\n', ' ')
        wrapped_label = _wrap_label(raw_label, width=25)
        
        if i in notes_node_map:
            wrapped_label = wrapped_label + f"\n(Note {notes_node_map[i]})"
        
        full_label = _escape_label(wrapped_label)
        ntype = n.get('type', 'Process')
        
        if ntype == 'Decision':
            shape, fill = 'diamond', '#F8CECC'
        elif ntype == 'Start':
            shape, fill = 'oval', '#D5E8D4'
        elif ntype == 'End':
            shape, fill = 'oval', '#D5E8D4'
        elif ntype == 'Reevaluation':
            shape, fill = 'box', '#FFCC80'
        else:
            shape, fill = 'box', '#FFF2CC'
        
        role = n.get('role', '')
        if role:
            fill = _role_fill(role, fill)
        
        g.node(nid, full_label, shape=shape, style='filled', fillcolor=fill)
    
    # LAYOUT CONSTRAINTS
    # 1. Start node at source rank
    if start_node_idx is not None:
        start_nid = node_id_map.get(start_node_idx)
        if start_nid:
            with g.subgraph() as s:
                s.attr(rank='source')
                s.node(start_nid)
    
    # 2. End nodes at sink rank
    if end_node_indices:
        with g.subgraph() as s:
            s.attr(rank='sink')
            for end_idx in end_node_indices:
                end_nid = node_id_map.get(end_idx)
                if end_nid:
                    s.node(end_nid)
    
    # 3. Group Decision branch targets at same rank for side-by-side branching
    for dec_idx in decision_node_indices:
        dec_node = valid_nodes[dec_idx]
        branches = dec_node.get('branches', [])
        if len(branches) >= 2:
            branch_targets = []
            for b in branches:
                t = b.get('target')
                if isinstance(t, (int, float)) and 0 <= int(t) < len(valid_nodes):
                    target_nid = node_id_map.get(int(t))
                    if target_nid:
                        branch_targets.append(target_nid)
            if len(branch_targets) >= 2:
                with g.subgraph() as s:
                    s.attr(rank='same')
                    for tnid in branch_targets:
                        s.node(tnid)
    
    # Add notes legend
    if notes_list:
        legend_lines = ["NOTES:"]
        for note_num, note_text in notes_list:
            wrapped_note = _wrap_label(note_text, max_width=60)
            legend_lines.append(f"[{note_num}] {wrapped_note}")
        legend_text = _escape_label("\n".join(legend_lines))
        g.node('NotesLegend', legend_text, shape='box', style='filled', fillcolor='#B3D9FF', fontsize='10')
        with g.subgraph() as s:
            s.attr(rank='max')
            s.node('NotesLegend')
    
    # EDGES - Critical for proper branching
    for i, n in enumerate(valid_nodes):
        src = node_id_map.get(i)
        if not src:
            continue
        
        ntype = n.get('type', 'Process')
        
        if ntype == 'Decision' and 'branches' in n:
            for b in n.get('branches', []):
                t = b.get('target')
                lbl = _escape_label(b.get('label', ''))
                if isinstance(t, (int, float)) and 0 <= int(t) < len(valid_nodes):
                    dst = node_id_map.get(int(t))
                    if dst:
                        if lbl:
                            g.edge(src, dst, label=lbl)
                        else:
                            g.edge(src, dst)
        elif ntype == 'End':
            pass  # Terminal - no outgoing edges
        else:
            explicit_target = n.get('target')
            if explicit_target is not None and isinstance(explicit_target, (int, float)):
                target_idx = int(explicit_target)
                if 0 <= target_idx < len(valid_nodes):
                    dst = node_id_map.get(target_idx)
                    if dst:
                        g.edge(src, dst)
            elif i + 1 < len(valid_nodes):
                next_node = valid_nodes[i + 1]
                if next_node.get('type') != 'End' or len(valid_nodes) <= i + 2:
                    dst = node_id_map.get(i + 1)
                    if dst:
                        g.edge(src, dst)
    
    return g

def render_graphviz_bytes(graph, fmt="svg"):
    """Render a graphviz.Digraph to bytes if possible, else return None."""
    if graphviz is None or graph is None:
        return None
    try:
        return graph.pipe(format=fmt)
    except Exception:
        return None

def get_smart_model_cascade(requires_vision=False, requires_json=False):
    """Return prioritized list of models for Auto mode based on task requirements.
    
    Model names use -latest aliases for automatic updates:
    - gemini-flash-latest: Primary model (maps to gemini-3-flash-preview after Jan 30, 2026)
    - gemini-pro-latest: Fallback (maps to gemini-3-pro-preview after Jan 30, 2026)
    
    See: https://ai.google.dev/gemini-api/docs/thought-signatures
    Thought signature validation required for Gemini 3+ models with function calling.
    
    Strategy: Try models from most to least sophisticated. Auto mode cascades through
    available models until quota is found. User-selected models fall back to alternatives.
    """
    if model_choice == "Auto":
        # Cascade using -latest aliases for automatic model updates
        return [
            "gemini-flash-latest",         # Primary: auto-updates to latest flash
            "gemini-pro-latest",           # Fallback: auto-updates to latest pro
        ]
    else:
        # Use user-selected model with intelligent fallback
        return [model_choice, "gemini-flash-latest", "gemini-pro-latest"]

def get_gemini_response(
    prompt, 
    json_mode=False, 
    stream_container=None, 
    image_data=None, 
    timeout=30,
    function_declaration=None,
    enable_thinking=True,
    thinking_budget=1024,
    contents=None
):
    """
    Send a prompt (with optional image) to Gemini and get a response.
    Supports native function calling and thought signature validation for Gemini 3+ models.
    
    Per official API: https://ai.google.dev/gemini-api/docs/api-key
    Thought signatures: https://ai.google.dev/gemini-api/docs/thought-signatures
    
    Args:
        prompt: Text prompt string
        json_mode: If True, extract JSON from response (legacy mode, prefer function_declaration)
        stream_container: Deprecated (v1 API)
        image_data: Optional dict with 'mime_type' and 'data' (base64 bytes) for image
        timeout: Seconds to wait per model before moving to next
        function_declaration: Optional FunctionDeclaration for native function calling
        enable_thinking: Enable thought signature validation (required for Gemini 3+ function calling)
        thinking_budget: Token budget for internal reasoning (256-4096)
        contents: Optional pre-built contents array (for file URIs, etc.)
    
    Returns:
        - If function_declaration provided: dict with 'function_name' and 'arguments'
        - If json_mode: parsed JSON dict/list
        - Otherwise: text string
    """
    client = get_genai_client()
    if not client:
        st.error("AI Error. Please check API Key.")
        return None

    # Smart cascade: if image provided, prioritize vision models
    candidates = get_smart_model_cascade(requires_vision=bool(image_data), requires_json=json_mode)

    # Build contents array per official API structure
    # https://ai.google.dev/gemini-api/docs/api-overview#request-body
    if contents:
        # Use pre-built contents (e.g., with file URIs)
        pass
    elif image_data:
        # Multimodal: text + image
        contents = [
            {
                "parts": [
                    {"text": prompt},
                    {
                        "inline_data": {
                            "mime_type": image_data.get('mime_type', 'image/jpeg'),
                            "data": image_data['data']
                        }
                    }
                ]
            }
        ]
    else:
        # Text-only: wrap in proper structure
        contents = [
            {
                "parts": [
                    {"text": prompt}
                ]
            }
        ]

    # Build generation config with thinking and optional function calling
    config_kwargs = {}
    
    # Configure thinking for thought signature validation (Gemini 3+ requirement)
    if enable_thinking:
        config_kwargs["thinking_config"] = types.ThinkingConfig(
            thinking_budget=thinking_budget
        )
    
    # Configure function calling if declaration provided
    if function_declaration:
        config_kwargs["tools"] = [types.Tool(function_declarations=[function_declaration])]
        # Force function calling mode
        config_kwargs["tool_config"] = types.ToolConfig(
            function_calling_config=types.FunctionCallingConfig(mode="AUTO")
        )
    
    config = types.GenerateContentConfig(**config_kwargs) if config_kwargs else None

    response = None
    last_error = None
    skipped_models = []
    
    for model_name in candidates:
        try:
            # Build API call arguments
            call_kwargs = {
                "model": model_name,
                "contents": contents,
            }
            if config:
                call_kwargs["config"] = config
            
            response = client.models.generate_content(**call_kwargs)
            
            # Check for function call response
            if function_declaration and response and response.candidates:
                result = extract_function_call_result(response)
                if result:
                    return result
            
            # Check if response has valid text
            if response and hasattr(response, 'text'):
                break
        except Exception as e:
            error_str = str(e)
            last_error = error_str

            # Check if error is quota exhaustion (429 RESOURCE_EXHAUSTED)
            if "429" in error_str or "RESOURCE_EXHAUSTED" in error_str:
                skipped_models.append(f"{model_name} (quota)")
            else:
                skipped_models.append(f"{model_name} (unavailable)")

            # Continue to next candidate
            time.sleep(0.3)
            continue

    if not response:
        # Silently return None - let the calling code handle the UI feedback
        return None

    try:
        # Access response.text per official API
        text = response.text if hasattr(response, 'text') else ""
        
        if not text:
            return None

        if json_mode:
            # Clean markdown code blocks (legacy JSON extraction mode)
            text = text.replace('```json', '').replace('```', '').strip()
            # Extract JSON object or array
            match = re.search(r'(\{[\s\S]*\}|\[[\s\S]*\])', text)
            if match:
                text = match.group(0)
            try:
                return json.loads(text)
            except Exception:
                return None
        return text
    except Exception:
        st.error("AI response parsing error. Please retry.")
        return None

@st.cache_data(ttl=3600)
def get_available_models(api_key):
    """Fetch list of available models from Gemini API.
    Uses client.models.list() per official SDK.
    https://ai.google.dev/gemini-api/docs/models
    """
    try:
        client = genai.Client(api_key=api_key)
        models = client.models.list()
        model_names = []
        for m in models:
            # Extract model name
            name = getattr(m, 'name', '')
            if name:
                # Model names come as 'models/gemini-2.5-flash' format
                # Extract just the model ID
                if 'models/' in name:
                    model_id = name.split('models/')[-1]
                    model_names.append(model_id)
                else:
                    model_names.append(name)
        return sorted(list(set(model_names))) if model_names else None
    except Exception as e:
        # Silently fail for caching purposes
        return None

def validate_ai_connection() -> bool:
    """Attempt a minimal generate_content call to verify the API key/model works.
    Returns True if a response is obtained, else False.
    Per official API: https://ai.google.dev/gemini-api/docs/quickstart
    """
    client = get_genai_client()
    if not client:
        return False
    try:
        # Use proper content structure per official API with thinking config
        resp = client.models.generate_content(
            model="gemini-flash-latest",
            contents=[{"parts": [{"text": "ping"}]}],
            config=get_generation_config(enable_thinking=True, thinking_budget=256)
        )
        return bool(resp and hasattr(resp, 'text') and resp.text)
    except Exception as e:
        # Store concise summary instead of verbose raw error
        err = str(e)
        if "RESOURCE_EXHAUSTED" in err or "429" in err or "quota" in err:
            st.session_state["ai_error"] = "Rate limit exceeded. Please try again later."
        elif "UNAUTHENTICATED" in err or "invalid" in err:
            st.session_state["ai_error"] = "Invalid API key. Check and retry."
        elif "PERMISSION_DENIED" in err:
            st.session_state["ai_error"] = "Permission denied for model. Try another model."
        else:
            st.session_state["ai_error"] = "AI service error. Please try again."
        return False

@st.cache_data(ttl=3600)
def search_pubmed(query):
    base_url = "https://eutils.ncbi.nlm.nih.gov/entrez/eutils/"
    try:
        search_params = {'db': 'pubmed', 'term': f"{query}", 'retmode': 'json', 'retmax': 50, 'sort': 'relevance'}
        url = base_url + "esearch.fcgi?" + urllib.parse.urlencode(search_params)
        with urllib.request.urlopen(url) as response: id_list = json.loads(response.read().decode()).get('esearchresult', {}).get('idlist', [])
        if not id_list: return []
        fetch_params = {'db': 'pubmed', 'id': ','.join(id_list), 'retmode': 'xml'}
        with urllib.request.urlopen(base_url + "efetch.fcgi?" + urllib.parse.urlencode(fetch_params)) as response: root = ET.fromstring(response.read().decode())
        citations = []
        for article in root.findall('.//PubmedArticle'):
            medline = article.find('MedlineCitation')
            pmid = medline.find('PMID').text
            title = medline.find('Article/ArticleTitle').text
            # Authors
            author_list = article.findall('.//Author')
            authors_str = "Unknown"
            if author_list:
                authors = []
                for auth in author_list[:3]:
                    lname = auth.find('LastName')
                    init = auth.find('Initials')
                    if lname is not None and init is not None: authors.append(f"{lname.text} {init.text}")
                authors_str = ", ".join(authors)
                if len(author_list) > 3: authors_str += ", et al."
            # Year
            year_node = article.find('.//PubDate/Year')
            year = year_node.text if year_node is not None else "N/A"
            # Journal
            journal = medline.find('Article/Journal/Title').text if medline.find('Article/Journal/Title') is not None else "N/A"
            # Abstract
            abs_node = medline.find('Article/Abstract')
            abstract_text = " ".join([e.text for e in abs_node.findall('AbstractText') if e.text]) if abs_node is not None else "No abstract."
            citations.append({
                "id": pmid, "title": title, "authors": authors_str, "year": year, "journal": journal,
                "url": f"https://pubmed.ncbi.nlm.nih.gov/{pmid}/", "abstract": abstract_text,
                "grade": "Un-graded", "rationale": "Not yet evaluated."
            })
        return citations
    except Exception as e: st.error(f"PubMed Search Error: {e}"); return []

def fetch_single_pmid(pmid):
    """Fetch metadata for a single PMID from PubMed. Returns evidence dict with default grade."""
    base_url = "https://eutils.ncbi.nlm.nih.gov/entrez/eutils/"
    try:
        # With rate limiting: respect NCBI 3 requests/second limit
        time.sleep(0.4)
        
        fetch_params = {'db': 'pubmed', 'id': pmid, 'retmode': 'xml'}
        with urllib.request.urlopen(base_url + "efetch.fcgi?" + urllib.parse.urlencode(fetch_params)) as response:
            root = ET.fromstring(response.read().decode())
        
        article = root.find('.//PubmedArticle')
        if article is None:
            return None
        
        medline = article.find('MedlineCitation')
        if medline is None:
            return None
        
        pmid_elem = medline.find('PMID')
        title_elem = medline.find('Article/ArticleTitle')
        
        if pmid_elem is None or title_elem is None:
            return None
        
        # Authors
        author_list = article.findall('.//Author')
        authors_str = "Unknown"
        if author_list:
            authors = []
            for auth in author_list[:3]:
                lname = auth.find('LastName')
                init = auth.find('Initials')
                if lname is not None and init is not None:
                    authors.append(f"{lname.text} {init.text}")
            authors_str = ", ".join(authors)
            if len(author_list) > 3:
                authors_str += ", et al."
        
        # Year
        year_node = article.find('.//PubDate/Year')
        year = year_node.text if year_node is not None else "N/A"
        
        # Journal
        journal = medline.find('Article/Journal/Title').text if medline.find('Article/Journal/Title') is not None else "N/A"
        
        # Abstract
        abs_node = medline.find('Article/Abstract')
        abstract_text = " ".join([e.text for e in abs_node.findall('AbstractText') if e.text]) if abs_node is not None else "No abstract."
        
        return {
            "id": pmid_elem.text,
            "title": title_elem.text,
            "authors": authors_str,
            "year": year,
            "journal": journal,
            "url": f"https://pubmed.ncbi.nlm.nih.gov/{pmid_elem.text}/",
            "abstract": abstract_text,
            "grade": "Un-graded",
            "rationale": "Auto-added; pending review",
            "source": "enriched_from_phase3"
        }
    except Exception as e:
        return None

def extract_pmids_from_nodes(nodes):
    """Extract all unique PMIDs from pathway nodes that are not 'N/A'."""
    pmids = set()
    for node in nodes:
        pmid = node.get('evidence', 'N/A')
        if pmid and pmid not in ['N/A', '', None]:
            pmids.add(str(pmid).strip())
    return pmids

def enrich_phase2_with_new_pmids(new_pmids, existing_pmids):
    """
    Compare Phase 3 PMIDs with Phase 2 PMIDs.
    Fetch and add any new PMIDs to Phase 2 evidence.
    Returns list of newly added evidence entries.
    """
    pmids_to_add = new_pmids - existing_pmids
    if not pmids_to_add:
        return []
    
    new_evidence = []
    for pmid in sorted(pmids_to_add):
        evidence_entry = fetch_single_pmid(pmid)
        if evidence_entry:
            new_evidence.append(evidence_entry)
    
    return new_evidence

def format_as_numbered_list(items):
    """Ensure numbered list formatting with a blank line between items.
    - Accepts list or string; outputs a string with "1. ..." and blank lines.
    """
    # If already a list, normalize and build with spacing
    if isinstance(items, list):
        clean_items = [re.sub(r'^\s*[\d\.\-\*]+\s*', '', str(item)).strip() for item in items if str(item).strip()]
        return "\n\n".join([f"{i+1}. {item}" for i, item in enumerate(clean_items)])

    # If it's a string, try to detect existing items and rebuild
    text = str(items or "").strip()
    if not text:
        return ""

    lines = [ln.rstrip() for ln in text.split("\n")]
    # Extract lines that look like items (start with number. or bullet)
    item_lines = []
    for ln in lines:
        m = re.match(r"^\s*(\d+\.|[-\*])\s+(.*)$", ln)
        if m:
            item_lines.append(m.group(2).strip())
        elif ln.strip():
            item_lines.append(ln.strip())
    # If we collected multiple logical items, rebuild with numbering + spacing
    if len(item_lines) >= 2:
        return "\n\n".join([f"{i+1}. {it}" for i, it in enumerate(item_lines)])

    # Otherwise, ensure at least single paragraph return
    # Also insert blank lines before any subsequent "N." occurrences
    text = re.sub(r"\n(?=(\d+\.)\s)", "\n\n", text)
    return text

def compute_textarea_height(text: str, min_rows: int = 10, max_rows: int = 100, line_px: int = 22, padding_px: int = 18, chars_per_line: int = 70) -> int:
    """
    Estimate a textarea height in pixels based on current text content.
    Streamlit doesn't auto-resize text_areas, so we approximate by rows.
    """
    if text is None:
        text = ""
    # Estimate visual lines including wrapping
    # Split on explicit newlines, then estimate wrapped rows per line
    est_lines = 0
    for block in text.split("\n\n"):
        for ln in block.split("\n"):
            length = len(ln.strip()) or 1
            est_lines += max(1, (length + chars_per_line - 1) // chars_per_line)
        # add one extra line between paragraphs for spacing
        est_lines += 1
    lines = max(min_rows, min(max_rows, est_lines or min_rows))
    return lines * line_px + padding_px

# ==========================================
# 3A. CHAT & FEEDBACK HELPER FUNCTIONS
# ==========================================

def initialize_chat_state():
    """Initialize chat messages in session state if not exists."""
    if "chat_messages" not in st.session_state:
        st.session_state["chat_messages"] = []
    if "chat_expanded" not in st.session_state:
        # Auto-open if query param ?chat=1 is present
        qp = _get_query_param("chat")
        st.session_state["chat_expanded"] = str(qp).lower() in ("1", "true", "yes", "y")


def save_feedback_response(rating: int, feedback_text: str, phase: str = ""):
    """Save feedback response as JSON to data/feedback/ folder."""
    import datetime
    os.makedirs("data/feedback", exist_ok=True)
    
    # Safely get pathway condition from session state
    pathway_condition = ""
    if hasattr(st.session_state, 'data') and isinstance(st.session_state.data, dict):
        pathway_condition = st.session_state.data.get("phase1", {}).get("condition", "")
    
    feedback_data = {
        "timestamp": datetime.datetime.utcnow().isoformat(),
        "rating": rating,
        "feedback_text": feedback_text,
        "phase": phase,
        "pathway_condition": pathway_condition
    }
    
    timestamp_ms = int(datetime.datetime.utcnow().timestamp() * 1000)
    filename = f"data/feedback/feedback_{timestamp_ms}.json"
    with open(filename, "w") as f:
        json.dump(feedback_data, f, indent=2)


def get_carepathiq_scoped_response(user_question: str) -> str:
    """
    Get a response from Gemini scoped strictly to CarePathIQ app questions.
    Enhanced with comprehensive knowledge of app structure, files, and capabilities.
    """
    # Try local FAQ first (lightweight, always works)
    response = get_local_faq_answer(user_question)
    if response:
        return response
    
    # Only try AI if client is available
    client = get_genai_client()
    if not client:
        return "Please enter your Gemini API key in the sidebar to enable AI-powered answers."
    
    scope_constraint = """You are the CarePathIQ AI Agent - an expert assistant with complete knowledge of the CarePathIQ clinical pathway development application.

**The 5 Phases in Detail:**

1. **Phase 1: Define Scope & Charter** — Input clinical condition and care setting. AI auto-drafts inclusion/exclusion criteria, problem statement, objectives, and population.
   
2. **Phase 2: Appraise Evidence** — Intelligent PubMed query generation with MeSH terms, field tags ([tiab], [mesh]), and AI optimization. Evidence table with PMID, abstract, and GRADE quality assessment. PICO framework support and interactive query refinement tools.
   
3. **Phase 3: Build Decision Tree** — Node-based pathway (decision, action, outcome nodes) with Graphviz rendering and AI-powered regeneration with user refinement.
   
4. **Phase 4: Design Interface** — Nielsen's 10 Usability Heuristics analysis with AI recommendations and visual preview of improvements.
   
5. **Phase 5: Operationalize** — Expert Panel Form, Beta Testing Form, Interactive Education (quizzes + certificates), Executive Summary. All standalone HTML/DOCX.

**Key Features:** Gemini API integration, admin mode (?admin=CODE), feedback system, CSV/DOCX/HTML/PNG/SVG exports, responsive pink/brown/teal design.

**Your Capabilities:** Explain phases/features, guide workflow, troubleshoot (API, models, exports), clarify inputs/outputs, explain Phase 5 deliverables.

**Response Style:** Concise 2-4 sentences, reference specific phases/features, conversational expert tone.

**User Question:** {user_question}"""

    response = get_gemini_response(scope_constraint)
    if not response:
        response = get_local_faq_answer(user_question)
    return response or "I'm not sure how to help with that. Please ask me about features in the CarePathIQ app!"


def get_local_faq_answer(user_question: str) -> str:
    """Provide lightweight built-in answers when AI is unavailable."""
    q = (user_question or "").strip().lower()
    if "5 phases" in q or "five phases" in q or "phases" in q:
        return (
            "CarePathIQ has 5 phases: 1) **Define Scope** — clarify condition, context, and goals. "
            "2) **Appraise Evidence** — gather and grade studies with structured PICO/MESH support. "
            "3) **Build Decision Tree** — design pathway logic and branches. "
            "4) **Design Interface** — preview and optimize using Nielsen's usability heuristics. "
            "5) **Operationalize** — export expert forms, beta testing, education modules, and executive summary."
        )
    if "decision tree" in q or "phase 3" in q or "nodes" in q or "pathway" in q:
        return (
            "In Phase 3, add nodes (decisions, actions, outcomes), connect them to form branches, "
            "and iterate using evidence and heuristics. Use the refinement box to request changes, "
            "then regenerate to update the pathway structure. The AI uses Phase 1 & 2 context to build logical flows."
        )
    if "evidence" in q or "phase 2" in q or "pubmed" in q or "mesh" in q:
        return (
            "Phase 2 enables intelligent PubMed searches with AI-powered query optimization. "
            "The system auto-generates optimized queries using MeSH terms and field tags ([tiab], [mesh]) for precise results. "
            "Include study PMID, title, abstract, GRADE quality assessment, and relevance notes. "
            "Features: proximity search guidance ([tiab:~N] syntax), AI query optimization, PICO framework, and CSV export. "
            "See PubMed Search Tips in the Refine search expander for advanced techniques."
        )
    if "phase 5" in q or "operationalize" in q or "export" in q or "expert" in q or "beta" in q:
        return (
            "Phase 5 generates 4 deliverables: Expert Panel Feedback Form (HTML), Beta Testing Form (HTML), "
            "Interactive Education Module (HTML with quizzes & certificates), and Executive Summary (Word doc). "
            "All are standalone files with no backend—users download, share, and collect responses via CSV."
        )
    if "heuristics" in q or "phase 4" in q or "usability" in q or "nielsen" in q:
        return (
            "Phase 4 applies Nielsen's 10 Usability Heuristics to your pathway. The AI analyzes each heuristic "
            "(visibility, error prevention, consistency, etc.) and suggests improvements. You can apply or reject "
            "each recommendation to optimize the user experience."
        )
    if "api" in q or "key" in q or "gemini" in q or "model" in q:
        return (
            "CarePathIQ uses Google Gemini API (get free key at https://aistudio.google.com/app/apikey). "
            "Uses gemini-flash-latest and gemini-pro-latest aliases for automatic model updates. "
            "Thought signature validation enabled for Gemini 3+ models. Enter your API key in the sidebar."
        )
    if "files" in q or "structure" in q or "code" in q:
        return (
            "Main files: streamlit_app.py (4200+ lines, 5-phase workflow), phase5_helpers.py (HTML generators), "
            "education_template.py (quiz system). Documentation: PHASE5_GUIDE.md, API_ALIGNMENT.md, README.md. "
            "Uses Streamlit, google-genai, Graphviz, python-docx."
        )
    return (
        "I can help with: the 5 phases workflow, building decision trees, evidence appraisal, "
        "usability heuristics, Phase 5 exports, API setup, or technical structure. What would you like to know?"
    )


def render_satisfaction_survey():
    """Render satisfaction survey with rating slider and feedback textarea."""
    with st.expander("📋 Feedback & Satisfaction Survey", expanded=False):
        st.markdown("### How satisfied are you with this app?")
        
        rating = st.slider(
            "Rating (10=Very Satisfied, 0=Not Satisfied)",
            min_value=0,
            max_value=10,
            value=5,
            step=1
        )
        
        feedback_text = st.text_area(
            "Additional feedback (optional)",
            placeholder="Tell us what you think...",
            height=100
        )
        
        if st.button("Submit Feedback"):
            phase = "Phase 3"
            save_feedback_response(rating, feedback_text, phase)
            st.success("✅ Thank you for your feedback!")


def load_admin_feedback_dashboard():
    """Admin-only dashboard to view and export feedback responses."""
    if not is_admin():
        return
    
    with st.expander("🔧 Admin: Feedback Dashboard", expanded=False):
        feedback_dir = "data/feedback"
        
        if not os.path.exists(feedback_dir):
            st.info("No feedback collected yet.")
            return
        
        feedback_files = [f for f in os.listdir(feedback_dir) if f.endswith(".json")]
        
        if not feedback_files:
            st.info("No feedback collected yet.")
            return
        
        st.markdown(f"**Total Responses:** {len(feedback_files)}")
        
        all_ratings = []
        for filename in feedback_files:
            try:
                with open(os.path.join(feedback_dir, filename)) as f:
                    data = json.load(f)
                    all_ratings.append(data.get("rating", 0))
            except Exception:
                continue
        
        if all_ratings:
            avg_rating = sum(all_ratings) / len(all_ratings)
            st.metric("Average Rating", f"{avg_rating:.1f}/10")
        
        st.markdown("### Feedback Responses")
        
        selected_file = st.selectbox(
            "View feedback:",
            feedback_files,
            format_func=lambda x: x.replace("feedback_", "").replace(".json", "")
        )
        
        if selected_file:
            try:
                with open(os.path.join(feedback_dir, selected_file)) as f:
                    feedback = json.load(f)
                    st.json(feedback)
            except Exception as e:
                st.error(f"Error loading feedback: {e}")
        
        if st.button("Export All Feedback as CSV"):
            import csv
            import io
            
            output = io.StringIO()
            writer = csv.writer(output)
            writer.writerow(["timestamp", "rating", "feedback", "phase", "pathway_condition"])
            
            for filename in feedback_files:
                try:
                    with open(os.path.join(feedback_dir, filename)) as f:
                        data = json.load(f)
                        writer.writerow([
                            data.get("timestamp", ""),
                            data.get("rating", ""),
                            data.get("feedback_text", ""),
                            data.get("phase", ""),
                            data.get("pathway_condition", "")
                        ])
                except Exception:
                    continue
            
            csv_data = output.getvalue()
            st.download_button(
                label="Download CSV",
                data=csv_data,
                file_name="feedback_export.csv",
                mime="text/csv"
            )

# ==========================================
# 3B. SIDEBAR & SESSION INITIALIZATION
# ==========================================
with st.sidebar:
    try:
        if "CarePathIQ_Logo.png" in os.listdir():
            with open("CarePathIQ_Logo.png", "rb") as f:
                logo_data = base64.b64encode(f.read()).decode()
            st.markdown(f"""<div style="text-align: center; margin-bottom: 20px;"><a href="https://carepathiq.org/" target="_blank"><img src="data:image/png;base64,{logo_data}" width="200" style="max-width: 100%;"></a></div>""", unsafe_allow_html=True)
    except Exception: pass

    st.divider()
    
    # API Key and Model - at top for easy access
    # Do not prefill from secrets so landing shows on first load
    gemini_api_key = st.text_input("Gemini API Key", value="", type="password", key="gemini_key")
    
    # Dynamically fetch available models from API
    available_models = []
    if gemini_api_key:
        available_models = get_available_models(gemini_api_key)
    
    model_options = ["Auto"] + (available_models if available_models else ["gemini-flash-latest", "gemini-pro-latest"])
    model_choice = st.selectbox("Model", model_options, index=0)
    
    st.divider()
    
    # Chat expander - plain text, no icon
    initialize_chat_state()
    with st.expander("Question about CarePathIQ?", expanded=False):
        # Initialize welcome message if this is the first load
        if not st.session_state.get("chat_initialized", False):
            st.session_state["chat_messages"] = [
                {
                    "role": "assistant",
                    "content": "Hi! Ask me anything about using CarePathIQ — the 5 phases, building pathways, using evidence, designing interfaces, or deploying your pathway."
                }
            ]
            st.session_state["chat_initialized"] = True
        
        # Display chat messages
        for message in st.session_state.get("chat_messages", []):
            with st.chat_message(message["role"]):
                st.write(message["content"])
        
        # Show suggested questions buttons (only if chat just started)
        if len(st.session_state.get("chat_messages", [])) <= 1:
            st.markdown("---")
            suggested_questions = [
                "What are the 5 phases?",
                "How do I add evidence in Phase 2?",
                "How do I create a decision tree in Phase 3?",
            ]
            
            for i, question in enumerate(suggested_questions):
                if st.button(question, key=f"suggested_q_{i}"):
                    # Append suggested question as user message
                    st.session_state["chat_messages"].append({
                        "role": "user",
                        "content": question
                    })
                    # Generate response
                    response = get_carepathiq_scoped_response(question)
                    st.session_state["chat_messages"].append({
                        "role": "assistant",
                        "content": response
                    })
                    st.rerun()
        
        st.markdown("---")
        
        # Chat input for user questions
        if user_input := st.chat_input("Ask about CarePathIQ..."):
            st.session_state["chat_messages"].append({
                "role": "user",
                "content": user_input
            })
            
            # Generate scoped response
            response = get_carepathiq_scoped_response(user_input)
            st.session_state["chat_messages"].append({
                "role": "assistant",
                "content": response
            })
            
            st.rerun()
    
    # Help us improve - plain text, no icon
    with st.expander("Help us improve!", expanded=False):
        st.markdown("How satisfied are you with this app?")
        
        rating = st.slider(
            "Rating",
            min_value=0,
            max_value=10,
            value=5,
            step=1,
            help="0=Not Satisfied, 10=Very Satisfied",
            key="sidebar_feedback_rating"
        )
        
        feedback_text = st.text_area(
            "Additional feedback (optional)",
            placeholder="Tell us what you think...",
            height=80,
            key="sidebar_feedback_text"
        )
        
        if st.button("Submit Feedback", key="sidebar_submit_feedback"):
            if rating or feedback_text.strip():
                save_feedback_response(rating, feedback_text, "sidebar")
                st.success("✅ Thank you for your feedback!")
            else:
                st.info("Please provide a rating or feedback")
    
    load_admin_feedback_dashboard()

    if gemini_api_key:
        try:
            st.session_state["genai_client"] = genai.Client(api_key=gemini_api_key)
            should_validate = st.session_state.get("last_tested_key") != gemini_api_key
            if should_validate:
                st.session_state["last_tested_key"] = gemini_api_key
                ok = validate_ai_connection()
                if ok:
                    st.success("AI Connected")
                    st.session_state["ai_valid"] = True
                    st.session_state.pop("ai_error", None)
                else:
                    # Show concise message only
                    err = st.session_state.get("ai_error")
                    short_msg = err or "API key invalid or model unavailable."
                    st.error(short_msg)
                    st.session_state["ai_valid"] = False
            else:
                # Preserve prior validation result
                if st.session_state.get("ai_valid"):
                    st.success("AI Connected")
                else:
                    st.info("Key entered — awaiting first AI call")
        except Exception as e:
            st.error(f"Failed to initialize Gemini client: {str(e)[:120]}")
            st.stop()

# LANDING PAGE LOGIC — SHOW WELCOME INSTEAD OF BLANK STOP
if not gemini_api_key:
    st.markdown(
        "<h2 style='color:#5D4037;font-style:italic;'>Intelligently Build & Deploy Care Pathways</h2>",
        unsafe_allow_html=True,
    )
    st.markdown(
        """
        <div style="
            background-color: #FFB0C9;
            border-left: 5px solid #5D4037;
            padding: 16px;
            border-radius: 6px;
            color: #3E2723;
            margin-bottom: 20px;">
            <strong>Welcome!</strong> Enter your <strong>Gemini API token</strong> on the left to activate the AI Agent. If you don't have one, you can get one for free <a href="https://aistudio.google.com/app/apikey" target="_blank"><strong>here</strong></a>.
        </div>
        """,
        unsafe_allow_html=True,
    )

    st.markdown(COPYRIGHT_HTML_FOOTER, unsafe_allow_html=True)
    st.stop()

if "current_phase_label" not in st.session_state:
    st.session_state.current_phase_label = PHASES[0]

if "data" not in st.session_state:
    st.session_state.data = {
        "phase1": {"condition": "", "setting": "", "inclusion": "", "exclusion": "", "problem": "", "objectives": "", "schedule": [], "population": ""},
        "phase2": {"evidence": [], "mesh_query": ""},
        "phase3": {"nodes": []},
        "phase4": {"heuristics_data": {}},
        "phase5": {"exec_summary": "", "beta_html": "", "expert_html": "", "edu_html": ""}
    }
if "suggestions" not in st.session_state:
    st.session_state.suggestions = {}
# --- FORCE MIGRATION FOR OLD DATA ---
if "pico_p" in st.session_state.data.get("phase2", {}):
    # Old PICO structure detected; clear Phase 2 data to force new layout
    st.session_state.data["phase2"] = {"evidence": [], "mesh_query": ""}

# ==========================================
# 4. MAIN WORKFLOW LOGIC
# ==========================================
# Display tagline at top (always visible when logged in)
st.markdown(
    "<h2 style='color:#5D4037;font-style:italic;'>Intelligently Build & Deploy Care Pathways</h2>",
    unsafe_allow_html=True,
)

## --- PHASE NAVIGATION ---

# Calculate progress for visual indicator
try:
    progress_value = calculate_granular_progress()
    progress_pct = int(progress_value * 100)
except:
    progress_value = 0.0
    progress_pct = 0

# Display progress bar
st.markdown(
    f"""
    <div class="phase-nav-container">
        <div class="phase-nav-label">Pathway Development Progress: {progress_pct}% Complete</div>
        <div class="phase-progress-bar">
            <div class="phase-progress-fill" style="width: {progress_pct}%;"></div>
        </div>
    </div>
    """,
    unsafe_allow_html=True
)

# Create horizontal navigation with phase numbers
phase = st.session_state.get("current_phase_label", PHASES[0])
current_phase_index = PHASES.index(phase) if phase in PHASES else 0

# Calculate completion status for each phase
data = st.session_state.get('data', {})
phase_completion = []

# Phase 1: Check if main fields are filled
p1 = data.get('phase1', {})
p1_complete = bool(p1.get('condition') and p1.get('setting') and p1.get('inclusion') and p1.get('exclusion'))
phase_completion.append(p1_complete)

# Phase 2: Check if evidence is collected
p2 = data.get('phase2', {})
p2_complete = bool(p2.get('evidence') and len(p2.get('evidence', [])) > 0)
phase_completion.append(p2_complete)

# Phase 3: Check if nodes exist
p3 = data.get('phase3', {})
p3_complete = bool(p3.get('nodes') and len(p3.get('nodes', [])) > 0)
phase_completion.append(p3_complete)

# Phase 4: Check if heuristics analyzed
p4 = data.get('phase4', {})
p4_complete = bool(p4.get('heuristics_data'))
phase_completion.append(p4_complete)

# Phase 5: Check if any deliverable generated
p5 = data.get('phase5', {})
p5_complete = bool(p5.get('beta_html') or p5.get('expert_html') or p5.get('edu_html'))
phase_completion.append(p5_complete)

# Compact navigation with numbered phases and status indicators
st.caption("**Phase**")

# Use ultra-short labels for navigation buttons
phase_short_labels = [
    "Scope",
    "Evidence",
    "Logic",
    "Design",
    "Deploy"
]

# Sidebar quick navigation mirroring the short labels
with st.sidebar:
    st.markdown("### Quick Navigation")
    for i, p in enumerate(PHASES):
        is_active = (i == current_phase_index)
        short_label = phase_short_labels[i]
        button_label = f"{i + 1}. {short_label}"
        button_type = "primary" if is_active else "secondary"
        if st.button(button_label, key=f"side_nav_{p.replace(' ', '_').replace('&', 'and')}", type=button_type):
            st.session_state.current_phase_label = p
            st.rerun()

# Create columns with arrows between buttons for forward flow visualization
# Pattern: [button] → [button] → [button] → [button] → [button]
num_buttons = len(PHASES)
num_arrows = num_buttons - 1
# Create columns: button, arrow, button, arrow, ..., button
col_specs = []
for i in range(num_buttons):
    col_specs.append(3)  # Button column (wider)
    if i < num_arrows:
        col_specs.append(0.5)  # Arrow column (narrow)

nav_cols = st.columns(col_specs)

col_idx = 0
for i, p in enumerate(PHASES):
    is_active = (i == current_phase_index)
    is_complete = phase_completion[i] if i < len(phase_completion) else False
    button_type = "primary" if is_active else "secondary"
    
    # Create button label with phase number using shorter label
    phase_num = i + 1
    short_label = phase_short_labels[i]
    button_label = f"{phase_num}. {short_label}"
    
    with nav_cols[col_idx]:
        if st.button(button_label, key=f"nav_{p.replace(' ', '_').replace('&', 'and')}", type=button_type, use_container_width=True):
            st.session_state.current_phase_label = p
            st.rerun()
    
    col_idx += 1
    
    # Add arrow between buttons (except after the last button)
    if i < num_arrows:
        with nav_cols[col_idx]:
            st.markdown("<div style='text-align: center; padding-top: 4px; font-size: 20px; color: #666;'>→</div>", unsafe_allow_html=True)
        col_idx += 1

st.markdown("---")

# --- PHASE 1 ---
if "Scope" in phase:
    # 1) Draft helper
    def trigger_p1_draft():
        """Run AI draft - call this from main script, not from callback."""
        c = st.session_state.get('p1_cond_input', '').strip()
        s = st.session_state.get('p1_setting', '').strip()
        if not (c and s):
            return False
        prompt = f"""
        Act as a Chief Medical Officer creating a clinical care pathway. For "{c}" in "{s}", return a JSON object with exactly these keys: inclusion, exclusion, problem, objectives.
        
        CRITICAL REQUIREMENTS:
        - inclusion: ONLY 3-5 brief patient characteristics that INCLUDE them in the pathway (e.g., age range, presentation type, risk factors). Concise phrases, not detailed descriptions.
        - exclusion: ONLY 3-5 brief characteristics that EXCLUDE patients (e.g., contraindications, alternative diagnoses, comorbidities). Concise phrases, not detailed descriptions.
        - problem: One brief clinical problem statement (1-2 sentences). Describe the gap or challenge, not educational content.
        - objectives: ONLY 3-4 brief clinical objectives for the pathway (e.g., "Reduce time to diagnosis", "Standardize treatment decisions"). Short statements, not detailed goals.
        
        Format each list as a simple newline-separated text, NOT as a JSON array. Do not use markdown formatting (no asterisks, dashes for bullets). Use plain text only.
        """
        try:
            with st.spinner("Generating pathway scope..."):
                result = get_gemini_response(
                    prompt, 
                    function_declaration=DEFINE_PATHWAY_SCOPE,
                    thinking_budget=1024
                )
                # Extract from function call or fall back
                if isinstance(result, dict) and 'arguments' in result:
                    data = result['arguments']
                elif isinstance(result, dict):
                    data = result
                else:
                    data = get_gemini_response(prompt, json_mode=True)
            if data and isinstance(data, dict):
                st.session_state.data['phase1']['inclusion'] = format_as_numbered_list(data.get('inclusion', ''))
                st.session_state.data['phase1']['exclusion'] = format_as_numbered_list(data.get('exclusion', ''))
                st.session_state.data['phase1']['problem'] = str(data.get('problem', ''))
                st.session_state.data['phase1']['objectives'] = format_as_numbered_list(data.get('objectives', ''))
                return True
        except Exception as e:
            st.error(f"Failed to generate pathway scope: {e}")
        return False

    # 2) Sync helpers
    def sync_p1_widgets():
        st.session_state.data['phase1']['condition'] = st.session_state.get('p1_cond_input', '')
        st.session_state.data['phase1']['setting'] = st.session_state.get('p1_setting', '')
        st.session_state.data['phase1']['inclusion'] = st.session_state.get('p1_inc', '')
        st.session_state.data['phase1']['exclusion'] = st.session_state.get('p1_exc', '')
        st.session_state.data['phase1']['problem'] = st.session_state.get('p1_prob', '')
        st.session_state.data['phase1']['objectives'] = st.session_state.get('p1_obj', '')

    def sync_and_request_draft():
        """Callback: sync values and set flag for main script to run AI draft."""
        sync_p1_widgets()
        # Set flag - main script will detect this and run the AI call
        st.session_state['p1_needs_draft'] = True

    # 3) Seed widget values
    p1 = st.session_state.data['phase1']
    st.session_state.setdefault('p1_cond_input', p1.get('condition', ''))
    st.session_state.setdefault('p1_setting',    p1.get('setting', ''))
    st.session_state.setdefault('p1_inc',        p1.get('inclusion', ''))
    st.session_state.setdefault('p1_exc',        p1.get('exclusion', ''))
    st.session_state.setdefault('p1_prob',       p1.get('problem', ''))
    st.session_state.setdefault('p1_obj',        p1.get('objectives', ''))

    # 4) CHECK FLAG AND RUN AI DRAFT (in main script, not callback - allows spinner)
    if st.session_state.get('p1_needs_draft'):
        st.session_state['p1_needs_draft'] = False  # Clear flag immediately
        if trigger_p1_draft():
            # Copy results to widget keys so they display
            p1 = st.session_state.data['phase1']
            st.session_state['p1_inc'] = p1.get('inclusion', '')
            st.session_state['p1_exc'] = p1.get('exclusion', '')
            st.session_state['p1_prob'] = p1.get('problem', '')
            st.session_state['p1_obj'] = p1.get('objectives', '')
            st.rerun()  # Rerun to display new values in text areas

    # 5) UI: Inputs
    st.header(f"Phase 1. {PHASES[0]}")
    styled_info("<b>Tip:</b> Enter Clinical Condition and Care Setting; the agent generates the rest automatically.")

    col1, col2 = columns_top(2)
    with col1:
        st.subheader("1. Clinical Focus")
        st.text_input("Clinical Condition", placeholder="e.g., Chest Pain", key="p1_cond_input", on_change=sync_and_request_draft)
        st.text_input("Care Setting", placeholder="e.g., Emergency Department", key="p1_setting", on_change=sync_and_request_draft)

        st.subheader("2. Target Population")
        st.text_area("Inclusion Criteria", key="p1_inc", height=compute_textarea_height(st.session_state.get('p1_inc',''), 14), on_change=sync_p1_widgets)
        st.text_area("Exclusion Criteria", key="p1_exc", height=compute_textarea_height(st.session_state.get('p1_exc',''), 14), on_change=sync_p1_widgets)

    with col2:
        st.subheader("3. Clinical Gap / Problem Statement")
        st.text_area("Problem Statement / Clinical Gap", key="p1_prob", height=compute_textarea_height(st.session_state.get('p1_prob',''), 12), on_change=sync_p1_widgets, label_visibility="collapsed")

        st.subheader("4. Goals")
        st.text_area("Project Goals", key="p1_obj", height=compute_textarea_height(st.session_state.get('p1_obj',''), 14), on_change=sync_p1_widgets, label_visibility="collapsed")

    st.divider()
    st.subheader("5. Project Timeline (Gantt Chart)")
    styled_info("<b>Tip:</b> Hover over the top right of the chart to download the image or table. You can also directly edit the timeline table below to adjust start/end dates and task owners.")
    if not st.session_state.data['phase1']['schedule']:
        today = date.today()
        def add_weeks(start, w): return start + timedelta(weeks=w)
        d1 = add_weeks(today, 2); d2 = add_weeks(d1, 4); d3 = add_weeks(d2, 2); d4 = add_weeks(d3, 2)
        d5 = add_weeks(d4, 4); d6 = add_weeks(d5, 4); d7 = add_weeks(d6, 2); d8 = add_weeks(d7, 4)
        st.session_state.data['phase1']['schedule'] = [
            {"Stage": "1. Project Charter", "Owner": "PM", "Start": today, "End": d1},
            {"Stage": "2. Pathway Draft", "Owner": "Clinical Lead", "Start": d1, "End": d2},
            {"Stage": "3. Expert Panel", "Owner": "Expert Panel", "Start": d2, "End": d3},
            {"Stage": "4. Iterative Design", "Owner": "Clinical Lead", "Start": d3, "End": d4},
            {"Stage": "5. Informatics Build", "Owner": "IT", "Start": d4, "End": d5},
            {"Stage": "6. Beta Testing", "Owner": "Quality", "Start": d5, "End": d6},
            {"Stage": "7. Go-Live", "Owner": "Ops", "Start": d6, "End": d7},
            {"Stage": "8. Optimization", "Owner": "Clinical Lead", "Start": d7, "End": d8},
            {"Stage": "9. Monitoring", "Owner": "Quality", "Start": d8, "End": add_weeks(d8, 12)}
        ]
    df_sched = pd.DataFrame(st.session_state.data['phase1']['schedule'])
    edited_sched = st.data_editor(df_sched, num_rows="dynamic", width="stretch", key="sched_editor", column_config={"Stage": st.column_config.TextColumn("Stage", width="medium")})
    if not edited_sched.empty:
        st.session_state.data['phase1']['schedule'] = edited_sched.to_dict('records')
        chart_data = edited_sched.copy()
        chart_data.dropna(subset=['Start', 'End', 'Stage'], inplace=True)
        chart_data['Start'] = pd.to_datetime(chart_data['Start'])
        chart_data['End'] = pd.to_datetime(chart_data['End'])
        if not chart_data.empty:
            # Define consistent color scheme for all owners
            owner_colors = {
                'PM': '#5f9ea0',           # Cadet blue
                'Clinical Lead': '#4169e1',  # Royal blue
                'Expert Panel': '#b0c4de',   # Light steel blue
                'IT': '#dc143c',             # Crimson
                'Ops': '#ffb6c1',            # Light pink
                'Quality': '#5D4037'         # Brown (consistent with brand)
            }
            
            chart = alt.Chart(chart_data).mark_bar().encode(
                x=alt.X('Start', title='Date'),
                x2='End',
                y=alt.Y('Stage', sort=None),
                color=alt.Color('Owner', 
                    scale=alt.Scale(
                        domain=list(owner_colors.keys()),
                        range=list(owner_colors.values())
                    ),
                    legend=alt.Legend(title='Owner')
                ),
                tooltip=['Stage', 'Start', 'End', 'Owner']
            ).properties(height=300).interactive()
            st.altair_chart(chart, width="stretch")
    else:
        st.info("Add timeline entries to see the Gantt chart.")

    st.divider()
    
    # Download Project Charter - generates on-the-fly when clicked
    sync_p1_widgets()
    d = st.session_state.data['phase1']
    
    if d.get('condition') and d.get('problem'):
        if st.button("Download Project Charter (.docx)", type="secondary"):
            with st.status("Generating Project Charter...", expanded=True) as status:
                st.write("Building project charter based on IHI Quality Improvement framework...")
                p_ihi = f"""You are a Quality Improvement Advisor using IHI's Model for Improvement.

Build a project charter for managing "{d['condition']}" in "{d.get('setting', '') or 'care setting'}".

**Phase 1 Context:**
Problem: {d['problem']}
Inclusion: {d['inclusion']}
Exclusion: {d['exclusion']}
Objectives: {d.get('objectives', '')}

**Required JSON Output** (clean, no nested objects, all values as strings or string arrays):
Return ONLY a JSON object with these exact keys:
- problem: narrative (string, 1–2 paragraphs, use Phase 1 problem statement)
- project_description: scope statement (string, 1 paragraph: 1 sentence scope + 1 sentence approach)
- rationale: evidence basis (string, 1 paragraph, use Phase 1 context)
- expected_outcomes: list of strings (5–8 outcomes, one string per item)
- aim_statement: SMART goal (string, 1 line)
- outcome_measures: list of strings (safety/efficiency metrics, 4–6 items)
- process_measures: list of strings (protocol adherence, 4–6 items)
- balancing_measures: list of strings (staff impact, 3–4 items)
- initial_activities: narrative (string, 1 paragraph)
- change_ideas: list of strings (4–6 PDSA-ready interventions)
- stakeholders: comma-separated string
- barriers: list of strings (3–5 barriers)
- boundaries: object with in_scope (string) and out_of_scope (string)

Return clean JSON ONLY. No markdown, no explanation."""
                # Use native function calling for reliable structured output
                result = get_gemini_response(
                    p_ihi, 
                    function_declaration=CREATE_IHI_CHARTER,
                    thinking_budget=1024
                )
                if isinstance(result, dict) and 'arguments' in result:
                    res = result['arguments']
                elif isinstance(result, dict):
                    res = result
                else:
                    res = get_gemini_response(p_ihi, json_mode=True)
                if res:
                    st.session_state.data['phase1']['ihi_content'] = res
                    doc = create_word_docx(st.session_state.data['phase1'])
                    if doc:
                        status.update(label="Ready!", state="complete")
                        st.download_button("Download Project Charter (.docx)", doc, f"Project_Charter_{d['condition']}.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                    else:
                        status.update(label="Word export unavailable. Please ensure python-docx is installed.", state="error")
                else:
                    status.update(label="Failed to generate charter.", state="error")
    else:
        st.info("Complete Clinical Condition and Problem Statement to download the Project Charter.")

    # Refine & Regenerate (placed below charter so edits can be made, then re-generate)
    st.divider()
    submitted = False
    with st.expander("Refine & Regenerate", expanded=False):
        st.markdown("**Tip:** Describe any desired modifications and optionally attach supporting documents. Click \"Regenerate\" to automatically update all Phase 1 content and downloads")
        with st.form("p1_refine_form"):
            col_text, col_file = columns_top([2, 1])
            with col_text:
                st.text_area(
                    "Refinement Notes",
                    key="p1_refine_input",
                    placeholder="Clarify inclusion criteria; tighten scope; align objectives",
                    height=90,
                    help="Describe what to change. After applying, click Generate Project Charter above again."
                )

            with col_file:
                st.caption("Supporting Documents (optional)")
                p1_uploaded = st.file_uploader(
                    "Drag & drop or browse",
                    key="p1_file_upload",
                    accept_multiple_files=True,
                    label_visibility="collapsed",
                    help="Attach PDFs/DOCX files to inform pathway scope generation."
                )
                if p1_uploaded:
                    for uploaded_file in p1_uploaded:
                        file_result = upload_and_review_file(uploaded_file, f"p1_refine_{uploaded_file.name}", "clinical scope and charter")
                        if file_result:
                            st.success(f"✓ {file_result['filename']} uploaded")

            col1, col2 = st.columns([3, 1])
            with col2:
                submitted = st.form_submit_button("Regenerate", type="secondary", use_container_width=True)

    if submitted:
        refinement_text = st.session_state.get('p1_refine_input', '').strip()
        
        # Collect all uploaded file URIs for Gemini context
        file_uris = []
        file_texts = []
        for key in st.session_state.keys():
            if key.startswith("file_p1_refine_") and "_name" not in key:
                val = st.session_state.get(key, '')
                if val:
                    # Check if it's a file URI reference or extracted text
                    if "(" in val and val.endswith(")"):
                        # Extract URI from "File: name (uri)" format
                        uri = val.split("(")[-1].rstrip(")")
                        if uri.startswith("https://"):
                            file_uris.append(uri)
                    else:
                        # It's extracted text content
                        file_texts.append(val)
        
        # Build context from files
        has_files = bool(file_uris or file_texts)
        
        if refinement_text or has_files:
            current = st.session_state.data['phase1']
            
            # Build prompt with file context
            file_context = ""
            if file_texts:
                file_context = "\n\nContent from uploaded documents:\n" + "\n\n---\n\n".join(file_texts[:3])  # Limit to 3 docs
            
            user_input = refinement_text if refinement_text else "Use the uploaded documents to refine the pathway scope."
            
            prompt = f"""
            Update the following sections based on this user feedback: "{user_input}"
            {file_context}
            Current Data JSON: {json.dumps({k: current[k] for k in ['inclusion','exclusion','problem','objectives']})}
            
            IMPORTANT: Use information from the uploaded documents to make the inclusion criteria, exclusion criteria, problem statement, and objectives more specific and evidence-based.
            
            Return JSON with keys inclusion, exclusion, problem, objectives (use numbered lists where applicable).
            Do not use markdown formatting (no asterisks for bold). Use plain text only.
            """
            
            # Build contents with file URIs if available
            contents = None
            if file_uris:
                parts = [{"text": prompt}]
                for uri in file_uris[:3]:  # Limit to 3 files
                    parts.append({"file_data": {"file_uri": uri}})
                contents = [{"parts": parts}]
            
            with ai_activity("Applying refinements from documents…"):
                # Use native function calling for reliable structured output
                result = get_gemini_response(
                    prompt, 
                    function_declaration=DEFINE_PATHWAY_SCOPE,
                    thinking_budget=1024,
                    contents=contents
                )
                if isinstance(result, dict) and 'arguments' in result:
                    data = result['arguments']
                elif isinstance(result, dict):
                    data = result
                else:
                    data = get_gemini_response(prompt, json_mode=True, contents=contents)
            if data and isinstance(data, dict):
                st.session_state.data['phase1']['inclusion'] = format_as_numbered_list(data.get('inclusion', ''))
                st.session_state.data['phase1']['exclusion'] = format_as_numbered_list(data.get('exclusion', ''))
                st.session_state.data['phase1']['problem'] = str(data.get('problem', ''))
                st.session_state.data['phase1']['objectives'] = format_as_numbered_list(data.get('objectives', ''))
                # Auto-generate charter
                generate_p1_charter()
                st.success("Refinements applied and Project Charter auto-generated!")
            else:
                st.error("Failed to apply refinements. Please try again.")
    render_bottom_navigation()
    st.stop()

# --- PHASE 2 ---
elif "Evidence" in phase or "Appraise" in phase:
    st.header(f"Phase 2. {PHASES[1]}")

    # Build robust default query from Phase 1 if none saved
    # Format: "managing patients with [clinical condition] in [care setting]" using PubMed syntax
    default_q = st.session_state.data['phase2'].get('mesh_query', '')
    if not default_q and st.session_state.data['phase1']['condition']:
        c = st.session_state.data['phase1']['condition']
        s = st.session_state.data['phase1']['setting']
        
        # Use AI to build intelligent PubMed query with proximity searching
        client = get_genai_client()
        if client:
            proximity_prompt = f"""Create an optimized PubMed search query for clinical pathways and evidence-based practice guidelines.

**Clinical Context:**
- Condition: {c}
- Care Setting: {s if s else 'general care'}

**PubMed Search Syntax (CRITICAL - Use ONLY These Valid Operators):**
1. Phrase searching: Use quotes for exact phrases: "exact phrase"
2. Proximity searching: "term1 term2"[tiab:~N] where N = max words between terms
   - Example: "diabetes management"[tiab:~3] finds terms within 3 words of each other
   - Only works with [ti], [tiab], [ad] fields
3. Boolean operators: AND, OR, NOT (use parentheses for grouping)
4. Field tags (case-insensitive):
   - [mesh] or [mh] or [MeSH Terms] for MeSH headings
   - [tiab] or [Title/Abstract] for title/abstract
   - [ti] or [Title] for title only
   - [pt] for Publication Type
   - [lang] for Language
5. Publication type filters: Practice Guideline[pt], Review[pt], Systematic Review[pt]
6. Date filter: "last 5 years"[dp]
7. Language filter: english[lang]
8. DO NOT USE: NEAR, ADJ, NEAR/N - these are NOT valid PubMed syntax

**Requirements:**
- Use MeSH Terms: "{c}"[MeSH Terms]
- For proximity, use: "term1 term2"[tiab:~N] format
- Search for guidelines: "clinical pathway"[tiab] OR Practice Guideline[pt] OR "care protocol"[tiab]
- Return ONLY the raw query string, no explanations
- Example: ("diabetes mellitus"[MeSH Terms] OR "diabetes management"[tiab:~3]) AND ("clinical pathway"[tiab] OR Practice Guideline[pt]) AND english[lang]
- IMPORTANT: Every query MUST end with AND english[lang]"""
            
            with ai_activity("Building intelligent PubMed query..."):
                ai_query = get_gemini_response(proximity_prompt)
                if ai_query and isinstance(ai_query, str) and len(ai_query.strip()) > 10:
                    default_q = ai_query.strip()
                    # Validate query doesn't contain invalid operators (NEAR, ADJ are NOT valid PubMed syntax)
                    if 'ADJ' in default_q.upper() or 'NEAR' in default_q.upper():
                        # Fallback if AI generated invalid syntax
                        st.warning("AI-generated query contained unsupported PubMed operators (NEAR/ADJ are not valid). Using fallback query.")
                        default_q = None
                else:
                    # Fallback to basic query with proper PubMed syntax
                    cond_q = f'"{c}"[MeSH Terms]'
                    if s:
                        default_q = f'({cond_q} OR "{c}"[tiab]) AND ("{s}"[tiab]) AND ("clinical pathway"[tiab] OR Practice Guideline[pt] OR protocol[tiab]) AND english[lang]'
                    else:
                        default_q = f'({cond_q} OR "{c}"[tiab]) AND ("clinical pathway"[tiab] OR Practice Guideline[pt] OR protocol[tiab]) AND english[lang]'
        else:
            # Fallback if no AI client - use proper PubMed syntax
            cond_q = f'"{c}"[MeSH Terms]'
            if s:
                default_q = f'({cond_q} OR "{c}"[tiab]) AND ("{s}"[tiab]) AND ("clinical pathway"[tiab] OR Practice Guideline[pt] OR protocol[tiab]) AND english[lang]'
            else:
                default_q = f'({cond_q} OR "{c}"[tiab]) AND ("clinical pathway"[tiab] OR Practice Guideline[pt] OR protocol[tiab]) AND english[lang]'

    # Auto-run search once per distinct default query when evidence is empty
    if (
        not st.session_state.data['phase2']['evidence']
        and default_q
        and st.session_state.get('p2_last_autorun_query') != default_q
    ):
        st.session_state.data['phase2']['mesh_query'] = f"{default_q} AND (\"last 5 years\"[dp])"
        full_query = st.session_state.data['phase2']['mesh_query']
        with ai_activity("Searching PubMed and auto‑grading…"):
            results = search_pubmed(full_query)
            st.session_state.data['phase2']['evidence'] = results
            if results:
                prompt = (
                    "Assign GRADE quality of evidence (use EXACTLY one of: 'High (A)', 'Moderate (B)', 'Low (C)', or 'Very Low (D)') "
                    "and provide a brief Rationale (1-2 sentences) for each article. "
                    f"{json.dumps([{k:v for k,v in e.items() if k in ['id','title']} for e in results])}. "
                    "Return ONLY valid JSON object where keys are PMID strings and values are objects with 'grade' and 'rationale' fields. "
                    "Example: {\"12345678\": {\"grade\": \"High (A)\", \"rationale\": \"text here\"}}"
                )
                # Use native function calling for reliable structured output
                result = get_gemini_response(
                    prompt, 
                    function_declaration=GRADE_EVIDENCE,
                    thinking_budget=1024
                )
                if isinstance(result, dict) and 'arguments' in result:
                    grades = result['arguments'].get('grades', {})
                elif isinstance(result, dict) and 'grades' not in result:
                    grades = result
                else:
                    grades = get_gemini_response(prompt, json_mode=True)
                if grades and isinstance(grades, dict):
                    for e in st.session_state.data['phase2']['evidence']:
                        pmid_str = str(e['id'])
                        if pmid_str in grades:
                            grade_data = grades[pmid_str]
                            e['grade'] = grade_data.get('grade', 'Un-graded') if isinstance(grade_data, dict) else 'Un-graded'
                            e['rationale'] = grade_data.get('rationale', 'Not provided.') if isinstance(grade_data, dict) else 'Not provided.'
        st.session_state['p2_last_autorun_query'] = st.session_state.data['phase2']['mesh_query']

    # Summary banner for newly enriched evidence from Phase 3
    try:
        ev_list = st.session_state.data.get('phase2', {}).get('evidence', [])
        new_count = sum(1 for e in ev_list if bool(e.get('is_new')) or e.get('source') == 'enriched_from_phase3')
    except Exception:
        new_count = 0
    if new_count > 0:
        styled_info(f"<b>New evidence added from Phase 3:</b> {new_count} item(s) auto‑graded and added below. Use 'Show only new evidence' to focus review.")

    # Refinement with the current query prefilled
    with st.expander("Refine search", expanded=False):
        # Add proximity search guidance
        with st.expander("📚 PubMed Search Tips", expanded=False):
            st.markdown("""
            **Advanced PubMed Query Techniques:**
            
            🔍 **Proximity Searching** (find terms near each other):
            - Format: `"term1 term2"[field:~N]` where N = max words between terms
            - Only works with `[ti]`, `[tiab]`, or `[ad]` fields
            - Example: `"diabetes management"[tiab:~3]` finds terms within 3 words
            - Example: `"clinical pathway"[tiab:~0]` finds adjacent terms (any order)
            
            📋 **Field Tags** (search specific parts):
            - `[MeSH Terms]` or `[mesh]` — Medical Subject Headings
            - `[tiab]` or `[Title/Abstract]` — Title or abstract text
            - `[ti]` or `[Title]` — Title only
            - `[pt]` — Publication type (e.g., `Practice Guideline[pt]`)
            
            🎯 **Useful Filters:**
            - `english[lang]` — English only
            - `"last 5 years"[dp]` — Date published (auto-added)
            - `systematic[sb]` — Systematic reviews
            - `meta-analysis[pt]` — Meta-analyses
            
            💡 **Example Queries:**
            ```
            ("sepsis"[MeSH Terms] AND "emergency department"[tiab]) 
            AND ("clinical pathway"[tiab] OR Practice Guideline[pt])
            
            ("heart failure"[MeSH Terms] AND "care protocol"[tiab:~5]) 
            AND english[lang]
            ```
            
            ⚠️ **Note:** NEAR/N and ADJ operators are NOT valid PubMed syntax. Use `[field:~N]` for proximity.
            
            📖 **Full Guide:** [PubMed Help](https://pubmed.ncbi.nlm.nih.gov/help/#proximity-searching)
            """)
        
        current_q = st.session_state.data['phase2'].get('mesh_query', default_q)
        current_q_full = current_q or ""
        if current_q_full and '"last 5 years"[dp]' not in current_q_full:
            current_q_full = f"{current_q_full} AND (\"last 5 years\"[dp])"
        q = st.text_input(
            "PubMed Search Query (editable full query)",
            value=current_q_full,
            placeholder="Use MeSH Terms or [tiab] field tags for better results",
            key="p2_query_input",
        )
        q_clean = (q or "").strip()

        def ensure_time_filter(term: str) -> str:
            return term if '"last 5 years"[dp]' in term else f"{term} AND (\"last 5 years\"[dp])"

        def auto_enhance_query_with_proximity(query: str) -> str:
            """Automatically enhance user query with proximity operators if not already present."""
            # Check if query already has proximity operators (correct PubMed format: [field:~N])
            if ':~' in query:
                return query  # Already has proximity, don't modify
            
            # Check if it's a simple query that could benefit from enhancement
            c = st.session_state.data['phase1'].get('condition', '')
            s = st.session_state.data['phase1'].get('setting', '')
            
            if not c:
                return query  # Can't enhance without context
            
            # Remove any invalid NEAR/ADJ operators that may have been entered
            if 'NEAR' in query.upper() or 'ADJ' in query.upper():
                # These are invalid - strip them and rebuild
                query = query.replace('NEAR/', ' ').replace('ADJ', ' ')
                import re
                query = re.sub(r'\s+', ' ', query)
            
            optimize_prompt = f"""Optimize this PubMed query for better precision.

**Current Query:**
{query}

**Context:**
- Condition: {c}
- Setting: {s}

**Valid PubMed Syntax ONLY:**
1. Proximity: "term1 term2"[tiab:~N] where N = max words between terms
   - Example: "diabetes management"[tiab:~3]
2. MeSH: "term"[MeSH Terms] or "term"[mesh]
3. Field tags: [tiab], [ti], [pt], [lang]
4. Boolean: AND, OR, NOT with parentheses
5. DO NOT USE: NEAR, ADJ, NEAR/N - these are INVALID

**Requirements:**
- Return ONLY the query string, no explanations
- Use valid PubMed proximity syntax [tiab:~N] if beneficial
- If query is already well-formed, return it unchanged

**Example:**
Input: ("diabetes"[MeSH Terms]) AND (pathway OR guideline)
Output: ("diabetes"[MeSH Terms]) AND ("clinical pathway"[tiab] OR Practice Guideline[pt])"""
            
            try:
                enhanced = get_gemini_response(optimize_prompt)
                if enhanced and isinstance(enhanced, str) and len(enhanced.strip()) > 10:
                    # Extract just the query (remove any comments or explanations)
                    enhanced_clean = enhanced.strip().split('\n')[0].strip()
                    # Final validation - reject if it still contains NEAR or ADJ
                    if 'NEAR' not in enhanced_clean.upper() and 'ADJ' not in enhanced_clean.upper():
                        return enhanced_clean
            except Exception:
                pass
            
            return query  # Return original if enhancement fails

        col_run, col_open = st.columns([1, 1])
        with col_run:
            # Use secondary to render in app's brown style
            if st.button("Regenerate Evidence Table", type="secondary", key="p2_search_run"):
                search_term = q_clean or current_q_full or default_q or ""
                if not search_term:
                    st.warning("Please enter a PubMed search query first.")
                else:
                    # Auto-enhance with proximity operators in background
                    with ai_activity("Enhancing query with proximity operators..."):
                        search_term = auto_enhance_query_with_proximity(search_term)
                    search_term = ensure_time_filter(search_term)
                    st.session_state.data['phase2']['mesh_query'] = search_term
                    with ai_activity("Searching PubMed and auto‑grading…"):
                        results = search_pubmed(search_term)
                        st.session_state.data['phase2']['evidence'] = results
                        if results:
                            prompt = (
                                "Assign GRADE quality of evidence (use EXACTLY one of: 'High (A)', 'Moderate (B)', 'Low (C)', or 'Very Low (D)') "
                                "and provide a brief Rationale (1-2 sentences) for each article. "
                                f"{json.dumps([{k:v for k,v in e.items() if k in ['id','title']} for e in results])}. "
                                "Return ONLY valid JSON object where keys are PMID strings and values are objects with 'grade' and 'rationale' fields. "
                                "Example: {\"12345678\": {\"grade\": \"High (A)\", \"rationale\": \"text here\"}}"
                            )
                            # Use native function calling for reliable structured output
                            result = get_gemini_response(
                                prompt, 
                                function_declaration=GRADE_EVIDENCE,
                                thinking_budget=1024
                            )
                            if isinstance(result, dict) and 'arguments' in result:
                                grades = result['arguments'].get('grades', {})
                            elif isinstance(result, dict) and 'grades' not in result:
                                grades = result
                            else:
                                grades = get_gemini_response(prompt, json_mode=True)
                            if grades and isinstance(grades, dict):
                                for e in st.session_state.data['phase2']['evidence']:
                                    pmid_str = str(e['id'])
                                    if pmid_str in grades:
                                        grade_data = grades[pmid_str]
                                        e['grade'] = grade_data.get('grade', 'Un-graded') if isinstance(grade_data, dict) else 'Un-graded'
                                        e['rationale'] = grade_data.get('rationale', 'Not provided.') if isinstance(grade_data, dict) else 'Not provided.'
                        # Ensure defaults if AI response missing
                        for e in st.session_state.data['phase2']['evidence']:
                            e.setdefault('grade', 'Un-graded')
                            e.setdefault('rationale', 'Not yet evaluated.')
                    st.session_state['p2_last_autorun_query'] = search_term
                    st.rerun()
        with col_open:
            # Show PubMed link using the best-available query (user input, current saved, or default)
            effective_q = q_clean or current_q_full or default_q or ""
            if effective_q:
                full_q = ensure_time_filter(effective_q)
                st.link_button(
                    "Open in PubMed ↗",
                    f"https://pubmed.ncbi.nlm.nih.gov/?term={urllib.parse.quote(full_q)}",
                    type="secondary",
                )
            else:
                st.link_button("Open in PubMed ↗", "https://pubmed.ncbi.nlm.nih.gov", disabled=True, type="secondary")

    if st.session_state.data['phase2']['evidence']:
        grade_help = (
            "GRADE Criteria\n"
            "- High (A): Further research is very unlikely to change our confidence in the estimate of effect.\n"
            "- Moderate (B): Further research is likely to have an important impact on our confidence in the estimate of effect and may change the estimate.\n"
            "- Low (C): Further research is very likely to have an important impact on our confidence in the estimate of effect and is likely to change the estimate.\n"
            "- Very Low (D): We are very uncertain about the estimate."
        )

        # Use native Streamlit subheader with help parameter for consistent alignment
        st.subheader("Filter by GRADE", help=grade_help)

        # Default filters set to show all grades initially + 'new only' toggle
        col_g1, col_g2 = st.columns([3, 2])
        with col_g1:
            selected_grades = st.multiselect(
                "",
                ["High (A)", "Moderate (B)", "Low (C)", "Very Low (D)", "Un-graded"],
                default=["High (A)", "Moderate (B)", "Low (C)", "Very Low (D)", "Un-graded"],
                key="grade_filter_multiselect"
            )
        with col_g2:
            st.checkbox("Show only new evidence", value=False, key="p2_show_new_only")
        
        # Sort Logic: High to Low
        grade_order = {"High (A)": 0, "Moderate (B)": 1, "Low (C)": 2, "Very Low (D)": 3, "Un-graded": 4}
        
        evidence_data = st.session_state.data['phase2']['evidence']
        # Normalize flags/grades so new enrichments always surface and filter correctly
        for e in evidence_data:
            is_enriched = e.get('source') == 'enriched_from_phase3'
            e['is_new'] = bool(e.get('is_new')) or is_enriched
            if e.get('grade') not in grade_order:
                e['grade'] = 'Un-graded'
        # Apply Sorting
        evidence_data.sort(key=lambda x: grade_order.get(x.get('grade', 'Un-graded'), 4))
        
        # Filter for Display
        display_data = [e for e in evidence_data if e.get('grade', 'Un-graded') in selected_grades]
        if st.session_state.get('p2_show_new_only'):
            display_data = [e for e in display_data if bool(e.get('is_new')) or e.get('source') == 'enriched_from_phase3']
        df_ev = pd.DataFrame(display_data)
        
        if not df_ev.empty:
            # Ensure all required columns exist in DataFrame
            required_cols = ["id", "title", "grade", "rationale", "url", "authors", "abstract", "year", "journal", "source"]
            for col in required_cols:
                if col not in df_ev.columns:
                    df_ev[col] = ""
            
            # Add visual indicator for newly auto-added entries without emojis or Phase mentions
            df_ev["source"] = df_ev.get("source", "").fillna("")
            # Prefer explicit boolean column for styling & export persistence
            if "is_new" not in df_ev.columns:
                df_ev["is_new"] = df_ev["source"].apply(lambda x: True if x == "enriched_from_phase3" else False)
            else:
                df_ev["is_new"] = df_ev["is_new"].fillna(False)
            # Insert a visible, disabled checkbox for "New" highlighting
            df_ev.insert(0, "new", df_ev["is_new"].astype(bool))

            # Neutral tip without Phase 3 reference
            styled_info("<b>Tip:</b> New evidence is marked with a checked box in the first table column. Review the auto-generated GRADE and rationale. Hover over the top right of the table for more options including CSV download.")

            edited_ev = st.data_editor(
                df_ev[["new", "id", "title", "grade", "rationale", "url"]], 
                column_config={
                    "new": st.column_config.CheckboxColumn("New", disabled=True, help="Auto-added and auto-graded from Phase 3."),
                    "id": st.column_config.TextColumn("PMID", disabled=True, width="small"),
                    "title": st.column_config.TextColumn("Title", width="large"),
                    "grade": st.column_config.SelectboxColumn("GRADE", options=["High (A)", "Moderate (B)", "Low (C)", "Very Low (D)", "Un-graded"], width="small"),
                    "rationale": st.column_config.TextColumn("GRADE Rationale", width="large"),
                    "url": st.column_config.LinkColumn("URL", width="small"),
                }, 
                hide_index=True, width="stretch", key="ev_editor"
            )

            # Persist edits back to session so downloads reflect the latest table
            try:
                edited_records = edited_ev.to_dict("records")
                by_id = {str(r.get("id", "")): r for r in edited_records}
                for e in st.session_state.data['phase2']['evidence']:
                    eid = str(e.get('id', ''))
                    if eid in by_id:
                        row = by_id[eid]
                        e["title"] = row.get("title", e.get("title", ""))
                        e["grade"] = row.get("grade", e.get("grade", "Un-graded"))
                        e["rationale"] = row.get("rationale", e.get("rationale", "Not yet evaluated."))
                        # carry forward markers for styling/exports
                        e["is_new"] = bool(row.get("new", e.get("is_new", False)))
                        if e.get("is_new"):
                            e["source"] = e.get("source", "auto_enriched")
            except Exception:
                pass
            
            # EXPORT OPTIONS SECTION
            st.divider()

            full_df = pd.DataFrame(evidence_data)
            c1, c2 = st.columns([1, 1])

            show_table = True
            show_citations = True

            with c1:
                if show_table:
                    st.subheader("Detailed Evidence Table", help="Includes journal, year, authors, and abstract for all results.")
                    full_export_df = full_df[["id", "title", "grade", "rationale", "url", "journal", "year", "authors", "abstract"]].copy()
                    full_export_df.columns = ["PMID", "Title", "GRADE", "GRADE Rationale", "URL", "Journal", "Year", "Authors", "Abstract"]
                    csv_data_full = full_export_df.to_csv(index=False).encode('utf-8')
                    # Centered download button beneath the section
                    dl_l, dl_c, dl_r = st.columns([1,2,1])
                    with dl_c:
                        st.download_button("Download (.csv)", csv_data_full, file_name="evidence_table.csv", mime="text/csv")

            with c2:
                if show_citations:
                    st.subheader("Formatted Citations", help="Generate Word citations in your preferred style. Pick a preset or type your own below.")
                    style_col, custom_col = columns_top([1, 2])
                    with style_col:
                        citation_style = st.selectbox(
                            "Citation style",
                            ["APA", "MLA", "Vancouver"],
                            key="p2_citation_style"
                        )
                        st.caption("Pick a preset or type your own below.")
                    with custom_col:
                        custom_style = st.text_input(
                            "Or enter custom style name",
                            value="",
                            placeholder="Harvard, Chicago",
                            max_chars=30,
                            key="p2_custom_citation_style"
                        )
                        st.caption("Leave blank to use the preset above.")

                    # Use custom style if provided, otherwise use selected preset
                    final_citation_style = custom_style if custom_style.strip() else citation_style
                    references_source = display_data if display_data else evidence_data
                    citations = references_source or []
                    lines = [format_citation_line(entry, final_citation_style) for entry in citations]
                    docx_bytes = create_references_docx(citations, style=final_citation_style)

                    no_citations = len(citations) == 0
                    dl2_l, dl2_c, dl2_r = st.columns([1,2,1])
                    with dl2_c:
                        if docx_bytes:
                            st.download_button(
                                "Download (.docx)",
                                docx_bytes,
                                file_name="citations.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                disabled=no_citations
                            )
                        else:
                            st.info("Word export unavailable (python-docx not installed)")
                        if no_citations:
                            st.info("Add evidence to enable download.")

    else:
        # If nothing to show, provide a helpful prompt
        styled_info("No results yet. Refine the search or ensure Phase 1 has a condition and setting.")
    render_bottom_navigation()
    st.stop()

# --- PHASE 3 ---
elif "Decision" in phase or "Tree" in phase:
    st.header(f"Phase 3. {PHASES[2]}")
    styled_info("<b>Tip:</b> The AI agent generated an evidence-based decision tree. You can update text, add/remove nodes, or refine using natural language below.")
    
    st.divider()
    
    # Reset enrichment flag each time Phase 3 is loaded (allows re-enrichment if new PMIDs added)
    # This is a safety mechanism to detect changes from previous session
    if 'p3_last_nodes_state' not in st.session_state:
        st.session_state['p3_last_nodes_state'] = None
    
    current_nodes_state = str(st.session_state.data.get('phase3', {}).get('nodes', []))
    if current_nodes_state != st.session_state['p3_last_nodes_state']:
        # Nodes have changed (new entry or edit); reset enrichment to allow re-detection
        st.session_state['p3_enrichment_performed'] = False
        st.session_state['p3_last_nodes_state'] = current_nodes_state
        # Reset auto-enhance tracking when the pathway changes
        st.session_state.data['phase3']['auto_enhance_state'] = {"attempted": False, "succeeded": False}
    
    # Auto-generate table on first entry to Phase 3
    cond = st.session_state.data['phase1']['condition']
    setting = st.session_state.data['phase1']['setting'] or "care setting"
    evidence_list = st.session_state.data['phase2']['evidence']
    
    if not st.session_state.data['phase3']['nodes'] and cond:
        ev_context = "\n".join([f"- PMID {e['id']}: {e['title']} | Abstract: {e.get('abstract', 'N/A')[:200]}" for e in evidence_list[:20]])
        prompt = f"""
        Act as a CLINICAL DECISION SCIENTIST with expertise in Medical Decision Analysis and evidence-based medicine.
        
        TASK: Build a SOPHISTICATED, COMPREHENSIVE decision-science pathway for managing {cond} in {setting}.
        
        FOUNDATIONAL FRAMEWORK (MANDATORY - Preserve All Principles):
        - CGT/Ad/it principles: Explicit decision structure, separate content from form
        - Users' Guide to Medical Decision Analysis (Dobler et al., Mayo Clin Proc 2021):
          * Make decision/chance/terminal flows EXPLICIT through DAG structure
          * Trade off BENEFITS vs HARMS at every decision point with evidence-backed rationales
          * Use evidence-based probabilities and utilities to guide branching
        - Ensure pathway reflects real clinical uncertainty and decision complexity

        SOPHISTICATED PATHWAY ELEMENTS (Adapt These Best-Practice Patterns to Your Specific Clinical Condition):
        
        1. VALIDATED RISK STRATIFICATION:
           - Use validated clinical prediction scores/tools relevant to THIS condition BEFORE diagnostic testing
           - Specify exact numerical thresholds for risk categories
           - Examples by condition type (adapt to your clinical scenario):
             * DVT/PE: Wells' Criteria (≤1, 2-6, ≥7), PERC rule, YEARS algorithm, PESI score
             * ACS/MI: HEART score, TIMI risk score, GRACE score
             * Stroke: NIHSS, ABCD2 score for TIA
             * Sepsis: qSOFA, SIRS criteria, Sepsis-3 definitions
             * Trauma: GCS, Trauma Score, Injury Severity Score
           - Apply age-adjusted or population-specific cutoffs where established (e.g., "D-dimer = age × 10 if >50" for VTE)
           - NOT generic "assess risk"—use published, validated tools with scoring for THIS condition
        
        2. SPECIAL POPULATION HANDLING:
           - Screen for special populations EARLY in pathway (before exposing to risks)
           - Pregnancy considerations: Check status before radiation imaging (X-ray, CT), teratogenic drugs, or procedures
           - Renal function: Assess before contrast, NSAIDs, or renally-cleared medications
           - Check for absolute contraindications before initiating high-risk treatments specific to THIS condition
           - Age-based considerations: Pediatric vs. geriatric dosing adjustments, atypical presentations
           - Comorbidity modifications: Active bleeding, immunosuppression, organ failure, drug allergies
        
        3. RESOURCE AVAILABILITY CONTINGENCIES:
           - Include explicit "What if preferred test/procedure is unavailable?" branches relevant to THIS pathway
           - Diagnostic alternatives: "If [preferred imaging] NOT available → [Alternative test] OR Transfer OR Empiric treatment"
           - Bed availability: "ED Observation bed available?", "ICU bed available?", "Telemetry bed available?"
           - Specialist availability: "Immediate consult available?" vs. "Scheduled follow-up" vs. "Transfer"
           - Equipment/supply constraints: Alternative diagnostic or therapeutic approaches
        
        4. MEDICATION SPECIFICITY (Critical - Adapt to THIS condition's treatments):
           - ALWAYS include brand AND generic names for common medications: "Drug (Brand)" or "Brand (Drug)"
           - Exact dosing with route, frequency, duration: "X mg [route] [frequency] × Y days/weeks"
           - Administration timing and location: "Give first dose in ED", "Start within X hours of symptom onset"
           - Population-specific preferences: "Preferred for CKD", "Avoid if [contraindication]", "Adjust for [condition]"
           - Practical prescribing details: "Prescribe starter pack", "Provide patient education sheet", "Use weight-based dosing"
           - Insurance and cost considerations: "Ensure Rx covered by insurance; provide coupon/assistance program link if needed"
           - Examples from anticoagulation (adapt format to THIS condition's drugs):
             * "Apixaban (Eliquis): 10 mg PO twice daily × 7 days, then 5 mg PO twice daily (74 tablets)"
             * "Enoxaparin (Lovenox): 1 mg/kg SQ q12h. Adjust for CrCl <30."
           - Never generic "start medication"—always specific drugs/classes with doses
        
        5. FOLLOW-UP PATHWAYS:
           - Specific timing appropriate to THIS condition: "[Provider type] within [timeframe]"
           - Specific provider types relevant to THIS condition: Cardiologist, Pulmonologist, Surgeon, etc.
           - Virtual care alternatives: "If unable to follow-up with [provider], advise Virtual Urgent Care" or telehealth options
           - Contingency plans: "If no [provider], provide [alternative resources]"
           - What to monitor specific to THIS condition: "Check [lab] in X days", "Repeat [test] in Y weeks"
        
        6. EDUCATIONAL CONTENT INTEGRATION:
           - Note opportunities for hyperlinks to validated clinical tools relevant to THIS condition
           - Link to medication information specific to drugs used in THIS pathway
           - Evidence citations: Relevant clinical guidelines, landmark trials, consensus statements
           - Patient education resources specific to THIS condition
           - Return precautions and red flag symptoms appropriate for THIS condition's discharge instructions
        
        Available Evidence Base:
        {ev_context}
        
        REQUIRED CLINICAL COVERAGE (4 Mandatory Stages - Each MUST Have Complexity):
        
        1. Initial Evaluation:
           - Chief complaint and symptom characterization
           - Vital signs assessment (with abnormality thresholds)
           - Physical examination findings and validated risk stratification
           - Validated clinical prediction scores relevant to THIS condition (with specific numerical thresholds)
           - Age-adjusted or population-specific thresholds where established in literature
           - Special population screening EARLY (pregnancy before radiation/teratogens, renal function before contrast/medications)
           - Early diagnostic workup (labs, imaging, monitoring)
        
        2. Diagnosis and Treatment:
           - Differential diagnosis decision trees (what tests rule in/out?)
           - Resource availability contingencies: "If preferred test unavailable → Alternative pathway or transfer"
           - Contraindication checks BEFORE treatment initiation specific to THIS condition's therapies
           - Therapeutic interventions with EXACT specificity (adapt examples below to THIS condition):
             * Brand AND generic names where commonly used
             * Exact dosing: "X mg [route] [frequency] × duration (total quantity)"
             * Administration details: "Give first dose in [location]", "Preferred for [population]", "Adjust for [condition]"
             * Insurance/cost considerations: "Ensure Rx covered by insurance" + assistance programs where applicable
             * IMPORTANT: Medication administration is a CLINICAL ACTION—create a Process node for each medication with the dosing IN the label
             * Example node: type="Process", label="Administer aspirin 325mg PO, clopidogrel 600mg IV bolus, heparin 70 U/kg", detail="Benefit: Prevents stent thrombosis..."
           - Risk-benefit analysis for major therapeutic choices specific to THIS condition
           - Edge cases and special populations (pregnant, elderly, immunocompromised, etc.) relevant to THIS condition
        
        3. Re-evaluation:
           - Monitoring criteria and frequency (vital signs, labs, imaging follow-ups)
           - Response to treatment assessment (improving vs. unchanged vs. deteriorating)
           - Escalation triggers and de-escalation pathways
           - When to repeat diagnostic testing or change therapy
           - Bed availability considerations: "ED Observation if available, else Medicine/SDU/MICU admit or Dispo navigator"
        
        4. Final Disposition:
           - Specific discharge instructions (medications with dose/route/duration, activity restrictions, dietary changes)
           - EXPLICIT follow-up pathways with timing and provider type:
             * "PCP follow-up within 2 weeks"
             * "OBGYN follow-up" (for pregnant patients)
             * "Vascular Surgery Referral" (specialty consult)
             * Virtual care alternatives: "If unable to follow-up with PCP, advise Virtual Urgent Care"
           - Educational content integration: Score calculators, evidence citations, patient resources
           - Admit/observation criteria with clear thresholds
           - Transfer to higher level of care (ICU, specialty unit) triggers
           - Return precautions and red flag symptoms for discharged patients
        
        OUTPUT FORMAT: JSON array of nodes with THESE EXACT FIELDS:
        - "type": "Start" | "Decision" | "Process" | "End" (no other types)
        - "label": Concise, specific clinical step using medical abbreviations (e.g., "ECG, troponin x2 at 0h/3h, IV access")
          * CRITICAL: If the clinical action includes medication administration, include it in the label itself
          * Example GOOD: "Administer aspirin 325mg PO + clopidogrel 600mg IV loading dose"
          * Example GOOD: "Start vancomycin 15-20 mg/kg IV q8-12h, adjust for renal function"
          * Example BAD: "Medication administration" (notes: "aspirin 325mg...") ← medication should be IN the label!
          * Clinical medications are ACTIONS (belong in label), not background notes
        - "evidence": PMID citation OR "N/A"
        - "notes": (optional) Actionable clinical details for pathway users:
          * RED FLAG SIGNS: Specific warning signs that require immediate action (e.g., "Red flags: syncope with exertion, family hx sudden death, abnormal ECG")
          * CLINICAL THRESHOLDS: Specific values triggering action (e.g., "Escalate if: HR>120, SBP<90, SpO2<92%")
          * MONITORING PARAMETERS: What to watch and when (e.g., "Monitor: troponin q3h, telemetry x24h")
          * SPECIAL CONSIDERATIONS: Population-specific notes (e.g., "Pregnancy: avoid CT, use MRI/US")
          * Do NOT include: Generic benefit/harm discussions, rationale explanations
        
        CRITICAL CONSTRAINTS (PRESERVE DECISION SCIENCE INTEGRITY):
        
        1. DECISION DIVERGENCE - Every Decision creates DISTINCT branches with MINIMUM SEPARATION:
           - "Is patient hemodynamically stable?" YES→Observation pathway | NO→ICU-level resuscitation
           - "Does EKG show STEMI?" YES→Cath lab pathway | NO→Serial troponin pathway
           
           MINIMUM DIVERGENCE RULE (CRITICAL):
           - Each branch from a Decision MUST have at least 2-3 unique steps BEFORE any convergence
           - NEVER have both branches immediately point to the same next node
           - If branches need to eventually converge (shared End node or shared later Process), they must first diverge meaningfully
           
           ALLOWED CONVERGENCE (later in pathway):
           ✓ Multiple pathways ending at shared End nodes: "Discharge with cardiology follow-up in 2 weeks"
           ✓ Branches meeting at a shared later Process step after meaningful divergence (3+ steps apart)
           ✓ Parallel workups that later merge for disposition decision
           
           FORBIDDEN CONVERGENCE (premature/immediate):
           ✗ Decision branches pointing to same immediate next node (renders decision meaningless)
           ✗ Branches merging within 1-2 steps of the decision (insufficient divergence)
           ✗ "Diamond" patterns where YES/NO both go to same Process immediately
           
           Example of WRONG premature convergence:
             Decision: "Fever present?" YES→(Cultures) | NO→(Cultures) → Same immediate step = BAD
           
           Example of CORRECT eventual convergence:
             Decision: "Fever >38.5°C?" 
               YES→(Blood cultures)→(Broad-spectrum antibiotics)→(ICU evaluation)→END: ICU admit
               NO→(Observation)→(Supportive care)→(Monitor 6h)→END: Discharge if stable
             Both pathways may share "Discharge planning" node AFTER their unique 3+ step sequences
        
        2. TERMINAL END NODES - Each pathway branch ends ONLY with End nodes:
           - No content after an End node
           - End nodes represent final disposition: "Discharged on aspirin/metoprolol x90 days with PCP follow-up"
           - Each clinical outcome gets its own End node; DO NOT use "or" (e.g., BAD: "Admit or ICU")
           - Even similar outcomes get separate End nodes if they represent distinct pathways
        
        3. EVIDENCE-BACKED STEPS:
           - Every Process and Decision node should have a PMID when available (from evidence list above)
           - If multiple PMIDs support a step, use one representative citation
           - Do NOT hallucinate PMIDs—use "N/A" if no supporting evidence in list
        
        4. COMPLEXITY AND SPECIFICITY:
           - Build comprehensive pathway (typically 15-40 nodes depending on clinical complexity)
           - Include ALL relevant special populations and edge cases:
             * Pregnancy status (check EARLY before radiation/teratogenic drugs)
             * Renal failure (CrCl-based dosing adjustments, contrast contraindications)
             * Drug allergies and absolute contraindications
             * Age extremes (pediatric vs. geriatric dosing/monitoring)
             * Active bleeding or high bleeding risk
             * Comorbidities affecting management (cancer, prior events, thrombophilia)
           - Medication specificity (CRITICAL - Never be vague):
             ✓ "Apixaban (Eliquis): 10 mg PO BID × 7d, then 5 mg PO BID. Give first dose in ED. Prescribe 74-tablet starter pack. Ensure Rx covered by insurance; provide Apixaban coupon link if needed."
             ✓ "Rivaroxaban (Xarelto): 15 mg PO BID × 21d, then 20 mg PO daily. Preferred for patients with CKD or ESRD."
             ✓ "Enoxaparin (Lovenox): 1 mg/kg SQ q12h OR 1.5 mg/kg SQ daily. Adjust for CrCl <30."
             ✗ "Start anticoagulation" (TOO VAGUE)
             ✗ "Treat with antibiotics" (TOO VAGUE)
           - Diagnostic test specificity with alternatives:
             ✓ "Order compression ultrasound of affected leg. If ultrasound NOT available → Hold anticoagulation, transfer to facility with imaging, OR give one-time therapeutic dose and arrange urgent outpatient imaging."
             ✗ "Order imaging" (TOO VAGUE)
           - Monitoring intervals with explicit timing:
             ✓ "Recheck troponin q3h × 2, then daily troponin × 2 if negative"
             ✓ "Vitals q15min × 1h, then q1h × 4h, then q4h if stable"
           - Clinical score thresholds with numerical cutoffs:
             ✓ "Wells' Score: ≤1 (low), 2-6 (intermediate), ≥7 (high)"
             ✓ "PESI Score: <86 (very low), 86-105 (low), 106-125 (intermediate), >125 (high)"
        
        5. DAG STRUCTURE (No cycles):
           - Pathway is a directed acyclic graph (DAG)—never loop back
           - Escalation only moves forward (ICU-bound patients don't move back to ED)
           - De-escalation is explicit: "Stable x 24h→Transfer to med/surg bed from ICU"
        
        6. ACTIONABILITY AND CLINICAL REALISM:
           - Every node represents an action or decision a clinician takes in real time
           - Include realistic clinical decision points: "Vitals stable x 2h" or "Troponin rising vs. falling?"
           - Timestamps and criteria matter: "Admit if BP <90 persistently AFTER 2L fluid bolus"
           - Resource availability branches: "ED Observation bed available?", "Ultrasound available now?"
        
        7. VISUAL DESIGN CUES (For Phase 4 Optimization):
           - Indicate risk levels for color coding: [Low Risk], [Intermediate Risk], [High Risk], [Alert/Critical]
           - Mark informational boxes: [Info], [Contraindication], [Special Population]
           - Suggest hyperlink candidates: validated scores, drug information, evidence citations
           - Note resource dependencies: [Requires Ultrasound], [Requires CT], [Requires Specialty Consult]
        
        Rules for Node Structure:
        - First node: type "Start", label "Patient present to {setting} with {cond}"
        - Last nodes: All type "End" (no Process/Decision after End)
        - Consecutive Decision nodes are OK (do NOT force Process nodes between them)
        - Use compound labels for clarity: "Assess troponin, CXR, EKG—any abnormality?" (Decision)
        - Notes field for actionable details (e.g., notes: "Red flags: syncope with exertion, chest pain, palpitations")
        
        LABEL CLARITY REQUIREMENTS (CRITICAL - Read Every Label Carefully):
        - Labels must be READABLE: max 120 characters per label
        - DO NOT use \\n or newline characters in labels - use plain text only
        - Use STANDARD MEDICAL ABBREVIATIONS only (not made-up symbols or extraneous characters)
        - Clean encoding: NO special Unicode characters, escaped sequences, or corrupted text
        - Prioritize clarity: Spell out potentially ambiguous terms (e.g., "Myocardial Infarction" not cryptic shorthand)
        - Each label should answer: "What does the clinician DO or DECIDE here?"
        - Example GOOD labels:
          ✓ "Elevated troponin AND chest pain symptoms: Admit to cardiac ICU"
          ✓ "Wells' Score 2-6 (intermediate DVT risk): Order compression ultrasound"
          ✓ "Age >65 AND renal failure (CrCl <30): Use reduced-dose enoxaparin"
        - Example BAD labels:
          ✗ "RuleOut!MI|EKG◊STEMIδ†neuro↔shock—→cath" (extraneous characters)
          ✗ "Ì÷ôè¢¨§ßþ" (corrupted encoding)
          ✗ "[Patient_with_multiple_comorbidities_age_>80_presenting_with_chest_pain_dyspnea_and_recent_fall]" (too long)
        
        Node Count Guidance:
        - MINIMUM 15 nodes (simple pathway structure)
        - TYPICAL 25-35 nodes (comprehensive with main branches)
        - MAXIMUM 50+ nodes (complex with edge cases, special populations, escalation/de-escalation)
        - Aim for depth over breadth: prefer explicit decision trees over oversimplification
        
        Generate a pathway that respects real clinical complexity and decision uncertainty. This is NOT a linear checklist—it's a decision tree that branches and evolves based on patient presentation and test results.
        
        CRITICAL ANTI-CONVERGENCE RULES:
        - Each Decision branch must have AT LEAST 2-3 unique nodes before any potential convergence
        - Branches may eventually share End nodes OR late-stage Process nodes, but ONLY after meaningful divergence
        - NEVER create "diamond" patterns where both branches immediately go to the same node
        - If you find yourself pointing two branches to the same next step, STOP and create distinct pathways first
        - Test: For every Decision, trace each branch forward 3 steps - they should be DIFFERENT steps
        
        DECISION NODE JSON STRUCTURE (CRITICAL - Follow This Exactly):
        
        Decision nodes MUST include a "branches" array with explicit "target" indices pointing to DIFFERENT nodes:
        
        EXAMPLE CORRECT STRUCTURE (10 nodes showing proper branching):
        ```json
        [
          {{"type": "Start", "label": "Patient presents to ED with chest pain", "evidence": "N/A"}},
          {{"type": "Process", "label": "Obtain ECG, troponin, vitals", "evidence": "N/A"}},
          {{"type": "Decision", "label": "STEMI on ECG?", "evidence": "N/A", "branches": [
            {{"label": "Yes", "target": 3}},
            {{"label": "No", "target": 6}}
          ]}},
          {{"type": "Process", "label": "Activate cath lab, give aspirin 325mg, heparin bolus", "evidence": "12345678"}},
          {{"type": "Process", "label": "Transfer to cath lab for PCI", "evidence": "N/A"}},
          {{"type": "End", "label": "Admit to CCU post-PCI with dual antiplatelet therapy", "evidence": "N/A"}},
          {{"type": "Process", "label": "Serial troponins q3h x2, telemetry monitoring", "evidence": "N/A"}},
          {{"type": "Decision", "label": "Troponin elevated or rising?", "evidence": "N/A", "branches": [
            {{"label": "Yes", "target": 8}},
            {{"label": "No", "target": 9}}
          ]}},
          {{"type": "End", "label": "Admit for NSTEMI workup, cardiology consult", "evidence": "N/A"}},
          {{"type": "End", "label": "Discharge with PCP follow-up in 72h, return precautions", "evidence": "N/A"}}
        ]
        ```
        
        KEY POINTS FROM THIS EXAMPLE:
        - Node 2 (Decision) branches to DIFFERENT nodes: target 3 (cath lab pathway) vs target 6 (serial troponin pathway)
        - Node 7 (Decision) branches to DIFFERENT End nodes: target 8 (admit) vs target 9 (discharge)
        - Each branch leads to its own distinct pathway
        - Start node is index 0, all paths eventually reach End nodes
        - The "target" values are 0-based indices into the node array
        
        COMMON MISTAKE TO AVOID:
        ```json
        {{"type": "Decision", "label": "Risk level?", "branches": [
          {{"label": "High", "target": 5}},
          {{"label": "Low", "target": 5}}  // WRONG! Both point to same node
        ]}}
        ```
        This renders the decision meaningless. Each branch MUST point to a different target.
        """
        with ai_activity("Generating..."):
            # Use native function calling for reliable structured output
            result = get_gemini_response(
                prompt, 
                function_declaration=GENERATE_PATHWAY_NODES,
                thinking_budget=2048  # Complex pathway generation
            )
            # Extract nodes from function call or fall back to raw result
            if isinstance(result, dict) and 'arguments' in result:
                nodes = result['arguments'].get('nodes', [])
            elif isinstance(result, dict) and 'nodes' in result:
                nodes = result.get('nodes', [])
            elif isinstance(result, list):
                nodes = result
            else:
                # Fallback to json_mode
                nodes = get_gemini_response(prompt, json_mode=True)
        if isinstance(nodes, list) and len(nodes) > 0:
            # Clean up common AI generation issues
            nodes = normalize_or_logic(nodes)  # Fix OR statements in End nodes
            nodes = fix_decision_flow_issues(nodes)  # Fix reconverging branches and non-terminal End nodes
            st.session_state.data['phase3']['nodes'] = nodes
            st.rerun()
        elif not isinstance(nodes, list):
            st.warning(f"❌ Could not parse decision tree. The AI returned: {type(nodes).__name__}. Please manually add nodes in the table below, or try again.")
    
    st.divider()
    st.markdown("### Decision Tree")
    
    # Build evidence options from Phase 2
    evidence_ids = ["N/A"] + [e['id'] for e in evidence_list]
    
    # Initialize with empty row if no nodes
    if not st.session_state.data['phase3']['nodes']:
        st.session_state.data['phase3']['nodes'] = [{"type": "Start", "label": "", "evidence": "N/A", "notes": ""}]
    
    df_nodes = pd.DataFrame(st.session_state.data['phase3']['nodes'])
    # Ensure notes column exists (may be 'detail' in older data)
    if 'notes' not in df_nodes.columns:
        if 'detail' in df_nodes.columns:
            df_nodes['notes'] = df_nodes['detail']
        else:
            df_nodes['notes'] = ""
    edited_nodes = st.data_editor(
        df_nodes,
        column_config={
            "type": st.column_config.SelectboxColumn(
                "Type",
                options=["Start", "Decision", "Process", "End"],
                required=True,
                width="small"
            ),
            "label": st.column_config.TextColumn(
                "Clinical Step",
                width="large",
                required=True
            ),
            "evidence": st.column_config.TextColumn(
                "Supporting Evidence (PMID)",
                width="medium",
                help="Enter PMID or 'N/A'"
            ),
            "notes": st.column_config.TextColumn(
                "Notes",
                width="large",
                help="Actionable clinical details: red flags, thresholds, monitoring parameters (displayed as blue trapezoid in visualization)"
            )
        },
        num_rows="dynamic",
        hide_index=True,
        width="stretch",
        key="p3_editor"
    )
    # Auto-save on edit
    st.session_state.data['phase3']['nodes'] = edited_nodes.to_dict('records')
    
    st.divider()
    
    # Display pathway metrics with evidence enrichment
    node_count = len(st.session_state.data['phase3']['nodes'])

    # Extract all PMIDs from Phase 3 nodes
    phase3_pmids = extract_pmids_from_nodes(st.session_state.data['phase3']['nodes'])
    
    # De-duplicate Phase 2 evidence first (in case there are existing duplicates)
    seen_pmids = set()
    deduplicated_evidence = []
    for e in evidence_list:
        pmid = e.get('id')
        if pmid not in seen_pmids:
            seen_pmids.add(pmid)
            deduplicated_evidence.append(e)
    
    # Update Phase 2 evidence list with deduplicated version
    st.session_state.data['phase2']['evidence'] = deduplicated_evidence
    evidence_list = deduplicated_evidence
    
    phase2_pmids = seen_pmids
    
    # Identify new PMIDs in Phase 3 not yet in Phase 2
    new_pmids_in_phase3 = phase3_pmids - phase2_pmids
    
    # Count evidence-backed nodes (nodes with non-'N/A' evidence field)
    evidence_backed_nodes = [n for n in st.session_state.data['phase3']['nodes'] 
                            if n.get('evidence', 'N/A') not in ['N/A', '', None]]
    evidence_backed_count = len(evidence_backed_nodes)
    
    # Track enrichment state: re-enrich only when new PMIDs appear that haven't been processed
    if 'p3_last_enriched_pmids' not in st.session_state:
        st.session_state['p3_last_enriched_pmids'] = set()
    
    enriched_count = 0
    pmids_to_enrich_now = new_pmids_in_phase3 - st.session_state['p3_last_enriched_pmids']
    
    # Auto-enrich Phase 2 with truly NEW PMIDs (not previously enriched)
    if pmids_to_enrich_now:
        with ai_activity("Enriching evidence and auto‑grading new PMIDs…"):
            new_evidence_list = enrich_phase2_with_new_pmids(pmids_to_enrich_now, phase2_pmids)
            if new_evidence_list:
                # Auto‑grade the newly added items
                auto_grade_evidence_list(new_evidence_list)
                # Mark as new for visual highlight and add to Phase 2
                for e in new_evidence_list:
                    e["is_new"] = True
                    if not e.get("source"):
                        e["source"] = "enriched_from_phase3"
                st.session_state.data['phase2']['evidence'].extend(new_evidence_list)
                enriched_count = len(new_evidence_list)
        # Update tracking to prevent re-enrichment of same PMIDs
        st.session_state['p3_last_enriched_pmids'] = new_pmids_in_phase3.copy()
    
    # Create metrics display
    if evidence_backed_count == 0:
        evidence_status = "No evidence citations"
    else:
        evidence_status = f"{evidence_backed_count} evidence-based nodes"
        if len(new_pmids_in_phase3) > 0:
            evidence_status += f" | {len(new_pmids_in_phase3)} new evidence in Phase 2"
    
    st.markdown(f"**Pathway Metrics:** {node_count} total nodes | {evidence_status}")
    
    # Show tip if new PMIDs detected (whether just enriched or already enriched)
    if new_pmids_in_phase3:
        styled_info(
            f"<b>New evidence identified:</b> {len(new_pmids_in_phase3)} PMID(s) from Phase 3 auto-graded and added to Phase 2 (marked with checked box in first column). "
            f"<b>Visit Phase 2</b> to review and finalize."
        )

    st.divider()

    # PDF Pathway Upload (new feature)
    with st.expander("🔍 Upload Clinical Guideline or Flowchart PDF", expanded=False):
        st.markdown("**Upload a PDF** containing clinical guidelines or flowchart diagrams. The AI agent will extract the pathway structure and replace the current pathway.")
        
        pdf_upload = st.file_uploader(
            "Upload PDF (guideline text, flowchart, or both)",
            type=["pdf"],
            key="p3_pathway_pdf",
            help="PDF will be analyzed for clinical decision logic and converted to pathway nodes"
        )
        
        if pdf_upload:
            # Check file size
            file_bytes = pdf_upload.read()
            file_size_mb = len(file_bytes) / (1024 * 1024)
            est_pages = int(file_size_mb * 20)
            
            col_info, col_action = st.columns([2, 1])
            with col_info:
                st.info(f"📄 **File:** {pdf_upload.name} ({file_size_mb:.1f} MB, ~{est_pages} pages)")
                
                if est_pages > 50:
                    st.warning("⚠️ Large PDF (>50 pages). Extraction may take 60-90 seconds or fail.")
            
            with col_action:
                extract_btn = st.button("Extract Pathway", type="primary", key="p3_extract_pdf")
            
            if extract_btn:
                # Generate file hash for caching
                import hashlib
                from io import BytesIO
                file_hash = hashlib.md5(file_bytes).hexdigest()
                st.session_state["pdf_file_hash"] = file_hash
                
                # Upload to Gemini
                with st.spinner("Uploading PDF to Gemini..."):
                    try:
                        from phase5_helpers import extract_pathway_from_pdf
                        file_obj = BytesIO(file_bytes)
                        uploaded = get_genai_client().files.upload(
                            file=(pdf_upload.name, file_obj, 'application/pdf')
                        )
                        
                        # Extract with progress
                        progress_bar = st.progress(0.0)
                        progress_text = st.empty()
                        
                        def update_progress(value, text):
                            progress_bar.progress(value)
                            progress_text.text(text)
                        
                        result = extract_pathway_from_pdf(
                            file_uri=uploaded.uri,
                            condition=cond,
                            setting=setting,
                            file_hash=file_hash,
                            progress_callback=update_progress,
                            genai_client=get_genai_client()
                        )
                        
                        progress_bar.empty()
                        progress_text.empty()
                        
                        st.session_state["pdf_extraction"] = result
                        st.rerun()
                    except Exception as e:
                        st.error(f"PDF extraction failed: {str(e)}")
        
        # Show results if extraction complete
        if "pdf_extraction" in st.session_state:
            result = st.session_state["pdf_extraction"]
            confidence = result["confidence"]
            
            col1, col2, col3, col4 = st.columns(4)
            with col1:
                st.metric("Nodes", len(result["nodes"]))
            with col2:
                st.metric("Confidence", f"{confidence*100:.0f}%")
            with col3:
                st.metric("PMIDs", result["pmids_extracted"])
            with col4:
                doc_icon = {"guideline": "📄", "flowchart": "📊", "hybrid": "📄+📊"}
                st.metric("Type", doc_icon.get(result["doc_type"]["type"], "❓"))
            
            # Show auto-fixes if any
            if result.get("fixes_applied"):
                with st.expander("⚙️ Auto-Fixes Applied", expanded=False):
                    for fix in result["fixes_applied"]:
                        st.write(f"- {fix}")
            
            # High confidence
            if confidence >= 0.8:
                st.success("✅ High-quality extraction")
                if st.button("✓ Use Extracted Pathway", type="primary", key="p3_use_extraction"):
                    st.session_state.data['phase3']['nodes'] = result["nodes"]
                    st.session_state.data['phase4']['viz_cache'] = {}
                    st.success("✓ Pathway replaced! Scroll down to review.")
                    del st.session_state["pdf_extraction"]
                    st.rerun()
            
            # Low confidence
            else:
                st.warning("⚠️ Moderate confidence — review or refine")
                
                with st.form("refine_extraction"):
                    st.write("**Preview (first 3 nodes):**")
                    st.json(result["nodes"][:3])
                    
                    refined_prompt = st.text_area(
                        "Refine Extraction Instructions",
                        placeholder="E.g., 'Focus on pages 8-15', 'Extract exact troponin thresholds', 'Use HEART score criteria'",
                        height=80,
                        key="p3_refined_prompt"
                    )
                    
                    col_retry, col_accept, col_cancel = st.columns(3)
                    with col_retry:
                        retry = st.form_submit_button("🔄 Retry", type="primary")
                    with col_accept:
                        accept = st.form_submit_button("✓ Accept")
                    with col_cancel:
                        cancel = st.form_submit_button("✕ Cancel")
                
                if retry and refined_prompt:
                    try:
                        from phase5_helpers import extract_pathway_from_pdf
                        progress_bar = st.progress(0.0)
                        progress_text = st.empty()
                        
                        def update_progress(value, text):
                            progress_bar.progress(value)
                            progress_text.text(text)
                        
                        result = extract_pathway_from_pdf(
                            file_uri=result["file_uri"],
                            condition=cond,
                            setting=setting,
                            custom_prompt=refined_prompt,
                            file_hash=st.session_state.get("pdf_file_hash"),
                            progress_callback=update_progress,
                            genai_client=get_genai_client()
                        )
                        
                        progress_bar.empty()
                        progress_text.empty()
                        st.session_state["pdf_extraction"] = result
                        st.rerun()
                    except Exception as e:
                        st.error(f"Retry failed: {str(e)}")
                
                if accept:
                    st.session_state.data['phase3']['nodes'] = result["nodes"]
                    st.session_state.data['phase4']['viz_cache'] = {}
                    st.info("✓ Pathway accepted. Review in data editor below.")
                    del st.session_state["pdf_extraction"]
                    st.rerun()
                
                if cancel:
                    del st.session_state["pdf_extraction"]
                    if "pdf_file_hash" in st.session_state:
                        del st.session_state["pdf_file_hash"]
                    st.rerun()

    # Refine & Regenerate (placed below the table)
    st.divider()
    submitted = False
    with st.expander("Refine & Regenerate", expanded=False):
        st.markdown("**Tip:** Describe any desired modifications and optionally attach supporting documents. Click \"Regenerate\" to automatically update all Phase 3 content and downloads")
        with st.form("p3_refine_form"):
            col_text, col_file = columns_top([2, 1])
            with col_text:
                st.text_area(
                    "Refinement Notes",
                    key="p3_refine_input",
                    placeholder="Add branch for renal impairment; include discharge meds for heart failure; clarify follow‑up",
                    height=90,
                    help="Describe what to change. After applying, the pathway will regenerate above."
                )

            with col_file:
                st.caption("Supporting Documents (optional)")
                p3_uploaded = st.file_uploader(
                    "Drag & drop or browse",
                    key="p3_file_upload",
                    accept_multiple_files=True,
                    label_visibility="collapsed",
                    help="Attach PDFs/DOCX files; the agent auto-summarizes them for context."
                )
                if p3_uploaded:
                    for uploaded_file in p3_uploaded:
                        file_result = upload_and_review_file(uploaded_file, f"p3_refine_{uploaded_file.name}", "decision tree pathway")
                        if file_result:
                            with st.expander(f"Review: {file_result['filename']}", expanded=False):
                                st.markdown(file_result["review"])

            col1, col2 = st.columns([3, 1])
            with col2:
                submitted = st.form_submit_button("Regenerate", type="secondary", use_container_width=True)

    if submitted:
        refinement_request = st.session_state.get('p3_refine_input', '').strip()
        if refinement_request:
            # Collect all uploaded document reviews
            doc_reviews = []
            for key in st.session_state.keys():
                if key.startswith("file_p3_refine_") and key != "file_p3_refine":
                    doc_reviews.append(st.session_state.get(key, ''))
            
            if doc_reviews:
                refinement_request += f"\n\nSupporting Documents:\n" + "\n\n".join(doc_reviews)
            
            if st.session_state.data['phase3']['nodes']:
                with ai_activity("Applying refinements to decision tree..."):
                    current_nodes = st.session_state.data['phase3']['nodes']
                    ev_context = "\n".join([f"- PMID {e['id']}: {e['title']} | Abstract: {e.get('abstract', 'N/A')[:200]}" for e in evidence_list[:20]])
                    prompt = f"""
                    Act as a Clinical Decision Scientist. Refine the existing pathway based on the user's request.

                    Current pathway for {cond} in {setting}:
                    {json.dumps(current_nodes, indent=2)}

                    Available Evidence:
                    {ev_context}

                    User's refinement request: "{refinement_request}"

                    Apply the requested changes while maintaining:
                    - CGT/Ad/it principles and Medical Decision Analysis best practices
                    - Coverage of: Initial Evaluation, Diagnosis/Treatment, Re-evaluation, Final Disposition
                    - Actionable steps with medical acronyms for brevity
                    - Specific discharge details (prescriptions with dose/route, referrals)
                    - Evidence citations (PMIDs where applicable)

                    Output: Complete revised JSON array of nodes with fields: type, label, evidence.
                    Rules:
                    - type in [Start, Decision, Process, End]
                    - First node: type "Start", label "patient present to {setting} with {cond}"
                    - NO node count limit - build complete clinical flow
                    - If >20 nodes, organize into sections or sub-pathways
                    """
                    # Use native function calling for reliable structured output
                    result = get_gemini_response(
                        prompt, 
                        function_declaration=GENERATE_PATHWAY_NODES,
                        thinking_budget=2048  # Complex refinement needs reasoning
                    )
                    # Extract nodes from function call or fall back
                    if isinstance(result, dict) and 'arguments' in result:
                        nodes = result['arguments'].get('nodes', [])
                    elif isinstance(result, dict) and 'nodes' in result:
                        nodes = result.get('nodes', [])
                    elif isinstance(result, list):
                        nodes = result
                    else:
                        # Fallback to json_mode
                        nodes = get_gemini_response(prompt, json_mode=True)
                    if isinstance(nodes, list) and len(nodes) > 0:
                        # Clean up common AI generation issues
                        nodes = normalize_or_logic(nodes)
                        nodes = fix_decision_flow_issues(nodes)
                        st.session_state.data['phase3']['nodes'] = nodes
                        # Clear Phase 4 visualization cache so regenerated views/downloads reflect updates
                        st.session_state.data.setdefault('phase4', {}).pop('viz_cache', None)
                        st.success("Refinements applied and pathway regenerated!")
                    else:
                        st.error("Failed to regenerate pathway. Please try again.")
                    st.rerun()
    
    # Navigation at the bottom
    render_bottom_navigation()
    st.stop()

# --- PHASE 4 ---
elif "Interface" in phase or "UI" in phase:
    st.header(f"Phase 4. {PHASES[3]}")
    styled_info("<b>Tip:</b> AI agent evaluates all Nielsen heuristics and applies those that meaningfully improve your pathway. Click 'Apply All Heuristics' to batch-apply intelligent recommendations. Review results below.")
    
    nodes = st.session_state.data['phase3']['nodes']
    p4_state = st.session_state.data.setdefault('phase4', {})

    # GUARD: If no nodes, continue rendering Phase 4 with guidance
    if not nodes:
        st.warning("No pathway nodes found. You can still review heuristics and apply custom refinements below.")

    # Initialize Phase 4 state defaults
    p4_state.setdefault('nodes_history', [])
    p4_state.setdefault('heuristics_data', {})
    p4_state.setdefault('auto_heuristics_done', False)
    p4_state.setdefault('ui_apply_flags', {})
    p4_state.setdefault('applied_status', False)
    p4_state.setdefault('applied_summary', "")
    p4_state.setdefault('applying_heuristics', False)  # Flag to prevent re-analysis during apply

    # Detect pathway changes and allow heuristics to re-run
    # BUT: Don't clear heuristics if we're currently applying them
    try:
        nodes_hash = hashlib.md5(json.dumps(nodes, sort_keys=True).encode('utf-8')).hexdigest()
    except Exception:
        nodes_hash = str(len(nodes))
    
    if p4_state.get('last_nodes_hash') != nodes_hash and not p4_state.get('applying_heuristics'):
        p4_state['last_nodes_hash'] = nodes_hash
        # Clear prior heuristics to refresh recommendations only if nodes changed externally (not from applying heuristics)
        p4_state['heuristics_data'] = {}
        p4_state['auto_heuristics_done'] = False
        p4_state['applied_status'] = False
        p4_state['applied_summary'] = ""
    elif p4_state.get('applying_heuristics'):
        # We just applied heuristics, update hash but keep heuristics data
        p4_state['last_nodes_hash'] = nodes_hash
        p4_state['applying_heuristics'] = False  # Reset flag

    # If heuristics never populated but auto-run flagged as done, allow rerun
    if not p4_state['heuristics_data'] and p4_state.get('auto_heuristics_done'):
        p4_state['auto_heuristics_done'] = False

    # Auto-run heuristics once if pathway data exists
    if nodes and not p4_state['heuristics_data'] and not p4_state['auto_heuristics_done']:
        with ai_activity("Analyzing usability heuristics…"):
            # Limit nodes in prompt to avoid token overflow, but include full content
            nodes_display = nodes[:15] if len(nodes) > 15 else nodes
            pathway_summary = f"Clinical Decision Pathway: {len(nodes)} nodes, {sum(1 for n in nodes if n.get('type') == 'Decision')} decision points"
            
            prompt = f"""You are a UX/Clinical Informatics expert specializing in Nielsen's 10 Usability Heuristics.

TASK: Analyze this clinical decision pathway against all 10 heuristics. For EACH heuristic (H1-H10):
1. Identify current state (strength or gap)
2. Provide 2-3 specific, actionable recommendations
3. Frame improvements in context of clinical decision support (clarity, safety, efficiency)

CRITICAL DISTINCTION - Do NOT confuse:
- Literature references (PMID xxxxx) = These are CORRECT and necessary; they are NOT jargon, they are evidence citations
- Medical jargon = Clinical abbreviations (SIRS, qSOFA, troponin) that may need explanation for non-specialist users or less common abbreviations that lack context
- Only flag jargon when it's unclear or lacks clinical definition; literature citations are appropriate

Pathway Overview: {pathway_summary}
Nodes analyzed ({len(nodes_display)} sample): {json.dumps(nodes_display, indent=2)}

HEURISTIC DEFINITIONS:
- H1 (Visibility): System keeps users informed of status, critical values, decision points
- H2 (Match system/real-world): Use clinician language, avoid unclear jargon or explain it; literature citations are appropriate
- H3 (User control & freedom): Provide escape routes, undo options, alternative pathways
- H4 (Consistency): Standardize terminology, node types, decision structures, formats
- H5 (Error prevention): Prevent wrong decisions through validation, constraints, alerts
- H6 (Recognition over recall): Make options visible; minimize memory load; use visual cues
- H7 (Flexibility & efficiency): Support both novice (guided) and expert (accelerated) use
- H8 (Aesthetic & minimalist): Remove clutter; keep essential clinical information prominent
- H9 (Error recovery): Help clinicians recognize, diagnose, and recover from decision errors
- H10 (Help & documentation): Provide context-specific guidance, evidence citations, rationale

EVALUATION FRAMEWORK:
For pathway-applicable heuristics (H2, H4, H5, H9): Focus on decision nodes, clinical specificity
For UI-design heuristics (H1, H3, H6, H7, H8, H10): Note as "UI layer concern" but still assess

Return ONLY valid JSON with exactly these keys: H1, H2, H3, H4, H5, H6, H7, H8, H9, H10
Each value: ONLY actionable recommendations (2-3 bullet points or numbered items). Do NOT include evaluation of current state.

EXAMPLE FORMAT:
{{
  "H1": "• Add node showing 'Reassess response at 24h' after treatment initiation\n• Include vital sign thresholds for escalation triggers\n• Surface critical alerts (e.g., sepsis criteria) prominently in decision labels",
  "H2": "• Expand less common abbreviations (e.g., SIRS definition) for novice users\n• Add node annotation explaining scoring systems (qSOFA)\n• Maintain all literature references—they strengthen credibility",
  ...
}}"""
            # Use native function calling for reliable structured heuristics output
            result = get_gemini_response(
                prompt, 
                function_declaration=ANALYZE_HEURISTICS,
                thinking_budget=1024
            )
            # Extract heuristics from function call or fall back
            if isinstance(result, dict) and 'arguments' in result:
                res = result['arguments']
            elif isinstance(result, dict) and 'H1' in result:
                res = result
            else:
                # Fallback to json_mode
                res = get_gemini_response(prompt, json_mode=True)
            if res and isinstance(res, dict) and len(res) >= 10:
                p4_state['heuristics_data'] = res
                p4_state['auto_heuristics_done'] = True
                st.rerun()  # Rerun to display heuristics immediately
            else:
                # Don't mark as done so it retries on next interaction
                st.warning(f"Heuristics generation incomplete. Received {len(res) if res else 0} of 10. Please try again or navigate away and back.")
    elif not nodes:
        # Debug: show why heuristics aren't generating
        st.info("Waiting for pathway nodes from Phase 3...")
    elif p4_state.get('auto_heuristics_done') and not p4_state.get('heuristics_data'):
        # Debug: generation was attempted but failed
        st.warning("Heuristics generation failed. Retrying...")
        p4_state['auto_heuristics_done'] = False
        st.rerun()

    # Prepare nodes for visualization with a lightweight cache
    nodes_for_viz = nodes if nodes else [
        {"label": "Start", "type": "Start"},
        {"label": "Add nodes in Phase 3", "type": "Process"},
        {"label": "End", "type": "End"},
    ]
    cache = p4_state.setdefault('viz_cache', {})
    sig = hashlib.md5(json.dumps(nodes_for_viz, sort_keys=True).encode('utf-8')).hexdigest()

    # FULLSCREEN-ONLY VISUALIZATION + RIGHT-SIDE HEURISTICS LAYOUT
    # Build or retrieve SVG for visualization (no inline preview)
    svg_bytes = cache.get(sig, {}).get("svg")
    
    # Always regenerate SVG fresh to ensure downloads work after heuristics are applied
    # The cache can sometimes get stale after rerun
    if svg_bytes is None or p4_state.get('applied_status'):
        g = build_graphviz_from_nodes(nodes_for_viz, "TD")
        if g:
            new_svg = render_graphviz_bytes(g, "svg")
            if new_svg:
                cache[sig] = {"svg": new_svg}
                svg_bytes = new_svg
                debug_log(f"SVG regenerated: {len(new_svg)} bytes")
            else:
                debug_log("render_graphviz_bytes returned None - graphviz may not be rendering")
        else:
            debug_log(f"build_graphviz_from_nodes returned None. graphviz module: {graphviz}")
    p4_state['viz_cache'] = {sig: cache.get(sig, {})}

    import base64
    svg_b64 = base64.b64encode(svg_bytes or b"").decode('utf-8') if svg_bytes else ""
    
    # Also decode SVG to string for direct embedding
    svg_str = svg_bytes.decode('utf-8') if svg_bytes else ""

    # Always prepare DOT source as a client-side fallback via Viz.js
    dot_src = dot_from_nodes(nodes_for_viz, "TD")

    # DEBUG: Log SVG generation status
    debug_log(f"SVG generation - graphviz: {graphviz is not None}, svg_bytes: {svg_bytes is not None}, svg_b64 len: {len(svg_b64)}, svg_str len: {len(svg_str)}")

    # SINGLE COLUMN LAYOUT for better UX flow
    # Order: Pathway Visualization → Nielsen Heuristics → Edit Pathway Data → Refine & Regenerate
    
    # ========== 1. PATHWAY VISUALIZATION ==========
    st.subheader("Pathway Visualization")
    
    if svg_bytes:
        st.download_button("📊 Download SVG", svg_bytes, file_name="pathway.svg", mime="image/svg+xml", help="High-quality vector graphic for presentations and email", use_container_width=False)
        st.caption("💡 Re-download after applying heuristics to get the updated pathway.")
    else:
        st.warning("SVG unavailable. Install Graphviz on the server and retry.")

    st.divider()

    # ========== 2. NIELSEN'S HEURISTICS EVALUATION ==========
    st.subheader("Nielsen's Heuristics Evaluation")
    h_data = p4_state.get('heuristics_data', {})

    if not h_data:
        # Manual generation button - user controls when to generate heuristics
        if nodes and st.button("Generate Heuristics Now", key="p4_manual_heuristics", type="secondary"):
            p4_state['auto_heuristics_done'] = False
            st.rerun()
    else:
        # Display ALL heuristics H1-H10 without pre-filtering
        # AI will intelligently evaluate each and apply only those that improve the pathway
        ordered_keys = sorted(h_data.keys(), key=lambda k: int(k[1:]) if k[1:].isdigit() else k)
        st.caption("Review each heuristic. AI agent will evaluate all and apply those that improve pathway structure. Results will show what was applied and what was skipped.")

        for heuristic_key in ordered_keys:
            insight = h_data.get(heuristic_key, "")
            # Format insight as bullet list if it's a list, or format string bullets properly
            if isinstance(insight, list):
                # Convert list to HTML bullet points
                formatted_insight = "<ul style='margin: 0; padding-left: 20px;'>" + "".join(f"<li>{item}</li>" for item in insight) + "</ul>"
            elif isinstance(insight, str):
                # Convert string with bullet characters or newlines to HTML list
                lines = [line.strip() for line in insight.replace('•', '\n').split('\n') if line.strip()]
                if len(lines) > 1:
                    formatted_insight = "<ul style='margin: 0; padding-left: 20px;'>" + "".join(f"<li>{line.lstrip('- •')}</li>" for line in lines) + "</ul>"
                else:
                    formatted_insight = insight
            else:
                formatted_insight = str(insight)
            # Get label from HEURISTIC_DEFS
            label_stub = HEURISTIC_DEFS.get(heuristic_key, "Heuristic").split(' (')[0].split(':')[0]

            with st.expander(f"**{heuristic_key}** - {label_stub}", expanded=False):
                st.caption(f"*{HEURISTIC_DEFS.get(heuristic_key, 'N/A')}*")
                st.markdown(
                    f"<div style='background-color: white; color: black; padding: 12px; border-radius: 5px; border: 1px solid #ddd; border-left: 4px solid #5D4037; margin-top: 8px;'>{formatted_insight}</div>",
                    unsafe_allow_html=True
                )

        # APPLY + UNDO BUTTONS
        has_heuristics = bool(h_data)
        if has_heuristics:
            if p4_state.get('applied_status') and p4_state.get('applied_summary_detail'):
                st.success("✅ Applied")
                with st.expander("View Changes Made", expanded=True):
                    st.markdown("**Changes Made:**")
                    st.markdown(p4_state['applied_summary_detail'])
            
            col_apply, col_undo = st.columns([1, 1])
            with col_apply:
                btn_applied = p4_state.get('applied_status', False)
                btn_label = "Applied ✓" if btn_applied else "Apply"
                btn_type = "primary" if btn_applied else "secondary"
                if st.button(btn_label, key="p4_apply_all_actionable", type=btn_type, disabled=btn_applied):
                    # Initialize history if needed
                    if 'nodes_history' not in p4_state:
                        p4_state['nodes_history'] = []
                    
                    # Save current state to history BEFORE applying
                    p4_state['nodes_history'].append(copy.deepcopy(nodes))
                    p4_state['applying_heuristics'] = True  # Set flag to prevent re-analysis
                    
                    with ai_activity("Applying heuristics…"):
                        improved_nodes, applied_heuristics, apply_summary = apply_actionable_heuristics_incremental(nodes, h_data)
                        if improved_nodes and len(improved_nodes) > 0:
                            st.session_state.data['phase3']['nodes'] = harden_nodes(improved_nodes)
                            p4_state['viz_cache'] = {}
                            p4_state['applied_status'] = True
                            p4_state['applied_heuristics'] = applied_heuristics
                            p4_state['applied_summary_detail'] = apply_summary
                            st.rerun()
                        else:
                            # Remove from history if apply failed
                            if p4_state.get('nodes_history'):
                                p4_state['nodes_history'].pop()
                            st.error("Could not process recommendations. AI returned no valid nodes. Please try again.")

            with col_undo:
                history_count = len(p4_state.get('nodes_history', []))
                undo_disabled = history_count == 0
                if st.button("Undo Last Changes", key="p4_undo_all", type="secondary", disabled=undo_disabled):
                    if p4_state.get('nodes_history') and len(p4_state['nodes_history']) > 0:
                        prev_nodes = p4_state['nodes_history'].pop()
                        st.session_state.data['phase3']['nodes'] = prev_nodes
                        p4_state['viz_cache'] = {}
                        p4_state['applied_status'] = False
                        p4_state['applied_summary_detail'] = ""
                        p4_state['applied_heuristics'] = []
                        st.success(f"✓ Reverted to previous version ({len(prev_nodes)} nodes)")
                        st.rerun()
                    else:
                        st.info("No changes to undo")
                if history_count > 0:
                    st.caption(f"{history_count} version(s) in history")

    st.divider()

    # ========== 3. EDIT PATHWAY DATA MANUALLY ==========
    st.subheader("Edit Pathway Data")
    with st.expander("Edit Pathway Data", expanded=False):
        df_p4 = pd.DataFrame(nodes)
        if 'node_id' not in df_p4.columns:
            df_p4.insert(0, 'node_id', range(1, len(df_p4) + 1))
        else:
            df_p4['node_id'] = range(1, len(df_p4) + 1)
        # Ensure notes column exists (may be 'detail' in older data)
        if 'notes' not in df_p4.columns:
            if 'detail' in df_p4.columns:
                df_p4['notes'] = df_p4['detail']
            else:
                df_p4['notes'] = ""
        # Show node_id, type, label, notes columns (remove evidence for cleaner view)
        display_cols = ['node_id', 'type', 'label', 'notes'] if 'notes' in df_p4.columns else ['node_id', 'type', 'label']
        display_cols = [col for col in display_cols if col in df_p4.columns]
        df_p4_display = df_p4[display_cols]
        edited_p4_display = st.data_editor(
            df_p4_display,
            column_config={
                "node_id": st.column_config.NumberColumn("ID", disabled=True, width="small"),
                "type": st.column_config.SelectboxColumn(
                    "Type",
                    options=["Start", "Decision", "Process", "End"],
                    required=True,
                    width="small"
                ),
                "label": st.column_config.TextColumn("Clinical Step", width="large"),
                "notes": st.column_config.TextColumn(
                    "Notes",
                    width="medium",
                    help="Actionable details: red flags, thresholds, monitoring (appears in notes legend)"
                )
            },
            num_rows="dynamic",
            key="p4_editor"
        )
        manual_changed = not df_p4_display.equals(edited_p4_display)
        if manual_changed:
            if 'node_id' in edited_p4_display.columns:
                edited_p4_display = edited_p4_display.drop('node_id', axis=1)
            # Preserve evidence column from original nodes
            for idx, row in edited_p4_display.iterrows():
                if 'evidence' in df_p4.columns and idx < len(df_p4):
                    edited_p4_display.at[idx, 'evidence'] = df_p4.at[idx, 'evidence']
            st.session_state.data['phase3']['nodes'] = edited_p4_display.to_dict('records')
            p4_state['viz_cache'] = {}
            st.info("Nodes updated. Click 'Regenerate Visualization & Downloads' to refresh.")

        regen_disabled = not manual_changed and not st.session_state.data['phase3'].get('nodes')
        if st.button("Regenerate Visualization & Downloads", key="p4_manual_regen", disabled=regen_disabled):
            p4_state['viz_cache'] = {}
            st.success("Visualization regenerated with latest edits. Open fullscreen or download updated SVG.")
            st.rerun()

    st.divider()

    # ========== 4. REFINE AND REGENERATE ==========
    with st.expander("Refine & Regenerate", expanded=False):
        st.markdown("**Tip:** Describe any desired modifications and optionally attach supporting documents. Click \"Regenerate\" to automatically update Phase 4 content and downloads.")
        regen_submitted = False
        with st.form("p4_refine_form"):
            col_text, col_file = columns_top([2, 1])
            with col_text:
                refine_notes = st.text_area(
                    "Refinement Notes",
                    placeholder="Consolidate redundant steps; add alerts for critical values; use patient-friendly terms",
                    key="p4_refine_notes",
                    height=90,
                    label_visibility="visible"
                )

            with col_file:
                st.caption("Supporting Documents (optional)")
                uploaded = st.file_uploader(
                    "Drag and drop file here",
                    key="p4_upload",
                    accept_multiple_files=False,
                    label_visibility="collapsed",
                    help="Limit 200MB per file"
                )
                if uploaded:
                    file_result = upload_and_review_file(uploaded, "p4_refine_file", "pathway")
                    if file_result:
                        with st.expander("File Review", expanded=False):
                            st.markdown(file_result["review"])

            col_form_gap, col_form_btn = st.columns([3, 1])
            with col_form_btn:
                regen_submitted = st.form_submit_button("Regenerate", type="secondary", use_container_width=True)

        apply_refine_clicked = st.button("Apply Refinements", key="p4_apply_refine", type="primary", use_container_width=True)

        if regen_submitted or apply_refine_clicked:
            refine_with_file = st.session_state.get('p4_refine_notes', '').strip()
            if refine_with_file:
                if st.session_state.get("file_p4_refine_file"):
                    refine_with_file += f"\n\n**Supporting Document:**\n{st.session_state.get('file_p4_refine_file')}"
                with st.spinner("Applying refinements..."):
                    refined = regenerate_nodes_with_refinement(nodes, refine_with_file, h_data) if 'regenerate_nodes_with_refinement' in globals() else None
                    if refined:
                        st.session_state.data['phase3']['nodes'] = refined
                        p4_state['viz_cache'] = {}
                        st.success("Refinements applied. Regenerated nodes below.")
                        st.rerun()
                    else:
                        st.warning("Could not apply refinements. Please try different notes or regenerate.")

    render_bottom_navigation()
    st.stop()

# --- PHASE 5 ---
elif "Operationalize" in phase or "Deploy" in phase:
    st.header(f"Phase 5. {PHASES[4]}")
    
    # Import Phase 5 helpers
    try:
        from phase5_helpers import (
            generate_expert_form_html,
            generate_beta_form_html,
            generate_education_module_html,
            create_phase5_executive_summary_docx,
            ensure_carepathiq_branding
        )
    except ImportError:
        st.error("Phase 5 helpers not found. Please ensure phase5_helpers.py is in the workspace.")
        st.stop()
    
    # Single info box at top
    styled_info("<b>Tip:</b> Click 'Generate' for each deliverable you need. Downloads appear after generation completes.")
    
    cond = st.session_state.data['phase1']['condition'] or "Pathway"
    setting = st.session_state.data['phase1'].get('setting', '') or ""
    nodes = st.session_state.data['phase3']['nodes'] or []
    
    if not nodes:
        st.warning("Complete Phase 3 first to generate deliverables.")
        render_bottom_navigation()
        st.stop()
    
    # 2x2 GRID LAYOUT - Each deliverable has Generate button + Download
    col1, col2 = st.columns(2)

    # ========== TOP LEFT: EXPERT PANEL FEEDBACK ==========
    with col1:
        st.markdown("<h3>Expert Panel Feedback</h3>", unsafe_allow_html=True)
        
        # Generate button
        if st.button("Generate Expert Feedback Form", key="p5_gen_expert", type="secondary"):
            with st.spinner("Generating expert feedback form..."):
                try:
                    g = build_graphviz_from_nodes(nodes, "TD")
                    svg_bytes = render_graphviz_bytes(g, "svg") if g else None
                    svg_b64 = base64.b64encode(svg_bytes).decode('utf-8') if svg_bytes else None
                    
                    expert_html = generate_expert_form_html(
                        condition=cond,
                        nodes=nodes,
                        organization=cond,
                        care_setting=setting,
                        pathway_svg_b64=svg_b64,
                        genai_client=get_genai_client()
                    )
                    st.session_state.data['phase5']['expert_html'] = ensure_carepathiq_branding(expert_html)
                    st.success("✓ Generated!")
                    st.rerun()
                except Exception as e:
                    if '429' in str(e) or 'quota' in str(e).lower():
                        st.error("⏳ API rate limit. Wait 15-30 seconds and try again.")
                    else:
                        st.error(f"Error: {str(e)[:100]}")
        
        # Download button (only shows if generated)
        if st.session_state.data['phase5'].get('expert_html'):
            st.download_button(
                "📥 Download Expert Feedback (.html)",
                st.session_state.data['phase5']['expert_html'],
                f"ExpertFeedback_{cond.replace(' ', '_')}.html",
                "text/html",
                key="p5_dl_expert"
            )

    # ========== TOP RIGHT: BETA TESTING GUIDE ==========
    with col2:
        st.markdown("<h3>Beta Testing Guide</h3>", unsafe_allow_html=True)
        
        # Generate button
        if st.button("Generate Beta Testing Guide", key="p5_gen_beta", type="secondary"):
            with st.spinner("Generating beta testing guide..."):
                try:
                    beta_html = generate_beta_form_html(
                        condition=cond,
                        nodes=nodes,
                        organization=cond,
                        care_setting=setting,
                        genai_client=get_genai_client(),
                        phase1_data=st.session_state.data.get('phase1', {}),
                        phase2_data=st.session_state.data.get('phase2', {}),
                        phase3_data=st.session_state.data.get('phase3', {}),
                        phase4_data=st.session_state.data.get('phase4', {})
                    )
                    st.session_state.data['phase5']['beta_html'] = ensure_carepathiq_branding(beta_html)
                    st.success("✓ Generated!")
                    st.rerun()
                except Exception as e:
                    if '429' in str(e) or 'quota' in str(e).lower():
                        st.error("⏳ API rate limit. Wait 15-30 seconds and try again.")
                    else:
                        st.error(f"Error: {str(e)[:100]}")
        
        # Download button (only shows if generated)
        if st.session_state.data['phase5'].get('beta_html'):
            st.download_button(
                "📥 Download Beta Testing (.html)",
                st.session_state.data['phase5']['beta_html'],
                f"BetaTestingGuide_{cond.replace(' ', '_')}.html",
                "text/html",
                key="p5_dl_beta"
            )
    
    st.divider()
    
    col3, col4 = st.columns(2)
    
    # ========== BOTTOM LEFT: EDUCATION MODULE ==========
    with col3:
        st.markdown("<h3>Education Module</h3>", unsafe_allow_html=True)
        # Hide "Press Enter to apply" hint for this section
        st.markdown("<style>div[data-testid='stTextInput'] div[data-testid='InputInstructions'] { display: none; }</style>", unsafe_allow_html=True)
        
        # Target audience input (use value="" to suppress "Press Enter to apply" hint)
        aud_edu = st.text_input(
            "Target Audience",
            value="",
            placeholder="e.g., Residents, Nursing Staff",
            key="p5_aud_edu_input",
            help="Enter the target audience for the education module"
        )
        
        # Generate button
        if st.button("Generate Education Module", key="p5_gen_edu", type="secondary"):
            if not aud_edu:
                st.warning("Please enter target audience first.")
            else:
                with st.spinner("Generating education module..."):
                    try:
                        edu_html = generate_education_module_html(
                            condition=cond,
                            nodes=nodes,
                            target_audience=aud_edu,
                            care_setting=setting,
                            genai_client=get_genai_client()
                        )
                        st.session_state.data['phase5']['edu_html'] = ensure_carepathiq_branding(edu_html)
                        st.success("✓ Generated!")
                        st.rerun()
                    except Exception as e:
                        if '429' in str(e) or 'quota' in str(e).lower():
                            st.error("⏳ API rate limit. Wait 15-30 seconds and try again.")
                        else:
                            st.error(f"Error: {str(e)[:100]}")
        
        # Download button (only shows if generated)
        if st.session_state.data['phase5'].get('edu_html'):
            st.download_button(
                "📥 Download Education Module (.html)",
                st.session_state.data['phase5']['edu_html'],
                f"EducationModule_{cond.replace(' ', '_')}.html",
                "text/html",
                key="p5_dl_edu"
            )

    # ========== BOTTOM RIGHT: EXECUTIVE SUMMARY ==========
    with col4:
        st.markdown("<h3>Executive Summary</h3>", unsafe_allow_html=True)
        
        # Generate button
        if st.button("Generate Executive Summary", key="p5_gen_exec", type="secondary"):
            with st.spinner("Generating executive summary..."):
                try:
                    docx_bytes = create_phase5_executive_summary_docx(
                        data=st.session_state.data,
                        condition=cond,
                        genai_client=get_genai_client()
                    )
                    st.session_state['exec_docx_bytes'] = docx_bytes
                    st.session_state.data['phase5']['exec_summary'] = f"Executive Summary for {cond}"
                    st.success("✓ Generated!")
                    st.rerun()
                except Exception as e:
                    if '429' in str(e) or 'quota' in str(e).lower():
                        st.error("⏳ API rate limit. Wait 15-30 seconds and try again.")
                    else:
                        st.error(f"Error: {str(e)[:100]}")
        
        # Download button (only shows if generated)
        if st.session_state.get('exec_docx_bytes'):
            st.download_button(
                "📥 Download Executive Summary (.docx)",
                st.session_state['exec_docx_bytes'],
                f"ExecutiveSummary_{cond.replace(' ', '_')}.docx",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="p5_dl_exec"
            )
    
    render_bottom_navigation()
    st.stop()

# Footer is now rendered within each phase via render_bottom_navigation()