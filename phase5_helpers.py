"""
Phase 5 Helper Functions for Shareable HTML & Document Generation
Generates standalone HTML files for Expert Panel, Beta Testing, and Education Module
Plus Word documents for Executive Summaries
"""

import json
import base64
import time
import hashlib
import re
from io import BytesIO
from datetime import datetime
from difflib import SequenceMatcher

# ==========================================
# HELPER FUNCTIONS 
# ==========================================


def merge_hybrid_intelligently(fc_nodes, gl_nodes):
    """Merge flowchart + guideline nodes without duplicates."""
    merged = []
    used_gl_indices = set()
    
    # PHASE 1: Match flowchart nodes to guideline nodes
    for fc_node in fc_nodes:
        best_match_idx = None
        best_score = 0
        
        for idx, gl_node in enumerate(gl_nodes):
            if idx in used_gl_indices:
                continue
            
            label_sim = SequenceMatcher(None, 
                                       fc_node.get('label', '').lower(), 
                                       gl_node.get('label', '').lower()).ratio()
            type_match = 1.0 if fc_node.get('type') == gl_node.get('type') else 0.5
            score = (label_sim * 0.7) + (type_match * 0.3)
            
            if score > best_score:
                best_score = score
                best_match_idx = idx
        
        # Merge if high similarity (>80%)
        if best_match_idx is not None and best_score > 0.8:
            gl_node = gl_nodes[best_match_idx]
            merged_node = {
                "type": fc_node.get("type", "Process"),
                "label": fc_node.get("label", ""),
                "evidence": gl_node.get("evidence", "N/A"),
                "detail": gl_node.get("detail", ""),
                "branches": fc_node.get("branches", [])
            }
            used_gl_indices.add(best_match_idx)
        elif best_match_idx is not None and best_score > 0.5:
            merged_node = fc_node.copy()
            merged_node["detail"] = merged_node.get("detail", "") + \
                                   f"\n[Similar to: {gl_nodes[best_match_idx].get('label', '')}]"
        else:
            merged_node = fc_node.copy()
        
        merged.append(merged_node)
    
    # PHASE 2: Add remaining guideline nodes
    for idx, gl_node in enumerate(gl_nodes):
        if idx not in used_gl_indices:
            merged.append(gl_node)
    
    # PHASE 3: Deduplicate
    seen = set()
    deduplicated = []
    for node in merged:
        key = (node.get('label', '').strip().lower(), node.get('type', ''))
        if key not in seen:
            deduplicated.append(node)
            seen.add(key)
    
    # PHASE 4: Reorder
    return reorder_nodes_topologically(deduplicated)


def reorder_nodes_topologically(nodes):
    """Reorder nodes: Start first, End last."""
    start_nodes = [n for n in nodes if n.get('type') == 'Start']
    end_nodes = [n for n in nodes if n.get('type') == 'End']
    mid_nodes = [n for n in nodes if n.get('type') not in ['Start', 'End']]
    
    if not start_nodes:
        start_nodes = [{"type": "Start", "label": "Patient presents", "evidence": "N/A"}]
    
    if not end_nodes:
        if mid_nodes:
            mid_nodes[-1]["type"] = "End"
            end_nodes = [mid_nodes.pop()]
        else:
            end_nodes = [{"type": "End", "label": "Disposition", "evidence": "N/A"}]
    
    return start_nodes[:1] + mid_nodes + end_nodes


def enrich_nodes_with_pmids(nodes, pmids):
    """Map extracted PMIDs to nodes based on section references."""
    for node in nodes:
        node_section = node.get("evidence", "")
        relevant_pmids = [p.get("pmid", "") for p in pmids 
                         if p.get("section", "").lower() in node_section.lower()]
        
        if relevant_pmids:
            node["evidence"] = relevant_pmids[0]
            if len(relevant_pmids) > 1:
                node["detail"] = node.get("detail", "") + \
                                f"\n\nAdditional evidence: PMIDs {', '.join(relevant_pmids[1:])}"
    
    return nodes


def auto_fix_pathway(nodes):
    """Aggressively fix validation issues."""
    fixes = []
    
    if not nodes:
        return nodes, fixes
    
    # Fix 1: Ensure Start at index 0
    if nodes[0].get('type') != 'Start':
        nodes.insert(0, {"type": "Start", "label": "Patient presents", "evidence": "N/A"})
        fixes.append("Added missing Start node at index 0")
    
    # Fix 2: Ensure at least one End
    if not any(n.get('type') == 'End' for n in nodes):
        nodes[-1]['type'] = 'End'
        fixes.append(f"Converted last node to End: {nodes[-1].get('label', 'N/A')}")
    
    # Fix 3: Remove invalid branch targets
    for node in nodes:
        if 'branches' in node:
            valid_branches = []
            for branch in node['branches']:
                target = branch.get('target')
                if isinstance(target, int) and 0 <= target < len(nodes):
                    valid_branches.append(branch)
                else:
                    fixes.append(f"Removed invalid branch target {target} from node '{node.get('label', 'N/A')}'")
            node['branches'] = valid_branches
    
    return nodes, fixes


def calculate_extraction_confidence(nodes, doc_type):
    """Score extraction quality (0.0-1.0)."""
    if not nodes:
        return 0.0
    
    score = 0.0
    
    # Node count (ideal: 8-35)
    if 8 <= len(nodes) <= 35:
        score += 0.3
    elif len(nodes) < 8:
        score += 0.1
    
    # Required types
    types = [n.get("type") for n in nodes]
    if "Start" in types:
        score += 0.2
    if "End" in types:
        score += 0.2
    if "Decision" in types:
        score += 0.1
    
    # Decision nodes have branches
    decision_count = sum(1 for n in nodes if n.get("type") == "Decision")
    valid_decisions = sum(1 for n in nodes 
                         if n.get("type") == "Decision" and n.get("branches"))
    if decision_count > 0:
        score += 0.1 * (valid_decisions / decision_count)
    
    # Evidence coverage
    with_evidence = sum(1 for n in nodes 
                       if n.get("evidence") and n["evidence"] != "N/A")
    score += 0.1 * (with_evidence / len(nodes))
    
    return round(score, 2)


# ==========================================
# SHARED STYLING & CONSTANTS
# ==========================================

CAREPATHIQ_COLORS = {
    "brown": "#5D4037",
    "brown_dark": "#3E2723",
    "teal": "#A9EED1",
    "light_gray": "#f5f5f5",
    "border_gray": "#ddd",
}

CAREPATHIQ_FOOTER = """
<div class="carepathiq-footer">
    <p style="margin: 10px 0;">
        <strong>CarePathIQ</strong> © 2024 by Tehreem Rehman
    </p>
    <p style="margin: 5px 0; font-size: 0.9em;">
        Licensed under <a href="https://creativecommons.org/licenses/by-sa/4.0/" target="_blank">CC BY-SA 4.0</a>
    </p>
</div>
"""

SHARED_CSS = """
:root {
    --brown: #5D4037;
    --brown-dark: #3E2723;
    --teal: #A9EED1;
    --light-gray: #f5f5f5;
    --border-gray: #ddd;
}

* {
    margin: 0;
    padding: 0;
    box-sizing: border-box;
}

body {
    font-family: 'Segoe UI', Tahoma, Verdana, sans-serif;
    line-height: 1.6;
    color: #333;
    background-color: #fafafa;
}

.container {
    max-width: 900px;
    margin: 0 auto;
    padding: 20px;
    background-color: white;
    border-radius: 8px;
    box-shadow: 0 2px 8px rgba(0,0,0,0.1);
}

.header {
    text-align: center;
    color: var(--brown-dark);
    border-bottom: 3px solid var(--teal);
    padding-bottom: 15px;
    margin-bottom: 30px;
}

.header h1 {
    font-size: 2em;
    margin-bottom: 5px;
}

.header p {
    font-size: 1.1em;
    color: #666;
}

.form-group {
    margin: 20px 0;
}

label {
    font-weight: 600;
    display: block;
    margin-bottom: 8px;
    color: var(--brown-dark);
}

input[type="text"],
input[type="email"],
input[type="number"],
textarea,
select {
    width: 100%;
    padding: 10px;
    border: 1px solid var(--border-gray);
    border-radius: 4px;
    font-family: inherit;
    font-size: 1em;
}

textarea {
    min-height: 100px;
    resize: vertical;
}

input:focus,
textarea:focus,
select:focus {
    outline: none;
    border-color: var(--teal);
    box-shadow: 0 0 5px rgba(169, 238, 209, 0.3);
}

button {
    background-color: var(--brown);
    color: white;
    padding: 12px 24px;
    border: none;
    border-radius: 4px;
    cursor: pointer;
    font-size: 1em;
    font-weight: 600;
    transition: background-color 0.3s ease;
    margin-right: 10px;
    margin-bottom: 10px;
}

button:hover {
    background-color: var(--brown-dark);
}

button:active {
    transform: scale(0.98);
}

.button-group {
    margin: 20px 0;
    display: flex;
    flex-wrap: wrap;
    gap: 10px;
}

.node-card {
    border: 1px solid var(--border-gray);
    padding: 15px;
    margin: 15px 0;
    border-radius: 6px;
    background-color: var(--light-gray);
}

.node-card h3 {
    color: var(--brown);
    margin-bottom: 8px;
    display: flex;
    align-items: center;
}

.node-badge {
    background-color: var(--teal);
    color: var(--brown-dark);
    padding: 2px 8px;
    border-radius: 12px;
    font-size: 0.85em;
    margin-right: 10px;
    font-weight: 600;
}

.node-type {
    font-size: 0.9em;
    color: #666;
    margin-bottom: 10px;
}

.checkbox-group {
    display: flex;
    align-items: center;
    margin: 10px 0;
}

.checkbox-group input[type="checkbox"] {
    width: auto;
    margin-right: 10px;
    cursor: pointer;
}

.checkbox-group label {
    margin: 0;
    cursor: pointer;
    font-weight: 500;
}

.expandable-section {
    margin-left: 20px;
    margin-top: 10px;
    padding: 10px;
    background-color: white;
    border-left: 3px solid var(--teal);
    display: none;
}

.expandable-section.show {
    display: block;
}

.progress-bar {
    width: 100%;
    height: 8px;
    background-color: var(--border-gray);
    border-radius: 4px;
    overflow: hidden;
    margin: 10px 0;
}

.progress-fill {
    height: 100%;
    background-color: var(--teal);
    transition: width 0.3s ease;
}

.carepathiq-footer {
    text-align: center;
    margin-top: 40px;
    padding-top: 20px;
    border-top: 1px solid var(--border-gray);
    color: #666;
    font-size: 0.9em;
}

.carepathiq-footer a {
    color: var(--brown);
    text-decoration: none;
}

.carepathiq-footer a:hover {
    text-decoration: underline;
}

@media (max-width: 768px) {
    .container {
        padding: 15px;
    }
    
    .header h1 {
        font-size: 1.5em;
    }
    
    .button-group {
        flex-direction: column;
    }
    
    button {
        width: 100%;
    }
}
"""

# ==========================================
# EXPERT PANEL FEEDBACK FORM
# ==========================================

def generate_expert_form_html(
    condition: str,
    nodes: list,
    organization: str = "CarePathIQ",
    care_setting: str = "",
    pathway_svg_b64: str = None,
    genai_client=None
) -> str:
    """
    Generate standalone expert panel feedback form with CSV download capability.
    
    Args:
        condition: Clinical condition being reviewed
        nodes: List of pathway nodes (dicts with 'type', 'label', 'evidence')
        organization: Organization name
        care_setting: Care setting/environment (e.g., "Emergency Department")
        genai_client: Optional Google Generative AI client (not used)
        
    Returns:
        Complete standalone HTML string
    """
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    nodes_json = json.dumps(nodes)
    condition_clean = (condition or "Pathway").strip()
    care_setting_clean = (care_setting or "").strip()
    if care_setting_clean:
        pathway_title = f"Pathway: Managing {condition_clean} in {care_setting_clean}"
        page_title = f"Expert Panel Feedback: {condition_clean} ({care_setting_clean})"
    else:
        pathway_title = f"Pathway: Managing {condition_clean}"
        page_title = f"Expert Panel Feedback: {condition_clean}"
    
    # Pathway visualization removed
    pathway_button_html = ""
    pathway_script = ""
    
    html = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{page_title}</title>
    <style>
        {SHARED_CSS}
        .info-grid {{display:grid;grid-template-columns:1fr 1fr 1fr;gap:12px;margin-bottom:20px}}
        @media (max-width: 768px) {{
            .info-grid {{grid-template-columns:1fr}}
        }}
        .compact-node {{border:1px solid var(--border-gray);border-radius:6px;padding:15px;margin-bottom:15px;background:#fafafa}}
        .compact-node h4 {{margin:0 0 8px 0;color:var(--brown-dark);font-size:1em}}
        .compact-node .node-meta {{font-size:0.85em;color:#666;margin-bottom:10px}}
        .feedback-section {{display:none;margin-top:12px;padding-top:12px;border-top:1px solid var(--border-gray)}}
        .feedback-section.show {{display:block}}
        .feedback-section textarea {{min-height:70px}}
        .feedback-section select {{margin-bottom:10px}}
    </style>
    {pathway_script}
</head>
<body>
    <div class="container">
        <div class="header">
            <h1>Expert Panel Feedback</h1>
            <p style="font-size: 1.05em; color: var(--brown-dark); margin-bottom: 8px; line-height: 1.6;">{pathway_title}</p>
            <p style="font-size:0.9em;margin-top:5px;color:#666;">Organization: {organization}</p>
            {pathway_button_html}
        </div>

        <form id="feedbackForm">
            <div class="info-grid">
                <div class="form-group">
                    <label for="reviewer_name">Your Name *</label>
                    <input type="text" id="reviewer_name" name="reviewer_name" required placeholder="Full name">
                </div>
                <div class="form-group">
                    <label for="reviewer_email">Your Email *</label>
                    <input type="email" id="reviewer_email" name="reviewer_email" required placeholder="email@institution.org">
                </div>
                <div class="form-group">
                    <label for="reviewer_role">Your Role/Title</label>
                    <input type="text" id="reviewer_role" name="reviewer_role" placeholder="e.g., EM Physician">
                </div>
            </div>

            <hr style="margin: 30px 0; border: none; border-top: 1px solid var(--border-gray);">

            <h2 style="color: var(--brown-dark); margin: 25px 0 15px 0;">Pathway Nodes</h2>
            <p style="margin-bottom: 15px; color: #666; font-size:0.95em;">Check nodes that need feedback. Only those will be included in the download.</p>

            <div id="nodesContainer"></div>

            <hr style="margin: 30px 0; border: none; border-top: 1px solid var(--border-gray);">

            <div class="form-group">
                <label for="overall_feedback">Overall Feedback (Optional)</label>
                <textarea id="overall_feedback" name="overall_feedback" placeholder="Any additional comments about the entire pathway..."></textarea>
            </div>

            <div class="form-group">
                <label for="implementation_barriers">Anticipated Implementation Barriers (Optional)</label>
                <textarea id="implementation_barriers" name="implementation_barriers" placeholder="What challenges do you foresee in implementing this pathway?"></textarea>
            </div>

            <div class="button-group">
                <button type="button" onclick="downloadAsCSV()" style="background:var(--brown);color:white;padding:10px 20px;border:none;border-radius:4px;cursor:pointer;font-size:1em;">Download Responses (CSV)</button>
                <button type="reset">Reset Form</button>
            </div>
        </form>

        {CAREPATHIQ_FOOTER}
    </div>

    <script>
        const pathwayNodes = {nodes_json};
        const condition = "{condition}";
        const timestamp = "{timestamp}";

        // Populate nodes on page load
        document.addEventListener('DOMContentLoaded', function() {{
            renderNodes();
        }});

        function renderNodes() {{
            const container = document.getElementById('nodesContainer');
            pathwayNodes.forEach((node, idx) => {{
                const evidence = node.evidence && node.evidence !== 'N/A' ? `PMID ${{node.evidence}}` : 'No evidence';
                const nodeHTML = `
                    <div class="compact-node">
                        <div style="display:flex;justify-content:space-between;align-items:center">
                            <h4><span class="node-badge">N${{idx + 1}}</span> ${{node.label || 'Step'}}</h4>
                            <label style="margin:0;cursor:pointer;font-weight:normal;display:flex;align-items:center;white-space:nowrap">
                                <input type="checkbox" id="feedback_check_${{idx}}" onchange="toggleExpansion(${{idx}})" style="width:auto;margin-right:6px;margin-top:0">
                                Provide Feedback
                            </label>
                        </div>
                        <div class="node-meta">Type: ${{node.type || 'Process'}} | Evidence: ${{evidence}}</div>

                        <div id="expansion_${{idx}}" class="feedback-section">
                            <label for="feedback_${{idx}}" style="font-size:0.9em"><strong>Change/Concern *</strong></label>
                            <textarea name="feedback_${{idx}}" id="feedback_${{idx}}" placeholder="Describe issue or suggested improvement..." required></textarea>
                            
                            <label for="source_${{idx}}" style="font-size:0.9em;margin-top:10px"><strong>Source *</strong></label>
                            <select name="source_${{idx}}" id="source_${{idx}}" required>
                                <option value="">-- Select Justification Source --</option>
                                <option value="Peer-Reviewed Literature">Peer-Reviewed Literature</option>
                                <option value="National Guideline">National Guideline (ACLS, AHA, etc.)</option>
                                <option value="Institutional Policy">Institutional Policy</option>
                                <option value="Patient Safety Concern">Patient Safety Concern</option>
                                <option value="Feasibility Issue">Feasibility Issue</option>
                                <option value="Other">Other</option>
                            </select>
                            
                            <label for="details_${{idx}}" style="font-size:0.9em;margin-top:10px"><strong>Details/Citation</strong></label>
                            <textarea name="details_${{idx}}" id="details_${{idx}}" placeholder="Reference, PMID, guideline, or rationale..." style="min-height:60px"></textarea>
                        </div>
                    </div>
                `;
                container.insertAdjacentHTML('beforeend', nodeHTML);
            }});
        }}

        function toggleExpansion(nodeIdx) {{
            const expansion = document.getElementById('expansion_' + nodeIdx);
            const checkbox = document.getElementById('feedback_check_' + nodeIdx);
            if (checkbox.checked) {{
                expansion.classList.add('show');
            }} else {{
                expansion.classList.remove('show');
            }}
        }}

        function downloadAsCSV() {{
            const rows = [
                ['Reviewer Name', 'Email', 'Role', 'Node ID', 'Node Label', 'Feedback', 'Source', 'Details']
            ];

            const reviewerName = document.getElementById('reviewer_name').value || 'Anonymous';
            const reviewerEmail = document.getElementById('reviewer_email').value || '';
            const reviewerRole = document.getElementById('reviewer_role').value || '';

            // Collect feedback from nodes
            pathwayNodes.forEach((node, idx) => {{
                const checkbox = document.getElementById('feedback_check_' + idx);
                if (checkbox && checkbox.checked) {{
                    const feedback = document.getElementById('feedback_' + idx)?.value || '';
                    const source = document.getElementById('source_' + idx)?.value || '';
                    const details = document.getElementById('details_' + idx)?.value || '';
                    
                    if (feedback) {{
                        rows.push([
                            reviewerName,
                            reviewerEmail,
                            reviewerRole,
                            'N' + (idx + 1),
                            node.label,
                            feedback,
                            source,
                            details
                        ]);
                    }}
                }}
            }});

            // Add overall feedback if present
            const overallFeedback = document.getElementById('overall_feedback').value;
            const implementationBarriers = document.getElementById('implementation_barriers').value;
            
            if (overallFeedback) {{
                rows.push(['', reviewerEmail, '', '', 'OVERALL FEEDBACK', overallFeedback, '', '']);
            }}
            
            if (implementationBarriers) {{
                rows.push(['', reviewerEmail, '', '', 'IMPLEMENTATION BARRIERS', implementationBarriers, '', '']);
            }}

            // Convert to CSV
            const csv = rows.map(row => 
                row.map(cell => '"' + (cell || '').replace(/"/g, '""') + '"').join(',')
            ).join('\\n');

            // Download
            const safeCondition = condition.replace(/\\s+/g, '_');
            const blob = new Blob([csv], {{ type: 'text/csv;charset=utf-8;' }});
            const link = document.createElement('a');
            const url = URL.createObjectURL(blob);
            link.setAttribute('href', url);
            link.setAttribute('download', `expert_feedback_${{safeCondition}}_${{timestamp}}.csv`);
            link.style.visibility = 'hidden';
            document.body.appendChild(link);
            link.click();
            document.body.removeChild(link);
        }}


    </script>
</body>
</html>"""
    
    return html


# ==========================================
# BETA TESTING FEEDBACK FORM
# ==========================================

def generate_beta_form_html(
    condition: str,
    nodes: list,
    organization: str = "CarePathIQ",
    care_setting: str = "",
    genai_client=None
) -> str:
    """
    Generate simplified beta testing form focused on:
    - 3 scenario-based end-to-end pathway tests
    - Nielsen's 10 heuristics evaluation
    - Overall usability feedback
    - CSV export of results
    
    Args:
        condition: Clinical condition being tested
        nodes: List of pathway nodes (for reference)
        organization: Organization name
        care_setting: Care setting/environment (e.g., "Emergency Department")
        genai_client: Optional Google Generative AI client (not used)
        
    Returns:
        Complete standalone HTML string
    """

    def build_beta_test_scenarios(condition: str, nodes: list, care_setting: str, genai_client=None):
        """Create three concise test scenarios using LLM context with safe fallback."""
        default_scenarios = [
            {
                "title": "Low-Risk Discharge",
                "vignette": "45M, pleuritic chest pain after URI, normal ECG, hs-trop <99th at 0/1h, HEAR 1.",
                "tasks": [
                    "Apply pathway criteria for low-risk presentation",
                    "Choose appropriate workup and confirm no escalation needed",
                    "Select correct disposition: discharge with NSAID + PCP follow-up",
                ],
                "success_criteria": "Did the pathway reach the intended discharge branch?",
                "notes_placeholder": "Describe any mismatch between vignette and end-node...",
            },
            {
                "title": "Moderate-Risk Observation",
                "vignette": "62F, substernal pressure, HTN/HLD, ECG non-ischemic, borderline hs-trop rise, HEART 5.",
                "tasks": [
                    "Follow the moderate-risk / observation branch",
                    "Confirm serial testing or CTA pathway is selected",
                    "Select correct disposition: observation/telemetry pending testing",
                ],
                "success_criteria": "Did the pathway reach the observation/admit branch?",
                "notes_placeholder": "Where did branching feel unclear or incorrect?",
            },
            {
                "title": "High-Risk Escalation",
                "vignette": "58M, diaphoresis, ECG with new ST depressions, elevated troponin.",
                "tasks": [
                    "Trigger the high-risk branch and required meds",
                    "Confirm escalation to cath lab/inpatient cardiology",
                    "Verify no pathway steps block time-sensitive care",
                ],
                "success_criteria": "Did the pathway reach escalation / cath lab branch?",
                "notes_placeholder": "Note any delays or wrong routing for high-risk cases...",
            },
        ]

        if not genai_client:
            return default_scenarios

        node_labels = [n.get("label", "") for n in (nodes or []) if n.get("label")]
        condensed_nodes = ", ".join(node_labels[:12])
        prompt = f"""You are building beta-testing scenarios for a clinical pathway.
Pathway: {condition}
Setting: {care_setting or 'general care'}
Key steps: {condensed_nodes or 'Use common pathway steps'}

Create 3 concise end-to-end test cases for pathway validation. Return JSON array with 3 items only:
[{{
  "title": "Short scenario title",
  "vignette": "45 words max, realistic clinical vignette",
  "tasks": ["3 clear actions for the tester"],
  "success_criteria": "One sentence describing correct end-state to verify",
  "notes_placeholder": "Short guidance for what to log if something breaks"
}}]
Keep language clinical and precise."""

        try:
            response = genai_client.models.generate_content(
                model="gemini-2.5-flash",
                contents=[{"text": prompt}]
            )
            scenarios_raw = json.loads(extract_json_from_response(response.text))
            scenarios = []
            for scenario in scenarios_raw[:3]:
                title = scenario.get("title") or "Scenario"
                vignette = scenario.get("vignette") or "Clinical scenario"
                tasks = [t for t in scenario.get("tasks", []) if t][:3] or ["Follow the pathway steps"]
                success = scenario.get("success_criteria") or "Did the pathway reach the intended outcome?"
                notes = scenario.get("notes_placeholder") or "Describe any mismatch or blockers..."
                scenarios.append({
                    "title": title.strip(),
                    "vignette": vignette.strip(),
                    "tasks": tasks,
                    "success_criteria": success.strip(),
                    "notes_placeholder": notes.strip(),
                })
            if len(scenarios) == 3:
                return scenarios
        except Exception:
            pass

        return default_scenarios
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    nodes_json = json.dumps(nodes or [])
    condition_clean = (condition or "Pathway").strip()
    care_setting_clean = (care_setting or "").strip()
    if care_setting_clean:
        pathway_title = f"Pathway: Managing {condition_clean} in {care_setting_clean}"
        page_title = f"Beta Testing Guide: {condition_clean} ({care_setting_clean})"
    else:
        pathway_title = f"Pathway: Managing {condition_clean}"
        page_title = f"Beta Testing Guide: {condition_clean}"
    
    # Pathway visualization removed
    pathway_button_html = ""
    pathway_script = ""
    
    scenarios = build_beta_test_scenarios(condition_clean, nodes or [], care_setting_clean, genai_client)
    scenario_blocks = []
    for idx, scenario in enumerate(scenarios, start=1):
        slug = re.sub(r'[^a-z0-9]+', '-', scenario.get("title", f"scenario-{idx}").lower()).strip('-') or f"scenario-{idx}"
        tasks_html = "\n".join([f"<li>{task}</li>" for task in scenario.get("tasks", [])])
        scenario_blocks.append(f"""
<div class=\"scenario-card\">
<h3>Scenario {idx}: {scenario.get('title', 'Scenario')}</h3>
<p><strong>Vignette:</strong> {scenario.get('vignette', '')}</p>
<ul class=\"tasks\">
{tasks_html}
</ul>
<div class=\"checklist\">
<strong>{scenario.get('success_criteria', 'Did the pathway land on the correct end-node?')}</strong>
<label><input type=\"checkbox\" class=\"scenario-check\" data-scenario=\"{slug}\" data-label=\"{scenario.get('title', 'Scenario')}\" data-notes-id=\"scenario{idx}_notes\"> ✓ Reached the intended outcome</label>
</div>
<label style=\"margin-top:10px\">Notes:</label>
<textarea id=\"scenario{idx}_notes\" placeholder=\"{scenario.get('notes_placeholder', 'Document issues or blockers...')}\"></textarea>
</div>
""")

    scenario_blocks_html = "\n".join(scenario_blocks)

    html = f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<title>{page_title}</title>
<style>
{SHARED_CSS}
.info-grid {{display:grid;grid-template-columns:1fr 1fr 1fr;gap:12px;margin-bottom:20px}}
@media (max-width: 768px) {{
    .info-grid {{grid-template-columns:1fr}}
}}
.scenario-card {{background:#fff;border:2px solid var(--border-gray);border-radius:8px;padding:20px;margin-bottom:20px}}
.scenario-card h3 {{color:var(--brown-dark);margin:0 0 10px 0}}
.scenario-card p {{margin:8px 0;color:#555}}
.scenario-card .tasks {{margin:12px 0;padding-left:20px}}
.scenario-card .tasks li {{margin:6px 0}}
.checklist {{margin-top:12px}}
.checklist label {{display:flex;align-items:center;margin:8px 0;cursor:pointer}}
.checklist input[type="checkbox"] {{width:20px;height:20px;margin-right:10px;cursor:pointer}}
.scenario-card textarea {{width:100%;min-height:80px;margin-top:8px;padding:10px;border:1px solid var(--border-gray);border-radius:4px;resize:vertical}}
table {{width:100%;border-collapse:collapse;margin:15px 0}}
th,td {{border:1px solid var(--border-gray);padding:12px;text-align:left}}
th {{background:#f8f9fa;font-weight:600;color:var(--brown-dark)}}
select,input[type="text"],input[type="email"],textarea {{width:100%;padding:10px;border:1px solid var(--border-gray);border-radius:4px}}
textarea {{min-height:70px;resize:vertical}}
label {{display:block;margin-bottom:6px;font-weight:500;color:var(--brown-dark)}}
.form-row {{display:grid;grid-template-columns:1fr 1fr;gap:16px;margin-bottom:16px}}
.rating-scale {{display:flex;gap:8px;align-items:center}}
.rating-scale input[type="radio"] {{width:auto;margin:0}}
.rating-scale label {{margin:0;font-weight:normal;cursor:pointer}}
.heuristics-grid {{display:grid;grid-template-columns:repeat(auto-fit,minmax(190px,1fr));gap:12px;margin-top:12px}}
.heuristic-card {{border:1px solid var(--border-gray);border-radius:8px;padding:12px;background:#fff;display:flex;flex-direction:column;gap:8px}}
.heuristic-head {{display:flex;justify-content:space-between;align-items:center;gap:8px}}
.heuristic-title {{font-weight:600;color:var(--brown-dark);font-size:0.95em;line-height:1.3}}
.heuristic-info {{width:18px;height:18px;display:inline-flex;align-items:center;justify-content:center;border-radius:50%;border:1px solid var(--border-gray);color:#555;font-size:0.8em;cursor:help;position:relative}}
.heuristic-info:hover::after {{content:attr(data-desc);position:absolute;top:28px;left:50%;transform:translateX(-50%);background:#fff;border:1px solid #ccc;padding:8px 10px;border-radius:6px;color:#333;white-space:normal;max-width:280px;box-shadow:0 2px 8px rgba(0,0,0,0.1);z-index:10;font-size:0.85em;line-height:1.4;font-weight:normal}}
.likert {{display:flex;gap:6px;align-items:center;flex-wrap:wrap;font-size:0.9em;color:#444}}
.likert label {{margin:0;display:inline-flex;align-items:center;gap:4px;font-weight:500;color:#444}}
.heuristic-card textarea {{min-height:50px;font-size:0.9em}}
</style>
{pathway_script}
</head>
<body>
<div class="container">
<div class="header">
<h1>Beta Testing Guide</h1>
<p style="margin-top:8px;font-size:1.1em">{pathway_title}</p>
<p style="font-size:0.9em;margin-top:5px">Organization: {organization} • Scenarios auto-generated with LLM context from the current pathway.</p>
{pathway_button_html}
</div>

<form id="betaForm">
<!-- Tester Info -->
<div class="info-grid">
<div class="form-group">
<label for="tester_name">Your Name *</label>
<input type="text" id="tester_name" required placeholder="Full name">
</div>
<div class="form-group">
<label for="tester_email">Email *</label>
<input type="email" id="tester_email" required placeholder="email@example.com">
</div>
<div class="form-group">
<label for="tester_role">Role *</label>
<input type="text" id="tester_role" required placeholder="e.g., RN, Physician, APP">
</div>
</div>

<hr style="margin:30px 0;border:none;border-top:2px solid var(--border-gray)">

<!-- Clinical Scenarios -->
<h2 style="color:var(--brown-dark);margin-bottom:20px">Clinical Scenarios — End-to-End Testing</h2>
<p style="margin-bottom:20px;color:#555">Use the data provided to follow the pathway to the correct terminal node (diagnosis → treatment → disposition).</p>
{scenario_blocks_html}

<hr style="margin:30px 0;border:none;border-top:2px solid var(--border-gray)">

<!-- Nielsen Heuristics -->
<h2 style="color:var(--brown-dark);margin-bottom:20px">Nielsen's Usability Heuristics</h2>
<p style="margin-bottom:20px;color:#555">Rate each heuristic from 1 (Poor) to 5 (Excellent) and provide comments.</p>

<div id="heuristicsGrid" class="heuristics-grid"></div>

<hr style="margin:30px 0;border:none;border-top:2px solid var(--border-gray)">

<!-- Overall Feedback -->
<h2 style="color:var(--brown-dark);margin-bottom:20px">Overall Feedback</h2>
<div class="form-row">
<div class="form-group">
<label for="overall_rating">Overall Pathway Quality (1–5)</label>
<select id="overall_rating">
<option value="1">1 - Poor</option>
<option value="2">2 - Fair</option>
<option value="3" selected>3 - Good</option>
<option value="4">4 - Very Good</option>
<option value="5">5 - Excellent</option>
</select>
</div>
<div class="form-group">
<label for="workflow_fit">Fits Clinical Workflow (1–5)</label>
<select id="workflow_fit">
<option value="1">1 - Poor Fit</option>
<option value="2">2 - Fair Fit</option>
<option value="3" selected>3 - Good Fit</option>
<option value="4">4 - Very Good Fit</option>
<option value="5">5 - Excellent Fit</option>
</select>
</div>
</div>

<div class="form-group">
<label for="strengths">What worked well? (Strengths)</label>
<textarea id="strengths" placeholder="Describe positive aspects, helpful features, clear sections..."></textarea>
</div>

<div class="form-group">
<label for="improvements">What needs improvement? (Issues & Suggestions)</label>
<textarea id="improvements" placeholder="Describe problems encountered, confusing areas, missing features..."></textarea>
</div>

<div class="button-group" style="margin-top:30px;gap:10px;align-items:center;flex-wrap:wrap">
    <button type="button" onclick="downloadCSV()" style="background:var(--brown);color:white;font-size:1.05em;padding:14px 28px">Download Responses (CSV)</button>
    <button type="reset" style="background:#999;color:white">Reset Form</button>
</div>
</form>

{CAREPATHIQ_FOOTER}
</div>

<script>
const HEURISTICS = [
  {{id: "h1", name: "Visibility of System Status", desc: "Does the pathway clearly show where you are and what's happening?"}},
  {{id: "h2", name: "Match Between System and Real World", desc: "Does it use familiar clinical language and concepts?"}},
  {{id: "h3", name: "User Control and Freedom", desc: "Can you easily undo mistakes or go back to previous steps?"}},
  {{id: "h4", name: "Consistency and Standards", desc: "Are terms, layouts, and actions consistent throughout?"}},
  {{id: "h5", name: "Error Prevention", desc: "Does it prevent errors before they happen (not just detect)?"}},
  {{id: "h6", name: "Recognition Rather Than Recall", desc: "Are options visible rather than requiring memorization?"}},
  {{id: "h7", name: "Flexibility and Efficiency", desc: "Does it accommodate both novice and expert users?"}},
  {{id: "h8", name: "Aesthetic and Minimalist Design", desc: "Is the interface clean without unnecessary information?"}},
  {{id: "h9", name: "Help Users Recognize and Recover from Errors", desc: "Are error messages clear with suggested solutions?"}},
  {{id: "h10", name: "Help and Documentation", desc: "Is help available when needed and easy to understand?"}}
];

const condition = "{condition}";
const nodes = {nodes_json};

// Initialize heuristics form (compact grid)
function initializeHeuristics() {{
    const container = document.getElementById('heuristicsGrid');
    container.innerHTML = HEURISTICS.map(h => `
        <div class="heuristic-card">
            <div class="heuristic-head">
                <span class="heuristic-title">${{h.name}}</span>
                <span class="heuristic-info" data-desc="${{h.desc}}">?</span>
            </div>
            <div class="likert" role="radiogroup" aria-label="${{h.name}} rating">
                <label><input type="radio" name="${{h.id}}_rating" value="1" required>1</label>
                <label><input type="radio" name="${{h.id}}_rating" value="2">2</label>
                <label><input type="radio" name="${{h.id}}_rating" value="3" checked>3</label>
                <label><input type="radio" name="${{h.id}}_rating" value="4">4</label>
                <label><input type="radio" name="${{h.id}}_rating" value="5">5</label>
            </div>
            <textarea id="${{h.id}}_comments" placeholder="Comment (optional)"></textarea>
        </div>
    `).join('');
}}

// Download CSV
function downloadCSV() {{
  const name = document.getElementById('tester_name').value;
  const email = document.getElementById('tester_email').value;
  const role = document.getElementById('tester_role').value;
  
  if (!name || !email || !role) {{
    alert('Please fill in your name, email, and role before downloading.');
    return;
  }}
  
  let csv = 'Category,Item,Value\\n';
  csv += `Tester Name,${{name}},"${{name}}"\\n`;
  csv += `Tester Email,${{email}},"${{email}}"\\n`;
  csv += `Tester Role,${{role}},"${{role}}"\\n`;
  csv += `Condition,"{condition}","{condition}"\\n`;
  csv += `Test Date,"${{new Date().toLocaleDateString()}}","${{new Date().toISOString()}}"\\n`;
  csv += '\\n';
  
  // Scenarios
  csv += 'Scenario Testing,Item,Notes\\n';
    const s1Check = document.querySelector('input[data-scenario="lowrisk"]').checked ? 'Completed' : 'Not Completed';
{scenario_blocks_html}
</body>
</html>
"""
    return html


def generate_education_module_html(
    condition: str,
    nodes: list = None,
    target_audience: str = "",
    care_setting: str = "",
    genai_client=None
) -> str:
    """
    Generate comprehensive single-page education module with AI-generated content,
    5 MC quiz questions, and CSV export.
    
    Args:
        condition: Clinical condition
        nodes: Pathway nodes for context
        target_audience: Who this education is for
        care_setting: Setting where pathway is used
        genai_client: Google Generative AI client (required)
        
    Returns:
        Complete standalone HTML string
        
    Raises:
        ValueError: If genai_client is not provided or nodes are empty
    """
    
    if not genai_client:
        raise ValueError("AI client is required to generate education module content")
    
    if not nodes or len(nodes) == 0:
        raise ValueError("Pathway nodes are required to generate education module")
    
    condition_clean = (condition or "Clinical Pathway").strip()
    care_setting_clean = (care_setting or "healthcare setting").strip()
    audience_clean = (target_audience or "clinical staff").strip()
    
    # Extract key pathway information for AI context - include more detail
    decision_nodes = []
    process_nodes = []
    all_node_labels = []
    for n in nodes:
        node_type = n.get('type', '')
        label = n.get('label', '')
        detail = n.get('detail', '')
        all_node_labels.append(f"{node_type}: {label}")
        if node_type == 'Decision':
            decision_nodes.append(label)
        elif node_type == 'Process':
            process_nodes.append(label)
    
    # Provide rich pathway context for AI
    nodes_context = f"""
FULL PATHWAY STRUCTURE ({len(nodes)} nodes):
{chr(10).join(all_node_labels[:15])}
{'... and ' + str(len(nodes) - 15) + ' more nodes' if len(nodes) > 15 else ''}

KEY DECISION POINTS ({len(decision_nodes)}):
{chr(10).join(['- ' + d for d in decision_nodes[:10]])}

KEY CLINICAL ACTIONS ({len(process_nodes)}):
{chr(10).join(['- ' + p for p in process_nodes[:12]])}
"""
    
    # Generate learning objectives using AI - AUDIENCE-SPECIFIC
    objectives_prompt = f"""You are a clinical educator creating learning objectives for **{audience_clean}** about {condition_clean} management in {care_setting_clean}.

TARGET AUDIENCE PROFILE - {audience_clean}:
- Tailor complexity and terminology to this specific role
- If residents/trainees: focus on clinical reasoning and decision-making skills
- If nursing staff: emphasize assessment, recognition, escalation, and care coordination
- If attendings/physicians: focus on evidence synthesis, risk stratification, and management decisions
- If pharmacists: emphasize medication selection, dosing, interactions, monitoring
- If allied health: focus on their specific scope and handoff communication

PATHWAY CONTEXT:
{nodes_context}

Generate EXACTLY 4 specific, measurable learning objectives using Bloom's taxonomy action verbs.

FORMAT: Return ONLY a JSON array of 4 strings.

REQUIREMENTS:
1. Each objective MUST be specific to what {audience_clean} needs to know and do
2. Use clinical terminology from the actual pathway nodes above
3. Include measurable outcomes (identify, apply, evaluate, demonstrate, calculate, differentiate, recognize, prioritize)
4. Reference SPECIFIC decision points, scores, thresholds, or interventions from the pathway
5. Write at the appropriate level for {audience_clean}

EXAMPLE for Emergency Medicine Residents learning Chest Pain pathway:
["Apply the HEART score criteria to risk-stratify patients with undifferentiated chest pain and determine appropriate disposition", "Differentiate between STEMI, NSTEMI, and unstable angina presentations using ECG findings and troponin trends", "Evaluate contraindications to anticoagulation and select appropriate agents based on patient-specific factors", "Demonstrate appropriate escalation to cardiology consultation based on pathway-defined high-risk criteria"]"""

    obj_response = genai_client.models.generate_content(
        model="gemini-2.0-flash",
        contents=objectives_prompt
    )
    json_match = re.search(r'\[.*\]', obj_response.text, re.DOTALL)
    if json_match:
        learning_objectives = json.loads(json_match.group())[:4]
    else:
        learning_objectives = [
            f"Apply evidence-based assessment criteria for {condition_clean}",
            f"Utilize pathway decision points to guide clinical management in {care_setting_clean}",
            "Recognize escalation triggers and appropriate response actions",
            "Integrate pathway protocols into clinical workflow"
        ]
    
    # Generate detailed teaching points using AI - AUDIENCE-SPECIFIC
    teaching_prompt = f"""You are a clinical educator creating KEY TEACHING CONTENT for **{audience_clean}** about {condition_clean} management in {care_setting_clean}.

TARGET AUDIENCE: {audience_clean}
- Write at the appropriate knowledge level for this role
- Emphasize what THIS role specifically needs to know and do
- Use terminology and examples relevant to their daily workflow

PATHWAY CONTEXT:
{nodes_context}

Create EXACTLY 5 detailed teaching points as complete paragraphs (3-4 sentences each).

STRUCTURE:
1. **Initial Recognition & Assessment**: Red flags, vital sign thresholds, history elements specific to this pathway
2. **Risk Stratification**: Specific scoring tools, cutoffs, and how to apply them
3. **Diagnostic Workup**: What tests to order, in what sequence, and interpretation pearls
4. **Treatment Protocols**: Specific medications with doses, timing, contraindications from the pathway
5. **Disposition & Follow-up**: Admission criteria, discharge requirements, follow-up timing

REQUIREMENTS:
- Extract SPECIFIC clinical details from the pathway nodes (scores, thresholds, drug names, doses)
- Each paragraph must teach something actionable for {audience_clean}
- Include clinical pearls and common pitfalls relevant to this pathway
- Be detailed enough that someone could apply this knowledge clinically

Return ONLY a JSON array of 5 strings (each string is one teaching paragraph)."""

    teach_response = genai_client.models.generate_content(
        model="gemini-2.0-flash",
        contents=teaching_prompt
    )
    json_match = re.search(r'\[.*\]', teach_response.text, re.DOTALL)
    if json_match:
        teaching_points = json.loads(json_match.group())
    else:
        teaching_points = [f"This pathway guides {condition_clean} management for {audience_clean} in {care_setting_clean}, covering assessment, treatment, and disposition decisions."]
    
    # Generate clinically accurate quiz questions using AI - DETAILED EXPLANATIONS
    quiz_prompt = f"""You are a clinical educator creating ASSESSMENT QUESTIONS for **{audience_clean}** learning about {condition_clean} pathway in {care_setting_clean}.

TARGET AUDIENCE: {audience_clean}
- Questions should test knowledge relevant to THIS role's responsibilities
- Scenarios should reflect situations {audience_clean} would actually encounter

PATHWAY CONTEXT:
{nodes_context}

Generate EXACTLY 5 multiple choice questions testing understanding of this pathway.

QUESTION REQUIREMENTS:
- Each question presents a realistic clinical scenario with specific patient details (age, vitals, presentation)
- Scenarios must reference actual decision points, scores, or thresholds from the pathway
- All 4 options must be clinically plausible (no obviously wrong answers)
- Questions progress from basic assessment to complex management decisions

EXPLANATION REQUIREMENTS (CRITICAL - make explanations DETAILED and EDUCATIONAL):
- Explain WHY the correct answer is right with specific clinical reasoning
- Reference the specific pathway criteria, score cutoffs, or evidence that supports this answer
- Briefly explain why each wrong answer is incorrect or suboptimal
- Include a clinical pearl or teaching point
- Explanations should be 3-5 sentences, not brief

Return ONLY a valid JSON array:
[
  {{
    "question": "A [age/gender] patient presents to {care_setting_clean} with [specific scenario]. Vital signs show [specifics]. Based on the {condition_clean} pathway, what is the most appropriate next step?",
    "options": ["A) Specific clinical action", "B) Another plausible action", "C) Third reasonable option", "D) Fourth consideration"],
    "correct": "B",
    "explanation": "According to the pathway, [detailed reasoning with specific criteria]. The patient meets criteria for [X] because [specific factors]. Option A is incorrect because [reason]. Option C would be appropriate if [different scenario]. Option D is not indicated because [reason]. Clinical pearl: [actionable teaching point]."
  }}
]

QUESTION TOPICS (one each):
Q1: Initial assessment - recognizing key findings that trigger pathway entry
Q2: Risk stratification - applying a specific score or criteria from the pathway
Q3: Diagnostic approach - selecting appropriate tests based on risk level
Q4: Treatment decision - choosing intervention based on pathway criteria
Q5: Disposition - determining appropriate level of care or follow-up"""

    quiz_response = genai_client.models.generate_content(
        model="gemini-2.0-flash",
        contents=quiz_prompt
    )
    json_match = re.search(r'\[.*\]', quiz_response.text, re.DOTALL)
    if json_match:
        questions = json.loads(json_match.group())[:5]
    else:
        raise ValueError("Failed to generate quiz questions. Please try again.")
    
    # Build teaching points HTML
    if isinstance(teaching_points, list):
        teaching_html = "".join([f"<p style='margin-bottom: 15px; line-height: 1.7;'>{point}</p>" for point in teaching_points])
    else:
        teaching_html = f"<p style='line-height: 1.7;'>{teaching_points}</p>"
    
    # Build learning objectives HTML
    objectives_html = "".join([f"<li>{obj}</li>" for obj in learning_objectives])
    
    # Pre-compute JSON for JavaScript embedding
    questions_json = json.dumps(questions)
    
    # Professional title formatting
    module_title = f"{condition_clean} Clinical Pathway Education Module"
    subtitle = f"Interactive Pathway-Based Learning for {audience_clean}"
    
    # Build HTML
    html = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{module_title}</title>
    <style>
        * {{ margin: 0; padding: 0; box-sizing: border-box; }}
        body {{
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            background: linear-gradient(135deg, #98d8c8 0%, #f7f7f7 100%);
            padding: 20px;
            min-height: 100vh;
        }}
        .container {{
            max-width: 900px;
            margin: 0 auto;
            background: white;
            border-radius: 8px;
            box-shadow: 0 2px 10px rgba(0,0,0,0.1);
            overflow: hidden;
        }}
        .header {{
            background: linear-gradient(135deg, #5D4037 0%, #6D4C41 100%);
            color: white;
            padding: 40px 20px;
            text-align: center;
        }}
        .header h1 {{
            font-size: 26px;
            margin-bottom: 10px;
            font-weight: 600;
            letter-spacing: 0.5px;
        }}
        .header p {{
            font-size: 15px;
            opacity: 0.9;
            font-style: italic;
        }}
        .header .setting-badge {{
            display: inline-block;
            background: rgba(255,255,255,0.2);
            padding: 6px 14px;
            border-radius: 20px;
            font-size: 13px;
            margin-top: 12px;
        }}
        .content {{
            padding: 40px;
        }}
        .section {{
            margin-bottom: 40px;
        }}
        .section h2 {{
            color: #5D4037;
            margin-bottom: 15px;
            font-size: 20px;
            border-bottom: 2px solid #98d8c8;
            padding-bottom: 10px;
        }}
        .learning-points {{
            background: #f9f9f9;
            padding: 20px;
            border-radius: 6px;
            border-left: 4px solid #98d8c8;
            margin-bottom: 20px;
        }}
        .learning-points ul {{
            list-style: none;
            padding-left: 20px;
        }}
        .learning-points li {{
            margin: 12px 0;
            line-height: 1.6;
        }}
        .learning-points li:before {{
            content: "→ ";
            color: #98d8c8;
            font-weight: bold;
            margin-right: 10px;
        }}
        .teaching-content {{
            background: #fafafa;
            padding: 25px;
            border-radius: 6px;
            border-left: 4px solid #5D4037;
        }}
        .teaching-content p {{
            color: #333;
            font-size: 15px;
        }}
        .quiz-question {{
            background: #fafafa;
            padding: 20px;
            margin: 20px 0;
            border-radius: 6px;
            border: 1px solid #e0e0e0;
        }}
        .quiz-question h3 {{
            color: #333;
            margin-bottom: 15px;
            font-size: 15px;
            line-height: 1.5;
        }}
        .option {{
            margin: 10px 0;
            display: flex;
            align-items: flex-start;
        }}
        .option input[type="radio"] {{
            margin-top: 4px;
            margin-right: 10px;
            cursor: pointer;
        }}
        .option label {{
            cursor: pointer;
            flex: 1;
            line-height: 1.5;
        }}
        .feedback {{
            margin-top: 15px;
            padding: 12px;
            border-radius: 4px;
            display: none;
            font-size: 14px;
            line-height: 1.5;
        }}
        .feedback.show {{
            display: block;
        }}
        .feedback.correct {{
            background: #d4edda;
            color: #155724;
            border: 1px solid #c3e6cb;
        }}
        .feedback.incorrect {{
            background: #f8d7da;
            color: #721c24;
            border: 1px solid #f5c6cb;
        }}
        .button-row {{
            display: flex;
            gap: 10px;
            margin-top: 20px;
            justify-content: center;
        }}
        button {{
            padding: 12px 24px;
            font-size: 14px;
            border: none;
            border-radius: 4px;
            cursor: pointer;
            transition: all 0.3s ease;
        }}
        .btn-submit {{
            background: #5D4037;
            color: white;
        }}
        .btn-submit:hover {{
            background: #4E342E;
        }}
        .btn-download {{
            background: #98d8c8;
            color: #333;
        }}
        .btn-download:hover {{
            background: #7ec9b6;
        }}
        .results {{
            background: #f0f8f5;
            padding: 20px;
            border-radius: 6px;
            margin-top: 30px;
            text-align: center;
            display: none;
        }}
        .results.show {{
            display: block;
        }}
        .results h3 {{
            color: #5D4037;
            margin-bottom: 10px;
        }}
        .score {{
            font-size: 32px;
            color: #98d8c8;
            font-weight: bold;
            margin: 10px 0;
        }}
        .footer {{
            background: #f5f5f5;
            padding: 20px;
            text-align: center;
            color: #666;
            font-size: 12px;
            border-top: 1px solid #e0e0e0;
        }}
    </style>
</head>
<body>
    <div class="container">
        <div class="header">
            <h1>{module_title}</h1>
            <p>{subtitle}</p>
            <span class="setting-badge">{care_setting_clean}</span>
        </div>
        
        <div class="content">
            <div class="section">
                <h2>Learning Objectives</h2>
                <div class="learning-points">
                    <ul>
                        {objectives_html}
                    </ul>
                </div>
            </div>
            
            <div class="section">
                <h2>Key Teaching Points</h2>
                <div class="teaching-content">
                    {teaching_html}
                </div>
            </div>
            
            <div class="section">
                <h2>Knowledge Assessment — 5 Questions</h2>
                <p style="color: #666; margin-bottom: 20px; font-size: 14px;">Answer all questions to complete this module. A score of 100% is required to receive your certificate.</p>
                <form id="quizForm">
"""
    
    # Add quiz questions
    for idx, q in enumerate(questions):
        correct_letter = q.get('correct', 'A')
        html += f"""
                    <div class="quiz-question">
                        <h3>Question {idx + 1}: {q.get('question', '')}</h3>
"""
        for i, option in enumerate(q.get('options', [])):
            option_letter = chr(65 + i)  # A, B, C, D
            html += f"""
                        <div class="option">
                            <input type="radio" id="q{idx}_opt{option_letter}" name="q{idx}" value="{option_letter}" required>
                            <label for="q{idx}_opt{option_letter}">{option}</label>
                        </div>
"""
        html += f"""
                        <div class="feedback" id="feedback{idx}"></div>
                    </div>
"""
    
    html += f"""
                <div class="button-row">
                    <button class="btn-submit" onclick="submitQuiz()">Submit Answers</button>
                </div>
                
                <div class="results" id="results">
                    <h3>Your Score</h3>
                    <div class="score" id="scoreDisplay">0%</div>
                    <p id="scoreMessage"></p>
                    <div id="perfectScoreSection" style="display:none; margin-top: 20px;">
                        <p style="color: #155724; font-weight: bold; margin-bottom: 15px;">🎉 Perfect Score! You have completed this education module successfully.</p>
                        <div style="margin: 15px 0;">
                            <label for="certName">Name for Certificate:</label><br>
                            <input type="text" id="certName" placeholder="Enter your full name" style="padding: 10px; width: 300px; max-width: 100%; margin: 10px 0; border: 1px solid #ddd; border-radius: 4px;">
                        </div>
                        <button class="btn-download" onclick="downloadCertificate()">Download Certificate (PNG)</button>
                    </div>
                </div>
            </div>
        </div>
        
        <div class="footer">
            <p>© 2024 CarePathIQ. All rights reserved. Licensed under CC BY-SA 4.0</p>
            <p>{condition_clean} | {care_setting_clean}</p>
        </div>
    </div>
    
    <script>
        const quizData = {questions_json};
        const pathwayName = "{condition}";
        const caregSetting = "{care_setting}";
        let userResponses = {{}};
        
        function submitQuiz() {{
            const form = document.getElementById('quizForm');
            const formData = new FormData(form);
            
            let correct = 0;
            let total = quizData.length;
            
            quizData.forEach((q, idx) => {{
                const userAnswer = formData.get('q' + idx);
                userResponses['q' + idx] = userAnswer;
                
                const feedbackEl = document.getElementById('feedback' + idx);
                if (userAnswer === q.correct) {{
                    correct++;
                    feedbackEl.className = 'feedback show correct';
                    feedbackEl.innerHTML = '<strong>✓ Correct!</strong> ' + q.explanation;
                }} else {{
                    feedbackEl.className = 'feedback show incorrect';
                    feedbackEl.innerHTML = '<strong>✗ Incorrect.</strong> Correct answer: ' + q.correct + '. ' + q.explanation;
                }}
            }});
            
            const percentage = Math.round((correct / total) * 100);
            document.getElementById('scoreDisplay').textContent = percentage + '%';
            document.getElementById('scoreMessage').textContent = 'You answered ' + correct + ' out of ' + total + ' questions correctly.';
            
            if (percentage === 100) {{
                document.getElementById('perfectScoreSection').style.display = 'block';
            }}
            
            document.getElementById('results').classList.add('show');
            
            window.scrollTo(0, document.getElementById('results').offsetTop - 100);
        }}
        
        function downloadCertificate() {{
            const name = document.getElementById('certName').value.trim();
            if (!name) {{
                alert('Please enter your name on the certificate.');
                return;
            }}
            
            // Create canvas for certificate
            const canvas = document.createElement('canvas');
            canvas.width = 1200;
            canvas.height = 800;
            const ctx = canvas.getContext('2d');
            
            // Background
            ctx.fillStyle = '#f5f5dc';
            ctx.fillRect(0, 0, canvas.width, canvas.height);
            
            // Border
            ctx.strokeStyle = '#5D4037';
            ctx.lineWidth = 8;
            ctx.strokeRect(20, 20, canvas.width - 40, canvas.height - 40);
            
            // Inner border
            ctx.strokeStyle = '#98d8c8';
            ctx.lineWidth = 3;
            ctx.strokeRect(40, 40, canvas.width - 80, canvas.height - 80);
            
            // Title
            ctx.fillStyle = '#5D4037';
            ctx.font = 'bold 48px Georgia, serif';
            ctx.textAlign = 'center';
            ctx.fillText('Certificate of Completion', canvas.width / 2, 120);
            
            // Logo placeholder (CarePathIQ text)
            ctx.font = '24px Arial, sans-serif';
            ctx.fillStyle = '#98d8c8';
            ctx.fillText('CarePathIQ', 80, 100);
            
            // Pathway name
            ctx.fillStyle = '#333';
            ctx.font = '32px Georgia, serif';
            ctx.fillText(pathwayName + ' Pathway', canvas.width / 2, 280);
            
            // Message
            ctx.font = '18px Arial';
            ctx.fillStyle = '#666';
            ctx.textAlign = 'center';
            ctx.fillText('This certifies that', canvas.width / 2, 380);
            
            // Name
            ctx.fillStyle = '#5D4037';
            ctx.font = 'bold 36px Georgia, serif';
            ctx.fillText(name, canvas.width / 2, 480);
            
            // Bottom message
            ctx.font = '16px Arial';
            ctx.fillStyle = '#666';
            ctx.fillText('has successfully completed the education module for', canvas.width / 2, 560);
            ctx.fillText(pathwayName + ' in ' + caregSetting, canvas.width / 2, 600);
            
            // Date
            const today = new Date();
            const dateStr = today.toLocaleDateString('en-US', {{ year: 'numeric', month: 'long', day: 'numeric' }});
            ctx.font = '14px Arial';
            ctx.fillStyle = '#999';
            ctx.fillText('Date: ' + dateStr, canvas.width / 2, 680);
            
            // Download
            canvas.toBlob(function(blob) {{
                const url = window.URL.createObjectURL(blob);
                const a = document.createElement('a');
                a.href = url;
                a.download = name.replace(/\\s+/g, '_') + '_Certificate.png';
                document.body.appendChild(a);
                a.click();
                window.URL.revokeObjectURL(url);
                document.body.removeChild(a);
            }});
        }}
    </script>
</body>
</html>"""
    
    return html

    
    html = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Education Module: {condition}</title>
    <style>
        {SHARED_CSS}
        
        .education-nav {{
            display: flex;
            gap: 10px;
            margin: 20px 0;
            flex-wrap: wrap;
        }}
        
        .module-button {{
            padding: 10px 15px;
            border: 2px solid var(--border-gray);
            background: white;
            cursor: pointer;
            border-radius: 4px;
            transition: all 0.3s ease;
            font-weight: 600;
        }}
        
        .module-button.active {{
            background-color: var(--teal);
            border-color: var(--teal);
            color: var(--brown-dark);
        }}
        
        .module-button:hover {{
            border-color: var(--teal);
        }}
        
        .module-content {{
            display: none;
            animation: fadeIn 0.3s ease;
        }}
        
        .module-content.active {{
            display: block;
        }}
        
        @keyframes fadeIn {{
            from {{ opacity: 0; }}
            to {{ opacity: 1; }}
        }}
        
        .quiz {{
            background-color: var(--light-gray);
            padding: 20px;
            border-radius: 6px;
            margin-top: 20px;
        }}
        
        .quiz-question {{
            margin-bottom: 15px;
        }}
        
        .quiz-option {{
            display: flex;
            align-items: center;
            margin: 10px 0;
            cursor: pointer;
        }}
        
        .quiz-option input[type="radio"] {{
            width: auto;
            margin-right: 10px;
        }}
        
        .quiz-option label {{
            margin: 0;
            cursor: pointer;
            flex: 1;
        }}
        
        .quiz-feedback {{
            margin-top: 15px;
            padding: 15px;
            border-radius: 4px;
            display: none;
        }}
        
        .quiz-feedback.show {{
            display: block;
        }}
        
        .quiz-feedback.correct {{
            background-color: #d4edda;
            color: #155724;
            border: 1px solid #c3e6cb;
        }}
        
        .quiz-feedback.incorrect {{
            background-color: #f8d7da;
            color: #721c24;
            border: 1px solid #f5c6cb;
        }}
        
        .progress-container {{
            margin: 20px 0;
        }}
        
        .progress-label {{
            display: flex;
            justify-content: space-between;
            margin-bottom: 8px;
            font-weight: 600;
        }}
        
        .certificate {{
            display: none;
            background: linear-gradient(135deg, var(--teal) 0%, white 100%);
            padding: 40px;
            border: 6px solid var(--brown);
            border-radius: 10px;
            text-align: center;
            margin: 30px 0;
            page-break-after: always;
        }}
        
        .certificate.show {{
            display: block;
        }}
        
        .certificate-content {{
            background: white;
            padding: 40px;
            border: 2px dashed var(--brown);
            border-radius: 6px;
        }}
        
        .certificate-title {{
            font-size: 2em;
            color: var(--brown-dark);
            margin-bottom: 20px;
            font-weight: bold;
        }}
        
        .certificate-text {{
            font-size: 1.1em;
            color: #333;
            margin: 15px 0;
            line-height: 1.6;
        }}
        
        .certificate-condition {{
            font-size: 1.5em;
            color: var(--brown);
            font-weight: bold;
            margin: 20px 0;
        }}
        
        .certificate-date {{
            margin-top: 30px;
            color: #666;
            font-size: 0.95em;
        }}
        
        .certificate-logo {{
            margin: 20px 0;
            color: var(--brown);
            font-weight: bold;
            font-size: 1.3em;
        }}
        
        @media print {{
            body {{
                background: white;
            }}
            
            .container {{
                box-shadow: none;
                padding: 0;
            }}
            
            .header {{
                page-break-after: always;
            }}
            
            .education-nav,
            .module-button,
            button:not(.print-hidden),
            .progress-container {{
                display: none;
            }}
        }}
        
        @media (max-width: 768px) {{
            .education-nav {{
                flex-direction: column;
            }}
            
            .module-button {{
                width: 100%;
            }}
        }}
    </style>
</head>
<body>
    <div class="container">
        <div class="header">
            <h1>Education Module</h1>
            <p><strong>{condition}</strong></p>
            <p style="font-size: 0.95em; margin-top: 10px;">Interactive Learning & Certification</p>
        </div>

        <div class="progress-container">
            <div class="progress-label">
                <span>Course Progress</span>
                <span id="progressText">0%</span>
            </div>
            <div class="progress-bar">
                <div class="progress-fill" id="progressFill" style="width: 0%"></div>
            </div>
            <p id="moduleStatus" style="text-align: center; color: #666; font-size: 0.9em; margin-top: 5px;">
                Completed 0 of {total_sections} modules
            </p>
        </div>

        <div class="education-nav" id="navContainer"></div>

        <div id="contentContainer"></div>

        <div class="certificate" id="certificateSection">
            <div class="certificate-content">
                <div class="certificate-logo">★ CarePathIQ ★</div>
                <div class="certificate-title">Certificate of Completion</div>
                
                <div class="certificate-text">
                    This certifies that
                </div>
                
                <input type="text" id="certificateName" placeholder="Your Name" style="font-size: 1.3em; text-align: center; border: none; border-bottom: 2px solid var(--brown); width: 100%; margin: 15px 0; padding: 10px 0;">
                
                <div class="certificate-text">
                    has successfully completed the
                </div>
                
                <div class="certificate-condition">{condition}</div>
                
                <div class="certificate-text">
                    Education Module offered by
                </div>
                
                <div style="font-weight: bold; color: var(--brown); margin: 15px 0;">{organization}</div>
                
                <div class="certificate-date">
                    <p style="margin: 10px 0;">Date Completed: <strong id="completionDate"></strong></p>
                    <p style="margin: 10px 0;">This certificate is valid as a record of learning completion.</p>
                    <p style="margin-top: 20px; font-size: 0.85em; color: #999;">
                        Certificate ID: <span id="certId"></span>
                    </p>
                </div>
            </div>
        </div>

        <div style="margin-top: 30px; display: flex; gap: 10px; flex-wrap: wrap;">
            <button type="button" onclick="downloadCertificateImage()" id="downloadCertBtn" style="display: none;">Download Certificate (PNG)</button>
            <button type="button" onclick="printCertificate()" id="printCertBtn" style="display: none;">Print Certificate</button>
            <button type="button" onclick="restartCourse()" id="restartBtn" style="display: none;">Restart Course</button>
        </div>

        {CAREPATHIQ_FOOTER}
    </div>

    <script>
        const modules = {modules_json};
        const condition = "{condition}";
        let completedModules = {{}};
        let currentModule = 0;

        document.addEventListener('DOMContentLoaded', function() {{
            renderNav();
            renderModule(0);
            generateCertificateID();
        }});

        function renderNav() {{
            const navContainer = document.getElementById('navContainer');
            modules.forEach((mod, idx) => {{
                const btn = document.createElement('button');
                btn.type = 'button';
                btn.className = 'module-button' + (idx === 0 ? ' active' : '');
                btn.textContent = `Module ${{idx + 1}}`;
                btn.onclick = () => switchModule(idx);
                navContainer.appendChild(btn);
            }});
        }}

        function switchModule(idx) {{
            currentModule = idx;
            const buttons = document.querySelectorAll('.module-button');
            buttons.forEach((btn, i) => {{
                btn.classList.toggle('active', i === idx);
            }});
            renderModule(idx);
        }}

        function renderModule(idx) {{
            const mod = modules[idx];
            const container = document.getElementById('contentContainer');
            
            let html = `
                <h2 style="color: var(--brown-dark); margin-top: 30px;">${{mod.title}}</h2>
                <div style="background: white; padding: 20px; margin: 15px 0; border-radius: 6px;">
                    ${{mod.content}}
                </div>
            `;
            
            if (mod.quiz && mod.quiz.length > 0) {{
                html += '<div class="quiz"><h3>Module Quiz</h3>';
                mod.quiz.forEach((q, qIdx) => {{
                    html += `
                        <div class="quiz-question">
                            <p style="font-weight: bold;">${{q.question}}</p>
                    `;
                    q.options.forEach((opt, optIdx) => {{
                        html += `
                            <div class="quiz-option">
                                <input type="radio" id="q${{idx}}_o${{optIdx}}" name="quiz_${{idx}}" value="${{optIdx}}" onchange="checkAnswer(${{idx}}, ${{optIdx}}, ${{q.correct}})">
                                <label for="q${{idx}}_o${{optIdx}}">${{opt}}</label>
                            </div>
                        `;
                    }});
                    html += `
                            <div id="feedback_${{idx}}_${{qIdx}}" class="quiz-feedback"></div>
                        </div>
                    `;
                }});
                html += '</div>';
            }}
            
            html += `
                <div style="margin-top: 20px; display: flex; gap: 10px;">
                    ${{idx > 0 ? '<button type="button" onclick="switchModule(' + (idx - 1) + ')">← Previous</button>' : ''}}
                    ${{idx < modules.length - 1 ? '<button type="button" onclick="switchModule(' + (idx + 1) + ')">Next →</button>' : ''}}
                </div>
            `;
            
            container.innerHTML = html;
        }}

        function checkAnswer(modIdx, selectedIdx, correctIdx) {{
            const feedback = document.getElementById(`feedback_${{modIdx}}_0`);
            feedback.classList.add('show');
            
            if (selectedIdx === correctIdx) {{
                feedback.classList.add('correct');
                feedback.classList.remove('incorrect');
                feedback.innerHTML = '<strong>✓ Correct!</strong> Great job.';
                completedModules[modIdx] = true;
            }} else {{
                feedback.classList.add('incorrect');
                feedback.classList.remove('correct');
                feedback.innerHTML = '<strong>✗ Incorrect.</strong> Please try again or review the module.';
            }}
            
            updateProgress();
        }}

        function updateProgress() {{
            const completed = Object.keys(completedModules).length;
            const total = modules.length;
            const percent = Math.round((completed / total) * 100);
            
            document.getElementById('progressFill').style.width = percent + '%';
            document.getElementById('progressText').textContent = percent + '%';
            document.getElementById('moduleStatus').textContent = `Completed ${{completed}} of ${{total}} modules`;
            
            if (completed === total) {{
                showCertificate();
            }}
        }}

        function showCertificate() {{
            document.getElementById('certificateSection').classList.add('show');
            document.getElementById('downloadCertBtn').style.display = 'inline-block';
            document.getElementById('printCertBtn').style.display = 'inline-block';
            document.getElementById('restartBtn').style.display = 'inline-block';
            document.getElementById('completionDate').textContent = new Date().toLocaleDateString();
        }}

        function generateCertificateID() {{
            const id = 'CPQ-' + Date.now().toString(36).toUpperCase() + '-' + Math.random().toString(36).substr(2, 9).toUpperCase();
            document.getElementById('certId').textContent = id;
        }}

        function downloadCertificateImage() {{
            const certName = document.getElementById('certificateName').value || 'Learner';
            // Create canvas and convert certificate to image
            const element = document.getElementById('certificateSection');
            
            // Use html2canvas if available, otherwise alert user to use print
            alert('To download as image, please use Print → Save as PDF or use a screenshot tool.\\n\\nAlternatively, click Print Certificate to save as PDF.');
        }}

        function printCertificate() {{
            const certName = document.getElementById('certificateName').value || 'Learner';
            window.print();
        }}

        function restartCourse() {{
            completedModules = {{}};
            currentModule = 0;
            updateProgress();
            switchModule(0);
            document.getElementById('certificateSection').classList.remove('show');
            document.getElementById('downloadCertBtn').style.display = 'none';
            document.getElementById('printCertBtn').style.display = 'none';
            document.getElementById('restartBtn').style.display = 'none';
            window.scrollTo(0, 0);
        }}
    </script>
</body>
</html>"""
    
    return html


# ==========================================
# HELPER FOR EXECUTIVE SUMMARY (using python-docx)
# ==========================================

def infer_audience_from_description(target_audience: str, genai_client=None) -> dict:
    """
    Infer audience metadata using AI to determine strategic vs operational focus,
    detail level, and emphasis areas from free-text description.
    
    Args:
        target_audience: Free-text description of target audience
        genai_client: Google Generative AI client for intelligent inference
        
    Returns:
        dict with keys: strategic_focus, operational_focus, detail_level, emphasis_areas, tone, priorities
    """
    audience_lower = (target_audience or "").lower().strip()
    
    # Default metadata
    metadata = {
        'strategic_focus': False,
        'operational_focus': True,
        'detail_level': 'moderate',
        'emphasis_areas': [],
        'tone': 'professional',
        'priorities': []
    }
    
    if not audience_lower:
        return metadata
    
    # Use AI for intelligent audience analysis if available
    if genai_client:
        try:
            inference_prompt = f"""Analyze this target audience description and return a JSON object:

TARGET AUDIENCE: "{target_audience}"

Determine:
1. strategic_focus (boolean): True if audience is executive/leadership (C-suite, board, directors, chairs, administrators)
2. operational_focus (boolean): True if audience is clinical/operational (physicians, nurses, staff, managers)
3. detail_level (string): "executive" for high-level strategic view, "detailed" for clinical specifics, "moderate" for balanced
4. tone (string): "executive" for formal business language, "clinical" for medical terminology, "accessible" for general professional
5. priorities (array of 3 strings): Top 3 things this audience cares about most (e.g., "ROI and cost savings", "patient safety outcomes", "workflow efficiency", "regulatory compliance", "staff training", "implementation timeline")
6. emphasis_areas (array of strings): Key topics to emphasize for this audience

Return ONLY valid JSON, no other text."""
            
            response = genai_client.models.generate_content(
                model="gemini-2.0-flash",
                contents=inference_prompt
            )
            json_match = re.search(r'\{.*\}', response.text, re.DOTALL)
            if json_match:
                ai_metadata = json.loads(json_match.group())
                metadata.update(ai_metadata)
                return metadata
        except Exception:
            pass
    
    # Fallback to keyword matching if AI unavailable
    executive_keywords = ['executive', 'c-suite', 'chief', 'director', 'chair', 'administrator', 'board', 'leadership', 'strategic', 'cfo', 'coo', 'ceo', 'vp', 'vice president']
    clinical_keywords = ['physician', 'doctor', 'nurse', 'rn', 'clinician', 'resident', 'fellow', 'staff', 'provider', 'practitioner', 'team', 'manager']
    
    is_executive = any(kw in audience_lower for kw in executive_keywords)
    is_clinical = any(kw in audience_lower for kw in clinical_keywords)
    
    if is_executive:
        metadata['strategic_focus'] = True
        metadata['operational_focus'] = False
        metadata['detail_level'] = 'executive'
        metadata['tone'] = 'executive'
        metadata['emphasis_areas'] = ['strategic impact', 'financial value', 'organizational outcomes']
        metadata['priorities'] = ['ROI and resource optimization', 'Quality and safety metrics', 'Strategic alignment']
    elif is_clinical:
        metadata['operational_focus'] = True
        metadata['strategic_focus'] = False
        metadata['detail_level'] = 'detailed'
        metadata['tone'] = 'clinical'
        metadata['emphasis_areas'] = ['clinical protocols', 'workflow integration', 'decision support']
        metadata['priorities'] = ['Patient outcomes', 'Workflow efficiency', 'Evidence-based practice']
    else:
        metadata['tone'] = 'professional'
        metadata['emphasis_areas'] = ['implementation readiness', 'stakeholder engagement']
        metadata['priorities'] = ['Project success', 'Stakeholder alignment', 'Measurable outcomes']
    
    return metadata


def create_phase5_executive_summary_docx(data: dict, condition: str, target_audience: str = "", genai_client=None):
    """
    Create a Word document executive summary for Phase 5 with audience-adaptive content.
    Requires python-docx to be installed.
    
    Args:
        data: Session data with phase1, phase2, phase3 info
        condition: Clinical condition name
        target_audience: Free-text target audience description for LLM-based inference
        genai_client: Optional Google Generative AI client for audience inference
        
    Returns:
        BytesIO buffer with .docx content, or None if python-docx unavailable
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()

        # Pull structured data once for clarity
        p1_data = data.get('phase1', {})
        p2_data = data.get('phase2', {})
        p3_data = data.get('phase3', {})
        p4_data = data.get('phase4', {})
        p5_data = data.get('phase5', {})
        
        # Infer audience focus from target audience text
        audience_metadata = infer_audience_from_description(target_audience, genai_client)
        is_executive = audience_metadata.get('strategic_focus', False)
        is_operational = audience_metadata.get('operational_focus', False)
        detail_level = audience_metadata.get('detail_level', 'moderate')
        emphasis_areas = audience_metadata.get('emphasis_areas', [])

        # Extract Phase 1 data early for SMART objectives
        setting_text = p1_data.get('setting', '')
        population = p1_data.get('population', 'N/A') or 'N/A'
        problem_text = p1_data.get('problem', 'Not provided')
        objectives_text = p1_data.get('objectives', 'Not provided')

        # Build SMART objectives scaffold from Phase 1 objectives text
        smart_source = objectives_text if objectives_text and objectives_text != 'Not provided' else "Clarify objectives in Phase 1 to align with SMART criteria."
        smart_objectives = [
            f"Specific: {smart_source}",
            "Measurable: Define baseline and target metrics (safety events, LOS, throughput, readmissions).",
            "Achievable: Resource plan and staffing validated through stakeholder review.",
            "Relevant: Aligned to organizational quality, safety, and access priorities.",
            "Time-bound: Go-live timeline with 30/60/90 day checkpoints and quarterly reviews."
        ]

        # Phase status snapshot for quick exec visibility
        phase_inputs = [
            "Phase 1 Charter captured" if p1_data else "Phase 1 Charter pending",
            f"Phase 2 Evidence items: {len(p2_data.get('evidence', []))}" if p2_data.get('evidence') else "Phase 2 Evidence pending",
            f"Phase 3 Pathway nodes: {len(p3_data.get('nodes', []))}" if p3_data.get('nodes') else "Phase 3 Pathway design pending",
            "Phase 4 Usability: completed" if p4_data.get('heuristics_data') else "Phase 4 Usability review pending"
        ]

        # Title reflects clinical condition and care setting
        setting_suffix = f" - {setting_text}" if setting_text else ""
        title = doc.add_heading(f"Executive Summary: {condition}{setting_suffix}", 0)
        title_format = title.paragraph_format
        title_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # AUDIENCE-ADAPTIVE CONTENT STRUCTURE
        # Note: Target audience is used to TAILOR content, not displayed in the document
        
        # Get audience priorities for tailoring
        priorities = audience_metadata.get('priorities', [])
        tone = audience_metadata.get('tone', 'professional')
        
        # Generate AI-tailored content sections if genai_client available
        ai_content = {}
        if genai_client:
            try:
                evidence = p2_data.get('evidence', [])
                nodes = p3_data.get('nodes', [])
                
                # Build context for AI
                evidence_summary = f"{len(evidence)} evidence items" if evidence else "evidence pending"
                nodes_summary = f"{len(nodes)} pathway nodes" if nodes else "pathway design pending"
                
                # Generate tailored executive summary content
                summary_prompt = f"""Generate a professional executive summary for a clinical pathway project.

PROJECT CONTEXT:
- Clinical Condition: {condition}
- Care Setting: {setting_text or 'healthcare setting'}
- Problem Statement: {problem_text}
- Project Objectives: {objectives_text}
- Evidence Base: {evidence_summary}
- Pathway Structure: {nodes_summary}
- Phase 4 Usability: {'completed' if p4_data.get('heuristics_data') else 'pending'}
- Expert Panel: {'completed' if p5_data.get('expert_html') else 'pending'}
- Beta Testing: {'completed' if p5_data.get('beta_html') else 'pending'}

TARGET AUDIENCE PRIORITIES: {', '.join(priorities) if priorities else 'general stakeholders'}
TONE: {tone}

Generate a JSON object with these sections (each section should be 2-3 professional paragraphs):

{{
  "executive_overview": "Opening paragraph establishing the pathway's strategic value and scope...",
  "strategic_rationale": "Why this pathway matters - problem, opportunity, organizational alignment...",
  "evidence_foundation": "Summary of evidence quality and clinical rigor supporting the pathway...",
  "value_proposition": "Expected outcomes - quality, safety, efficiency, financial impact...",
  "implementation_readiness": "Current status, validation completed, stakeholder engagement...",
  "success_metrics": "How success will be measured - specific metrics aligned to objectives...",
  "recommended_actions": "Clear asks for leadership - approvals, resources, next steps..."
}}

REQUIREMENTS:
- Write in formal, {tone} language appropriate for the target audience
- Emphasize: {', '.join(priorities) if priorities else 'project success and patient outcomes'}
- Be specific to {condition} management, not generic healthcare language
- Each section should be substantive (2-3 sentences minimum)
- Do NOT include audience name or "prepared for" statements

Return ONLY valid JSON."""
                
                response = genai_client.models.generate_content(
                    model="gemini-2.0-flash",
                    contents=summary_prompt
                )
                json_match = re.search(r'\{.*\}', response.text, re.DOTALL)
                if json_match:
                    ai_content = json.loads(json_match.group())
            except Exception:
                pass
        
        if is_executive:
            # === EXECUTIVE SUMMARY (C-SUITE FOCUS) - NARRATIVE PROSE ===
            doc.add_heading("Executive Overview", level=1)
            
            evidence = p2_data.get('evidence', [])
            nodes = p3_data.get('nodes', [])
            
            # Use AI-generated content if available, otherwise use template
            if ai_content.get('executive_overview'):
                doc.add_paragraph(ai_content['executive_overview'])
            else:
                doc.add_paragraph(
                    f"This executive summary presents a comprehensive clinical decision pathway for {condition} management "
                    f"in {setting_text or 'the clinical setting'}. The pathway represents a strategic initiative to standardize "
                    f"care delivery, improve patient outcomes, and optimize resource utilization through evidence-based clinical "
                    f"decision-making. Built on a foundation of {len(evidence)} peer-reviewed evidence sources and structured "
                    f"across {len(nodes)} decision points, this pathway addresses a critical opportunity to reduce practice variation "
                    f"while enhancing both clinical quality and operational efficiency."
                )
            
            doc.add_heading("Strategic Rationale", level=1)
            if ai_content.get('strategic_rationale'):
                doc.add_paragraph(ai_content['strategic_rationale'])
            else:
                doc.add_paragraph(
                    f"The current state presents significant challenges: {problem_text} This variability not only compromises "
                    f"patient safety and outcomes but also leads to inefficient resource allocation, extended length of stay, "
                    f"and increased operational costs. The proposed pathway establishes a systematic, evidence-driven approach "
                    f"that aligns clinical practice with organizational strategic priorities for quality, safety, and value-based care."
                )
            
            doc.add_heading("Evidence Foundation & Clinical Rigor", level=1)
            if ai_content.get('evidence_foundation'):
                doc.add_paragraph(ai_content['evidence_foundation'])
            elif evidence:
                grades = {}
                for e in evidence:
                    grade = e.get('grade', 'Un-graded')
                    grades[grade] = grades.get(grade, 0) + 1
                high_quality = sum(count for grade, count in grades.items() if grade in ['A', 'High', 'Level 1', 'High (A)'])
                
                doc.add_paragraph(
                    f"The pathway integrates current best evidence from {len(evidence)} peer-reviewed publications, with "
                    f"{high_quality} high-quality systematic reviews and clinical trials providing Level A evidence. This rigorous "
                    f"evidence base ensures that clinical recommendations reflect the most current understanding of optimal "
                    f"{condition} management, reducing liability exposure while supporting quality reporting and accreditation standards."
                )
            else:
                doc.add_paragraph(
                    "The pathway design incorporates current clinical guidelines and evidence-based practices, with structured "
                    "opportunities for expert review and validation prior to implementation."
                )
            
            doc.add_heading("Value Proposition & Expected Outcomes", level=1)
            if ai_content.get('value_proposition'):
                doc.add_paragraph(ai_content['value_proposition'])
            else:
                doc.add_paragraph(
                    f"Implementation of this standardized pathway is projected to deliver measurable value across multiple dimensions. "
                    f"From a quality perspective, the pathway promotes consistent application of evidence-based interventions, reducing "
                    f"diagnostic errors and ensuring appropriate escalation of care when indicated. Operationally, standardization "
                    f"enables more predictable resource utilization, reducing unnecessary testing and procedures while maintaining or "
                    f"improving clinical outcomes. The pathway also supports workforce development through clear protocols that enhance "
                    f"staff confidence and competency, particularly valuable for training programs and onboarding new providers."
                )
            
            doc.add_heading("Implementation Readiness", level=1)
            if ai_content.get('implementation_readiness'):
                doc.add_paragraph(ai_content['implementation_readiness'])
            else:
                usability_status = "completed rigorous usability testing" if p4_data.get('heuristics_data') else "planned comprehensive usability evaluation"
                expert_status = "completed expert panel review with integrated feedback" if p5_data.get('expert_html') else "scheduled expert panel validation"
                beta_status = "completed beta testing in the target environment" if p5_data.get('beta_html') else "planned pilot testing"
                
                doc.add_paragraph(
                    f"The pathway has {usability_status}, ensuring alignment with clinical workflow and practitioner needs. "
                    f"Stakeholder engagement includes {expert_status} and {beta_status}, creating a robust validation process "
                    f"that mitigates implementation risk. A comprehensive education module has been developed to support staff "
                    f"onboarding, competency assessment, and ongoing training, ensuring smooth adoption and sustained adherence."
                )
            
            doc.add_heading("Strategic Objectives & Success Metrics", level=1)
            if ai_content.get('success_metrics'):
                doc.add_paragraph(ai_content['success_metrics'])
            else:
                doc.add_paragraph(
                    f"The pathway is designed to achieve specific, measurable objectives aligned with organizational priorities: "
                    f"{objectives_text} Success will be monitored through defined metrics including clinical outcomes (safety events, "
                    f"readmission rates, diagnostic accuracy), operational efficiency (length of stay, throughput, resource utilization), "
                    f"and quality measures (adherence rates, patient satisfaction, staff competency). Implementation will proceed in "
                    f"phases with 30/60/90 day checkpoints and quarterly performance reviews to ensure sustained improvement and "
                    f"continuous optimization."
                )
            
            doc.add_heading("Recommended Actions", level=1)
            if ai_content.get('recommended_actions'):
                doc.add_paragraph(ai_content['recommended_actions'])
            else:
                doc.add_paragraph(
                    f"To advance this initiative, leadership approval is requested for: (1) final pathway validation and sign-off following "
                    f"expert review and beta testing; (2) allocation of implementation resources including staff education time, system "
                    f"integration support, and monitoring infrastructure; (3) establishment of a governance structure with defined "
                    f"accountability for pathway oversight, performance tracking, and continuous improvement; and (4) authorization to "
                    f"proceed with go-live planning, including communication strategy, training rollout, and performance dashboard "
                    f"deployment. These actions position the organization to deliver evidence-based, high-value care that advances "
                    f"both clinical excellence and operational sustainability."
                )
            
        else:
            # === OPERATIONAL/CLINICAL FOCUS ===
            doc.add_heading("Project Scope", level=1)
            doc.add_paragraph(
                f"Clinical pathway for {condition or 'N/A'} in {setting_text or 'care setting not specified'}, focused on {population}."
            )
            doc.add_paragraph(f"Problem statement: {problem_text}", style='List Bullet')
            inclusion = p1_data.get('inclusion', '')
            exclusion = p1_data.get('exclusion', '')
            if inclusion:
                doc.add_paragraph(f"Inclusion criteria: {inclusion}", style='List Bullet 2')
            if exclusion:
                doc.add_paragraph(f"Exclusion criteria: {exclusion}", style='List Bullet 2')

            doc.add_heading("Project Goals & Success Measures", level=1)
            doc.add_paragraph(objectives_text)
            doc.add_paragraph(
                "Outcome focus: safer, faster care delivery with clear resource stewardship in the specified care setting.",
                style='List Bullet'
            )

            doc.add_heading("SMART Objectives", level=1)
            for item in smart_objectives:
                doc.add_paragraph(item, style='List Bullet')

            doc.add_heading("Phase Inputs Synthesized (Phases 1-4)", level=1)
            for item in phase_inputs:
                doc.add_paragraph(item, style='List Bullet')
            
            # Evidence Summary (operational - detail level determines depth)
            doc.add_heading("Evidence Summary", level=1)
            evidence = p2_data.get('evidence', [])
            if evidence:
                doc.add_paragraph(f"Total evidence items reviewed: {len(evidence)}")
                
                # Grade breakdown
                grades = {}
                for e in evidence:
                    grade = e.get('grade', 'Un-graded')
                    grades[grade] = grades.get(grade, 0) + 1
                
                doc.add_paragraph("Evidence Quality Distribution:", style='List Bullet')
                for grade, count in sorted(grades.items()):
                    doc.add_paragraph(f"{grade}: {count} articles", style='List Bullet 2')
            else:
                doc.add_paragraph("No evidence reviewed yet")
            
            # Pathway Overview
            doc.add_heading("Pathway Design", level=1)
            nodes = p3_data.get('nodes', [])
            
            if nodes:
                doc.add_paragraph(f"Total pathway nodes: {len(nodes)}")
                
                node_types = {}
                for node in nodes:
                    node_type = node.get('type', 'Process')
                    node_types[node_type] = node_types.get(node_type, 0) + 1
                
                doc.add_paragraph("Pathway Node Distribution:", style='List Bullet')
                for node_type, count in sorted(node_types.items()):
                    doc.add_paragraph(f"{node_type}: {count}", style='List Bullet 2')
                
                # Add node list (always for operational summaries)
                if detail_level == "detailed":
                    doc.add_heading("Pathway Nodes", level=2)
                    table = doc.add_table(rows=1, cols=3)
                    table.style = 'Table Grid'
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = 'ID'
                    hdr_cells[1].text = 'Node'
                    hdr_cells[2].text = 'Type'
                    
                    for idx, node in enumerate(nodes):
                        row_cells = table.add_row().cells
                        row_cells[0].text = f"N{idx + 1}"
                        row_cells[1].text = node.get('label', 'N/A')
                        row_cells[2].text = node.get('type', 'N/A')
            else:
                doc.add_paragraph("Pathway not yet designed")
            
            # Usability Assessment
            doc.add_heading("Usability Assessment", level=1)
            heuristics = p4_data.get('heuristics_data', {})
            
            if heuristics:
                doc.add_paragraph("Nielsen Heuristics Review Completed", style='List Bullet')
                doc.add_paragraph(f"Total heuristics evaluated: {len(heuristics)}", style='List Bullet 2')
            else:
                doc.add_paragraph("Usability assessment pending")

            # Co-design, validation, and education
            doc.add_heading("Co-Design, Testing, and Education", level=1)
            expert_status = (
                "Expert panel feedback captured via structured review and integrated into the pathway." if p5_data.get('expert_html')
                else "Expert panel review planned using structured feedback form; integrate findings into the pathway."
            )
            beta_status = (
                "Beta testing completed with target users; updates folded into the current pathway version." if p5_data.get('beta_html')
                else "Beta testing planned in the target setting; incorporate usability findings into the pathway."
            )
            edu_status = (
                "Custom interactive education module prepared for staff onboarding (HTML, offline-capable)." if p5_data.get('edu_html')
                else "Interactive education module planned to support staff onboarding and competency validation."
            )
            doc.add_paragraph(expert_status, style='List Bullet')
            doc.add_paragraph(beta_status, style='List Bullet')
            doc.add_paragraph(edu_status, style='List Bullet')
            
            # Implementation Next Steps
            doc.add_heading("Implementation Next Steps", level=1)
            doc.add_paragraph(
                "Finalize and sign off the pathway after incorporating expert and beta feedback for the specified care setting.",
                style='List Number'
            )
            doc.add_paragraph(
                "Deploy the interactive education module; track staff completion and competency.",
                style='List Number'
            )
            doc.add_paragraph(
                "Go-live with monitoring for safety, throughput, and resource stewardship in the care setting.",
                style='List Number'
            )
            doc.add_paragraph(
                "Report performance to leadership and iterate monthly based on outcomes and user feedback.",
                style='List Number'
            )
            doc.add_paragraph(
                "Maintain continuous improvement cycles with refreshed education content as the pathway evolves.",
                style='List Number'
            )
        
        # Footer
        section = doc.sections[0]
        footer = section.footer
        footer_paragraph = footer.paragraphs[0]
        footer_paragraph.text = f"CarePathIQ Executive Summary | {condition} | Generated {datetime.now().strftime('%Y-%m-%d')}"
        footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Save to buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
        
    except ImportError:
        return None


# ==========================================
# UTILITY FUNCTIONS
# ==========================================

def ensure_carepathiq_branding(html: str) -> str:
    """Ensure HTML has CarePathIQ branding and footer."""
    if CAREPATHIQ_FOOTER not in html:
        if "</body>" in html:
            html = html.replace("</body>", CAREPATHIQ_FOOTER + "</body>")
        else:
            html += CAREPATHIQ_FOOTER
    return html


# ==========================================
# ROLE-SPECIFIC EDUCATION MODULE FUNCTIONS
# ==========================================

def get_role_depth_mapping(target_audience: str) -> dict:
    """
    Map clinical roles to content depth, node types, and expectations.
    Returns configuration for role-specific education module customization.
    
    Args:
        target_audience: Free-text role description (e.g., "Emergency Medicine Residents")
    
    Returns:
        Dict with:
        - 'role_type': Simplified role (Resident, Attending, Nurse, APP, Student, etc.)
        - 'depth_level': 'deep', 'moderate', or 'focused'
        - 'node_types': List of pathway node types to include
        - 'role_statement': Explicit role statement for headers/content
        - 'expectations': Key expectations for this role in the pathway
    """
    audience_lower = (target_audience or "").lower()
    
    # Resident/Fellow/Trainee (deep clinical reasoning)
    if any(word in audience_lower for word in ['resident', 'fellow', 'trainee', 'junior', 'pgy']):
        return {
            'role_type': 'Resident',
            'depth_level': 'deep',
            'node_types': ['Start', 'Decision', 'Process', 'Action', 'End'],
            'role_statement': f'As a {target_audience.strip()}, you are responsible for clinical decision-making and pathway execution.',
            'expectations': [
                'Understand the clinical reasoning behind each decision point',
                'Know when to escalate to attending staff',
                'Apply evidence-based protocols independently',
                'Recognize complications and deviations from pathway'
            ]
        }
    
    # Attending/Senior Physician (oversight, decision validation)
    elif any(word in audience_lower for word in ['attending', 'physician', 'doctor', 'consultant', 'provider', 'senior']):
        return {
            'role_type': 'Attending',
            'depth_level': 'deep',
            'node_types': ['Start', 'Decision', 'Process', 'End'],
            'role_statement': f'As a {target_audience.strip()}, you provide clinical oversight and validate key care decisions.',
            'expectations': [
                'Supervise resident implementation of pathway',
                'Make final decisions at critical junctures',
                'Ensure adherence to evidence-based protocols',
                'Address complex or unusual presentations'
            ]
        }
    
    # Nurse/RN (assessment, monitoring, communication)
    elif any(word in audience_lower for word in ['nurse', 'rn', 'lpn', 'nursing', 'cna']):
        return {
            'role_type': 'Nurse',
            'depth_level': 'moderate',
            'node_types': ['Start', 'Process', 'Action', 'End'],
            'role_statement': f'As a {target_audience.strip()}, you execute clinical assessments, monitoring, and care coordination.',
            'expectations': [
                'Perform initial and ongoing patient assessment',
                'Monitor for changes requiring escalation',
                'Execute orders and document findings',
                'Communicate status changes to medical team'
            ]
        }
    
    # APP/NP/PA (independent to semi-independent decision-making)
    elif any(word in audience_lower for word in ['app', 'np', 'nurse practitioner', 'pa', 'physician assistant', 'practitioner']):
        return {
            'role_type': 'APP',
            'depth_level': 'deep',
            'node_types': ['Start', 'Decision', 'Process', 'Action', 'End'],
            'role_statement': f'As a {target_audience.strip()}, you make independent decisions within scope and collaborate with physicians.',
            'expectations': [
                'Make autonomous clinical decisions per scope of practice',
                'Know when to consult physician colleagues',
                'Order and interpret diagnostic tests',
                'Implement therapeutic interventions'
            ]
        }
    
    # Student/Learner (foundational understanding)
    elif any(word in audience_lower for word in ['student', 'learner', 'medical student', 'nursing student', 'clerk']):
        return {
            'role_type': 'Student',
            'depth_level': 'focused',
            'node_types': ['Start', 'Decision', 'Process', 'End'],
            'role_statement': f'As a {target_audience.strip()}, you are learning the clinical pathway and evidence-based practice.',
            'expectations': [
                'Understand clinical presentation and initial assessment',
                'Learn how clinical decisions are made',
                'Recognize key decision points in care',
                'Ask questions and build clinical knowledge'
            ]
        }
    
    # Allied Health (specific clinical roles)
    elif any(word in audience_lower for word in ['therapist', 'tech', 'assistant', 'specialist', 'coordinator']):
        return {
            'role_type': 'Allied Health',
            'depth_level': 'focused',
            'node_types': ['Process', 'Action', 'End'],
            'role_statement': f'As a {target_audience.strip()}, you support core clinical care through specialized skills.',
            'expectations': [
                'Understand your role in the pathway',
                'Know when your intervention is needed',
                'Communicate findings to clinical team',
                'Document activities accurately'
            ]
        }
    
    # Default/unknown
    else:
        return {
            'role_type': 'Clinical Team Member',
            'depth_level': 'moderate',
            'node_types': ['Start', 'Decision', 'Process', 'Action', 'End'],
            'role_statement': f'As a {target_audience.strip()}, you are part of the clinical team implementing this evidence-based pathway.',
            'expectations': [
                'Understand the pathway and your role within it',
                'Follow evidence-based protocols',
                'Communicate effectively with team members',
                'Contribute to quality improvement'
            ]
        }


def generate_role_specific_module_header(
    target_audience: str,
    condition: str,
    care_setting: str,
    node: dict = None
) -> str:
    """
    Generate role-specific module header with explicit role clarity.
    
    Args:
        target_audience: Target learner role
        condition: Clinical condition
        care_setting: Care environment (e.g., Emergency Department)
        node: Optional pathway node for context
    
    Returns:
        String with role-specific header statement
    """
    role_mapping = get_role_depth_mapping(target_audience)
    role_statement = role_mapping['role_statement']
    
    if node:
        node_label = node.get('label', 'Clinical Decision Point')
        return f"{role_statement} In this module, you will learn your specific role when managing: {node_label}"
    else:
        return f"{role_statement} This module covers the {condition} pathway in {care_setting or 'your clinical setting'}."


def generate_role_specific_learning_objectives(
    target_audience: str,
    condition: str,
    nodes: list = None,
    module_idx: int = 0
) -> list:
    """
    Generate role-specific learning objectives based on audience and pathway nodes.
    
    Args:
        target_audience: Target learner role
        condition: Clinical condition
        nodes: List of pathway nodes for this module
        module_idx: Module number (0-indexed)
    
    Returns:
        List of 3-4 role-specific learning objectives
    """
    role_mapping = get_role_depth_mapping(target_audience)
    role_type = role_mapping['role_type']
    expectations = role_mapping['expectations']
    
    objectives = []
    
    if nodes and len(nodes) > 0:
        # Extract key actions from nodes
        key_actions = [n.get('label', '') for n in nodes if n.get('type') in ['Decision', 'Process', 'Action'] and n.get('label')]
        
        if key_actions:
            first_action = key_actions[0]
            if role_type == 'Resident':
                objectives.append(f"Explain the clinical reasoning behind {first_action.lower()}")
                objectives.append(f"Determine when to escalate {condition} cases to senior staff")
                objectives.append(f"Implement {condition} pathway protocols independently")
            elif role_type == 'Attending':
                objectives.append(f"Supervise resident execution of {first_action.lower()}")
                objectives.append(f"Validate clinical decisions at critical junctures")
                objectives.append(f"Address complex {condition} presentations")
            elif role_type == 'Nurse':
                objectives.append(f"Perform assessment for {condition}")
                objectives.append(f"Recognize findings that require escalation")
                objectives.append(f"Execute care orders and monitor patient status")
            elif role_type == 'APP':
                objectives.append(f"Make autonomous decisions regarding {first_action.lower()}")
                objectives.append(f"Know when to consult physician colleagues")
                objectives.append(f"Order and interpret tests within scope of practice")
            else:
                objectives.append(f"Understand the approach to {condition}")
                objectives.append(f"Recognize your role in pathway implementation")
                objectives.append(f"Know when to escalate or ask for help")
    
    # Add expectation-based objectives
    if not objectives and expectations:
        objectives = expectations[:3]
    
    # Fallback
    if not objectives:
        objectives = [
            f"Understand evidence-based management of {condition}",
            f"Recognize key decision points in the pathway",
            "Implement pathway protocols appropriately"
        ]
    
    return objectives[:4]  # Return max 4 objectives


def generate_role_specific_quiz_scenario(
    question_idx: int,
    node: dict,
    target_audience: str,
    evidence_citations: list = None
) -> dict:
    """
    Generate a realistic, role-specific quiz scenario based on pathway node.
    
    Args:
        question_idx: Question number
        node: Pathway node (Decision or Process)
        target_audience: Target learner role
        evidence_citations: List of dicts with 'pmid', 'title', 'grade' from Phase 2
    
    Returns:
        Dict with 'question', 'options', 'correct', 'explanation', 'evidence'
    """
    role_mapping = get_role_depth_mapping(target_audience)
    role_type = role_mapping['role_type']
    node_label = node.get('label', 'Clinical decision point')
    node_evidence = node.get('evidence', '')
    
    # Build role-specific scenario framing
    if role_type == 'Resident':
        scenario_stem = f"You are the resident caring for a patient with findings suggesting {node_label.lower()}. "
    elif role_type == 'Attending':
        scenario_stem = f"The resident presents a case where {node_label.lower()} is being considered. You need to "
    elif role_type == 'Nurse':
        scenario_stem = f"During your assessment, you observe {node_label.lower()}. "
    elif role_type == 'APP':
        scenario_stem = f"You are evaluating a patient and need to decide on {node_label.lower()}. "
    else:
        scenario_stem = f"When managing a patient with this presentation, you should consider {node_label.lower()}. "
    
    # Create role-appropriate answer options
    if role_type == 'Resident':
        options = [
            f"Follow the pathway protocol and prepare to discuss with attending",
            "Make a decision independently without consulting",
            "Wait for attending without taking any action",
            "Refer to another service immediately"
        ]
    elif role_type == 'Attending':
        options = [
            "Validate the pathway decision and provide oversight",
            "Overrule the resident without explanation",
            "Defer entirely to the resident without review",
            "Order additional tests that contradict pathway"
        ]
    elif role_type == 'Nurse':
        options = [
            "Alert the clinical team and monitor per protocol",
            "Document and wait without escalating",
            "Perform the intervention yourself",
            "Assume it will resolve without intervention"
        ]
    elif role_type == 'APP':
        options = [
            "Implement per scope of practice and document decision",
            "Always refer to physician regardless of scope",
            "Proceed without consulting any guidelines",
            "Delay care until physician is available"
        ]
    else:
        options = [
            f"Recognize {node_label.lower()} and alert appropriate staff",
            "Ignore the finding",
            "Handle independently without telling anyone",
            "Wait to see if it resolves"
        ]
    
    # Build evidence-backed explanation
    explanation = f"The pathway recommends {node_label.lower()}. "
    if node_evidence and node_evidence != 'N/A':
        explanation += f"(Evidence: {node_evidence}) "
    
    # Add evidence citation if available
    evidence_note = ""
    if evidence_citations:
        for cite in evidence_citations[:1]:  # Use first citation
            if cite.get('pmid'):
                grade = cite.get('grade', 'Grade unknown')
                evidence_note = f" [PMID: {cite.get('pmid')}, {grade}]"
                explanation += evidence_note
                break
    
    return {
        'question': scenario_stem + f"What should you do?",
        'options': options,
        'correct': 0,  # First option is always correct
        'explanation': explanation,
        'evidence': evidence_note
    }


def filter_nodes_by_role(nodes: list, target_audience: str) -> list:
    """
    Filter pathway nodes based on role depth and responsibilities.
    
    Args:
        nodes: Full list of pathway nodes
        target_audience: Target learner role
    
    Returns:
        Filtered list of nodes appropriate for the role
    """
    role_mapping = get_role_depth_mapping(target_audience)
    allowed_types = role_mapping['node_types']
    
    return [n for n in nodes if n.get('type') in allowed_types]
